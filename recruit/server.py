from dotenv import load_dotenv
from pathlib import Path

load_dotenv(Path(__file__).with_name(".env"))
load_dotenv()

from fastapi import FastAPI, APIRouter, BackgroundTasks, HTTPException, Request, Response, File, UploadFile, Form
from starlette.middleware.cors import CORSMiddleware
from starlette.middleware.gzip import GZipMiddleware
from supabase import create_client, Client
import inspect
try:
    from supabase import ClientOptions
except Exception:
    ClientOptions = None
import os
import logging
import jwt
import asyncio
import httpx
import resend
import io
import uuid
from datetime import datetime, timezone, timedelta
from typing import List, Optional
from pydantic import BaseModel, EmailStr, field_validator, model_validator
from enum import Enum
import re as _re
import json
import secrets
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ── PyMuPDF ─ resume masking ─────────────────────────────────
try:
    import fitz as _fitz
    _FITZ_OK = True
except ImportError:
    _fitz    = None
    _FITZ_OK = False

# ── Google Drive helpers ──────────────────────────────────────
try:
    from google_drive import upload_resume, delete_resume, download_resume, get_file_metadata, ALLOWED_MIME_TYPES, MAX_FILE_BYTES
    _GDRIVE_OK = True
except Exception as _gdrive_err:
    import logging as _l
    _l.getLogger(__name__).warning(f"[google_drive] import failed — resume upload disabled: {_gdrive_err}")
    upload_resume = delete_resume = download_resume = None
    _GDRIVE_OK = False
    ALLOWED_MIME_TYPES = {
        "application/pdf": "pdf",
        "application/msword": "doc",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    }
    MAX_FILE_BYTES = 10 * 1024 * 1024

# ── LLM helpers ──────────────────────────────────────────────
try:
    from llm_utils import extract_resume_insights, extract_jd_keywords
    LLM_AVAILABLE = True
except Exception as _llm_err:
    import logging as _l
    _l.getLogger(__name__).warning(f"[llm_utils] import failed — LLM disabled: {_llm_err}")
    LLM_AVAILABLE = False

# ── Logging ──────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# ── Supabase client ───────────────────────────────────────────
SUPABASE_URL: str = os.environ["SUPABASE_URL"]
SUPABASE_KEY: str = os.environ["SUPABASE_SERVICE_ROLE_KEY"]
if ClientOptions:
    _supabase_httpx = httpx.Client(http2=False)
    options = None
    try:
        init_sig = inspect.signature(ClientOptions.__init__)
        if "http_client" in init_sig.parameters:
            options = ClientOptions(http_client=_supabase_httpx)
        else:
            options = ClientOptions()
    except Exception:
        options = None
    if options is not None:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY, options=options)
    else:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
else:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# ── Resend email ──────────────────────────────────────────────
resend.api_key  = os.environ.get("RESEND_API_KEY", "")
SENDER_EMAIL    = os.environ.get("SENDER_EMAIL", "onboarding@resend.dev")

# ── JWT ───────────────────────────────────────────────────────
JWT_SECRET    = os.environ.get("JWT_SECRET", "change_me_in_production")
JWT_ALGORITHM = "HS256"

# ── App ───────────────────────────────────────────────────────
#
# NO APScheduler here — Recruit Space has no scheduled jobs.
# Scheduled tasks (timesheet reminders, digest emails) live in the Core Space only.
# This means Recruit can safely run multiple uvicorn workers in future if needed.
#
app        = FastAPI(title="Nexus CRM — Recruit & Onboarding Service")
api_router = APIRouter(prefix="/api")

# ── Rate limiter (protects /public/apply from spam bots) ──────
limiter = Limiter(key_func=get_remote_address, default_limits=[])
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ── Public resume folder (website applicants) ─────────────────
PUBLIC_RESUME_FOLDER_ID = os.environ.get(
    "PUBLIC_RESUME_FOLDER_ID", "1II9tu-fCUqAs63vWzoiamS_2YnXdE0g2"
)

logger.info("=== Nexus Recruit startup ===")
logger.info(f"  Google Drive : {'ok' if _GDRIVE_OK else 'DISABLED (check logs)'}")
logger.info(f"  LLM          : {'ok' if LLM_AVAILABLE else 'DISABLED (check logs)'}")


# ============================================================
# ENUMS
# ============================================================
class CandidateStatus(str, Enum):
    SOURCED              = "sourced"
    SCREENED             = "screened"
    SHORTLISTED          = "shortlisted"
    INTERVIEW_SCHEDULED  = "interview_scheduled"
    INTERVIEWED          = "interviewed"
    SELECTED             = "selected"
    REJECTED             = "rejected"
    ONBOARDED            = "onboarded"

class ActivityType(str, Enum):
    CALL          = "call"
    EMAIL         = "email"
    MEETING       = "meeting"
    NOTE          = "note"
    STATUS_CHANGE = "status_change"
    INTERVIEW     = "interview"


# ============================================================
# PYDANTIC MODELS (Jobs, Candidates, Interviews, Submissions)
# ============================================================
class JobCreate(BaseModel):
    title:             str
    department:        Optional[str]  = None
    location:          Optional[str]  = None
    employment_type:   Optional[str]  = None
    description:       Optional[str]  = None
    requirements:      Optional[str]  = None
    salary_range:      Optional[str]  = None
    skills:            List[str]      = []
    is_active:         bool           = True
    is_urgent:         bool           = False
    post_to_linkedin:  bool           = False   # trigger LinkedIn company page post

class JobUpdate(BaseModel):
    title:           Optional[str]  = None
    department:      Optional[str]  = None
    location:        Optional[str]  = None
    employment_type: Optional[str]  = None
    description:     Optional[str]  = None
    requirements:    Optional[str]  = None
    salary_range:    Optional[str]  = None
    skills:          Optional[List[str]] = None
    is_active:       Optional[bool] = None
    is_urgent:       Optional[bool] = None

class CandidateCreate(BaseModel):
    full_name:        str
    email:            Optional[str]  = None
    phone:            Optional[str]  = None
    current_company:  Optional[str]  = None
    candidate_role:   Optional[str]  = None
    experience_years: Optional[int]  = None
    source:           Optional[str]  = None
    job_id:           Optional[str]  = None
    status:           Optional[str]  = "sourced"
    notes:            Optional[str]  = None
    resume_url:       Optional[str]  = None
    linkedin_url:     Optional[str]  = None
    portfolio_url:    Optional[str]  = None
    skills:           List[str]      = []
    # Extended fields v3
    candidate_type:       Optional[str]  = "domestic"  # domestic | international
    visa_status:          Optional[str]  = None
    total_experience:     Optional[str]  = None
    relevant_experience:  Optional[str]  = None
    location:             Optional[str]  = None
    relocation:           Optional[str]  = None
    # v4 fields
    work_mode:            List[str]      = []   # e.g. ["Remote","Hybrid"]
    current_ctc:          Optional[float] = None  # in LPA / numeric
    expected_ctc:         Optional[float] = None
    # LLM-extracted fields (populated on resume upload)
    tech_stack:           List[str]      = []

    @field_validator("status", mode="before")
    @classmethod
    def coerce_cand_status(cls, v):
        valid = {"sourced","screened","shortlisted","interview_scheduled","interviewed","selected","rejected","onboarded"}
        if not v or str(v).strip() not in valid:
            return "sourced"
        return str(v).strip()

    @field_validator("email", mode="before")
    @classmethod
    def clean_cand_email(cls, v):
        if not v or not str(v).strip(): return None
        val = str(v).strip()
        if "@" not in val or "." not in val.split("@")[-1]: return None
        return val

class CandidateUpdate(BaseModel):
    full_name:            Optional[str]          = None
    email:                Optional[str]          = None
    phone:                Optional[str]          = None
    current_company:      Optional[str]          = None
    candidate_role:       Optional[str]          = None
    experience_years:     Optional[int]          = None
    source:               Optional[str]          = None
    job_id:               Optional[str]          = None
    status:               Optional[str]          = None
    notes:                Optional[str]          = None
    resume_url:           Optional[str]          = None
    linkedin_url:         Optional[str]          = None
    portfolio_url:        Optional[str]          = None
    skills:               Optional[List[str]]    = None
    assigned_recruiter_id:Optional[str]          = None
    # Extended fields v3
    candidate_type:       Optional[str]  = None
    visa_status:          Optional[str]  = None
    total_experience:     Optional[str]  = None
    relevant_experience:  Optional[str]  = None
    location:             Optional[str]  = None
    relocation:           Optional[str]  = None
    # v4 fields
    work_mode:            Optional[List[str]]  = None
    current_ctc:          Optional[float]      = None
    expected_ctc:         Optional[float]      = None
    # LLM-extracted fields
    tech_stack:           Optional[List[str]] = None

class SubmissionCreate(BaseModel):
    lead_id:      str
    candidate_id: str
    status:       str = "submitted"
    notes:        Optional[str] = None

class SubmissionUpdate(BaseModel):
    status: Optional[str] = None
    notes:  Optional[str] = None

class InterviewCreate(BaseModel):
    candidate_id:   str
    job_id:         str
    scheduled_at:   str           # ISO datetime
    interview_type: str
    interviewers:   List[str] = []
    notes:          Optional[str] = None

class InterviewUpdate(BaseModel):
    scheduled_at:   Optional[str]      = None
    interview_type: Optional[str]      = None
    interviewers:   Optional[List[str]]= None
    notes:          Optional[str]      = None
    feedback:       Optional[str]      = None
    rating:         Optional[int]      = None
    completed:      Optional[bool]     = None



# ============================================================
# HELPERS
# ============================================================


def create_access_token(user_id: str, email: str) -> str:
    return jwt.encode(
        {"sub": user_id, "email": email,
         "exp": datetime.now(timezone.utc) + timedelta(hours=24),
         "type": "access"},
        JWT_SECRET, algorithm=JWT_ALGORITHM
    )

def create_refresh_token(user_id: str) -> str:
    return jwt.encode(
        {"sub": user_id,
         "exp": datetime.now(timezone.utc) + timedelta(days=7),
         "type": "refresh"},
        JWT_SECRET, algorithm=JWT_ALGORITHM
    )

async def get_current_user(request: Request) -> dict:
    token = request.cookies.get("access_token")
    if not token:
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
    if not token:
        raise HTTPException(401, "Not authenticated")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        if payload.get("type") != "access":
            raise HTTPException(401, "Invalid token type")
    except jwt.ExpiredSignatureError:
        raise HTTPException(401, "Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(401, "Invalid token")

    user = await safe_single(
        lambda: supabase.table("users")
            .select("id,email,name,role,avatar_url,created_at")
            .eq("id", payload["sub"])
            .single()
            .execute()
    )
    if not user:
        raise HTTPException(401, "User not found")
    return user

# ── Module-level role guard ───────────────────────────────────
# Call after get_current_user() on any sales or recruitment endpoint.
# Workers are timesheet-only; this blocks them from all other data.
_MODULE_ROLES: dict = {
    "sales":       {"admin", "sales", "viewer"},
    "recruitment": {"admin", "sales", "viewer"},
}

def _require_module(user: dict, module: str) -> None:
    """Raise 403 if user's role is not permitted to access this module."""
    allowed = _MODULE_ROLES.get(module, set())
    role = user.get("role", "")
    if role not in allowed:
        raise HTTPException(
            status_code=403,
            detail=f"Your role '{role}' does not have access to the {module} module.",
        )


async def send_email(to: str, subject: str, html: str):
    if not resend.api_key:
        logger.info(f"[email-sim] To: {to} | {subject}")
        return {"status": "simulated"}
    try:
        result = await asyncio.to_thread(
            resend.Emails.send, {"from": SENDER_EMAIL, "to": [to], "subject": subject, "html": html}
        )
        return {"status": "success", "id": result.get("id")}
    except Exception as e:
        logger.error(f"Email failed: {e}")
        return {"status": "failed", "error": str(e)}

def sb(table: str):
    """Shorthand for supabase.table()"""
    return supabase.table(table)

def get_supabase() -> Client:
    """Return the shared Supabase client used by route handlers."""
    return supabase

async def run(fn):
    """Run a synchronous supabase call in a thread pool."""
    return await asyncio.to_thread(fn)

async def safe_single(fn):
    """Run a .single() query safely — returns None on 0 rows (PGRST116) instead of crashing."""
    try:
        res = await asyncio.to_thread(fn)
        return res.data
    except Exception as e:
        err = str(e)
        if "PGRST116" in err or "0 rows" in err or "406" in err:
            return None
        raise



# ============================================================
# AUDIT LOG HELPERS
# ============================================================
def _get_ip(request: Request) -> str:
    xff = request.headers.get("x-forwarded-for", "")
    if xff:
        return xff.split(",")[0].strip()
    return request.client.host if request.client else "unknown"

async def _audit(
    action: str,
    user: dict = None,
    entity_type: str = None,
    entity_id:   str = None,
    entity_name: str = None,
    old_value:   dict = None,
    new_value:   dict = None,
    ip:          str = None,
    ua:          str = None,
):
    """Fire-and-forget audit log entry. Never raises."""
    try:
        await run(lambda: sb("audit_logs").insert({
            "user_id":     user["id"]    if user else None,
            "user_email":  user["email"] if user else None,
            "user_name":   user["name"]  if user else None,
            "action":      action,
            "entity_type": entity_type,
            "entity_id":   str(entity_id) if entity_id else None,
            "entity_name": entity_name,
            "old_value":   old_value,
            "new_value":   new_value,
            "ip_address":  ip,
            "user_agent":  ua,
        }).execute())
    except Exception as e:
        logger.warning(f"Audit log write failed: {e}")


# ============================================================
# AUTH
# ============================================================

# ============================================================
# JOBS
# ============================================================
@api_router.post("/jobs")
async def create_job(job: JobCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    apply_key = secrets.token_urlsafe(8)
    res = await run(lambda: sb("jobs").insert({
        "title":           job.title,
        "department":      job.department,
        "location":        job.location,
        "employment_type": job.employment_type,
        "description":     job.description,
        "requirements":    job.requirements,
        "salary_range":    job.salary_range,
        "skills":          job.skills,
        "is_active":       job.is_active,
        "is_urgent":       job.is_urgent,
        "created_by":      user["id"],
        # URL-safe random key — used by the public /apply?key=… form.
        # Generated here (not in the DB) so the value is visible in the response immediately.
        "apply_key":       apply_key,
    }).execute())
    created_job = res.data[0]

    # Optionally post to LinkedIn (non-blocking — job creation succeeds even if LinkedIn fails)
    linkedin_result = {"success": False, "error": "not_requested"}
    if job.post_to_linkedin:
        # Apply URL uses the frontend origin — stored as env var or derived from request
        frontend_origin = os.environ.get("FRONTEND_URL", "https://jmdata-crm-application.jmdatatalent.com")
        apply_url = f"{frontend_origin}/apply?key={apply_key}"
        linkedin_result = await _post_job_to_linkedin(created_job, apply_url)

    asyncio.create_task(_audit("create", user=user, entity_type="job", entity_id=created_job.get("id"),
                                entity_name=created_job.get("title"),
                                new_value={"title": created_job.get("title"), "department": created_job.get("department")},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {**created_job, "linkedin_post": linkedin_result}


@api_router.get("/jobs")
async def get_jobs(request: Request, is_active: Optional[bool] = None, search: Optional[str] = None):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    q = sb("jobs").select("*").order("created_at", desc=True)
    if is_active is not None: q = q.eq("is_active", is_active)
    if search: q = q.or_(f"title.ilike.%{search}%,department.ilike.%{search}%")
    res = await run(lambda: q.execute())
    jobs = res.data or []

    # Attach candidate count
    for job in jobs:
        cnt = await run(lambda jid=job["id"]: sb("candidates").select("id", count="exact").eq("job_id", jid).execute())
        job["candidate_count"] = cnt.count or 0
    return jobs


@api_router.get("/jobs/{job_id}")
async def get_job(job_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    job = await safe_single(lambda: sb("jobs").select("*").eq("id", job_id).single().execute())
    if not job:
        raise HTTPException(404, "Job not found")
    cands = await run(lambda: sb("candidates").select("*").eq("job_id", job_id).execute())
    job["candidates"] = cands.data or []
    return job


@api_router.put("/jobs/{job_id}")
async def update_job(job_id: str, job: JobUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    patch = {k: v for k, v in job.model_dump().items() if v is not None}
    old_job = await run(lambda: sb("jobs").select("title").eq("id", job_id).execute())
    job_name = (old_job.data or [{}])[0].get("title") or job_id
    await run(lambda: sb("jobs").update(patch).eq("id", job_id).execute())
    asyncio.create_task(_audit("update", user=user, entity_type="job", entity_id=job_id,
                                entity_name=job_name, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Job updated"}


@api_router.delete("/jobs/{job_id}")
async def delete_job(job_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    old_job = await run(lambda: sb("jobs").select("title").eq("id", job_id).execute())
    job_name = (old_job.data or [{}])[0].get("title") or job_id
    await run(lambda: sb("jobs").delete().eq("id", job_id).execute())
    asyncio.create_task(_audit("delete", user=user, entity_type="job", entity_id=job_id,
                                entity_name=job_name,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Job deleted"}


# ============================================================

# RESUME TEXT EXTRACTION HELPER
# ============================================================
def _extract_resume_text(file_bytes: bytes, content_type: str) -> str:
    """Extract plain text from PDF or DOCX bytes for LLM analysis."""
    try:
        if content_type == "application/pdf":
            import pdfplumber, io as _io
            with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                return "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                ).strip()
        elif content_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            import docx, io as _io
            doc = docx.Document(_io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        elif content_type == "application/msword":
            # Legacy .doc — best-effort UTF-8 decode
            return file_bytes.decode("utf-8", errors="ignore")
    except Exception as exc:
        logger.warning(f"[resume-extract] text extraction failed: {exc}")
    return ""


# ============================================================
# CANDIDATES
# ============================================================
@api_router.post("/candidates")
async def create_candidate(candidate: CandidateCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    doc_c = {
        "full_name":            candidate.full_name,
        "email":                candidate.email,
        "phone":                candidate.phone,
        "current_company":      candidate.current_company,
        "candidate_role":       candidate.candidate_role,
        "experience_years":     candidate.experience_years,
        "source":               candidate.source,
        "job_id":               candidate.job_id,
        "status":               candidate.status or "sourced",
        "notes":                candidate.notes,
        "resume_url":           candidate.resume_url,
        "linkedin_url":         candidate.linkedin_url,
        "portfolio_url":        candidate.portfolio_url,
        "skills":               candidate.skills,
        "assigned_recruiter_id":user["id"],
        "created_by":           user["id"],
        # Extended v3 fields
        "candidate_type":       candidate.candidate_type or "domestic",
        "visa_status":          candidate.visa_status,
        "total_experience":     candidate.total_experience,
        "relevant_experience":  candidate.relevant_experience,
        "location":             candidate.location,
        "relocation":           candidate.relocation,
        # v4 fields
        "work_mode":            candidate.work_mode or [],
        "current_ctc":          candidate.current_ctc,
        "expected_ctc":         candidate.expected_ctc,
        # LLM-extracted fields
        "tech_stack":           candidate.tech_stack or [],
    }
    doc_c = {k: v for k, v in doc_c.items() if v is not None and v != [] }
    if candidate.skills: doc_c["skills"] = candidate.skills
    if candidate.tech_stack is not None: doc_c["tech_stack"] = candidate.tech_stack
    try:
        res = await run(lambda: sb("candidates").insert(doc_c).execute())
    except Exception as e:
        err_str = str(e)
        if "PGRST204" in err_str:
            import re as _re
            col_match = _re.search(r"Could not find the '(\w+)' column", err_str)
            if col_match:
                bad_col = col_match.group(1)
                logger.warning(f"[create_candidate] Column '{bad_col}' missing — skipping. Run add_features_v3.sql.")
                doc_c.pop(bad_col, None)
                res = await run(lambda: sb("candidates").insert(doc_c).execute())
            else:
                raise
        else:
            raise
    candidate_id = res.data[0]["id"]
    await _log_activity(candidate_id=candidate_id, user=user, atype="note",
                        desc=f"Candidate added by {user['name']}")
    asyncio.create_task(_audit("create", user=user, entity_type="candidate", entity_id=candidate_id,
                                entity_name=res.data[0].get("full_name"),
                                new_value={"role": res.data[0].get("candidate_role"), "source": res.data[0].get("source"), "status": res.data[0].get("status")},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return res.data[0]


@api_router.get("/candidates")
async def get_candidates(
    request:      Request,
    job_id:       Optional[str] = None,
    status:       Optional[str] = None,
    source:       Optional[str] = None,
    search:       Optional[str] = None,
    skip:         int = 0,
    limit:        int = 50,
):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    q = sb("candidates").select(
        "*",
        count="exact"
    ).order("created_at", desc=True).range(skip, skip + limit - 1)

    if job_id:  q = q.eq("job_id", job_id)
    if status:  q = q.eq("status", status)
    if source:  q = q.eq("source", source)
    if search:
        q = q.or_(
            f"full_name.ilike.%{search}%,"
            f"email.ilike.%{search}%,"
            f"current_company.ilike.%{search}%,"
            f"candidate_role.ilike.%{search}%"
        )

    res = await run(lambda: q.execute())
    return {"candidates": res.data or [], "total": res.count or 0}


@api_router.get("/candidates/pipeline")
async def get_pipeline(request: Request, job_id: Optional[str] = None):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    pipeline = {}
    for s in CandidateStatus:
        q = sb("candidates").select("*").eq("status", s.value)
        if job_id: q = q.eq("job_id", job_id)
        res = await run(lambda qq=q: qq.execute())
        pipeline[s.value] = res.data or []
    return pipeline


@api_router.get("/candidates/{candidate_id}")
async def get_candidate(candidate_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    candidate_data = await safe_single(lambda: sb("candidates").select("*").eq("id", candidate_id).single().execute())
    if not candidate_data:
        raise HTTPException(404, "Candidate not found")
    cand = candidate_data

    acts = await run(lambda: sb("activities").select("*").eq("candidate_id", candidate_id).order("created_at", desc=True).execute())
    cand["activities"] = acts.data or []

    ivs = await run(lambda: sb("interviews").select("*").eq("candidate_id", candidate_id).order("scheduled_at", desc=True).execute())
    cand["interviews"] = ivs.data or []

    hist = await run(lambda: sb("candidate_status_history").select("*").eq("candidate_id", candidate_id).order("created_at", desc=True).execute())
    cand["status_history"] = hist.data or []
    return cand


@api_router.put("/candidates/{candidate_id}")
async def update_candidate(candidate_id: str, candidate: CandidateUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    existing = await safe_single(lambda: sb("candidates").select("status").eq("id", candidate_id).single().execute())
    if not existing:
        raise HTTPException(404, "Candidate not found")

    patch = {k: v for k, v in candidate.model_dump().items() if v is not None}
    # Explicitly allow empty lists for array fields (model_dump filters None but keeps [])
    for arr_field in ("tech_stack", "work_mode", "skills"):
        val = candidate.model_dump().get(arr_field)
        if val is not None:   # includes []
            patch[arr_field] = val
    if "status" in patch and isinstance(patch["status"], CandidateStatus):
        old_status = existing["status"]
        new_status = patch["status"].value
        patch["status"] = new_status
        if old_status != new_status:
            await run(lambda: sb("candidate_status_history").insert({
                "candidate_id":   candidate_id,
                "old_status":     old_status,
                "new_status":     new_status,
                "changed_by":     user["id"],
                "changed_by_name":user["name"],
            }).execute())
            await _log_activity(candidate_id=candidate_id, user=user, atype="status_change",
                                 desc=f"Stage moved from {old_status} to {new_status}")

    old_cand = await run(lambda: sb("candidates").select("full_name").eq("id", candidate_id).execute())
    cand_name = (old_cand.data or [{}])[0].get("full_name") or candidate_id
    await run(lambda: sb("candidates").update(patch).eq("id", candidate_id).execute())
    asyncio.create_task(_audit("update", user=user, entity_type="candidate", entity_id=candidate_id,
                                entity_name=cand_name, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Candidate updated"}


# ── Resume: Masked download ────────────────────────────────────────────────────
_PHONE_RE = _re.compile(
    r'(\+?(?:91[\s\-]?)?(?:[6-9]\d{9})'                   # Indian mobile
    r'|\+?(?:1[\s\-]?)?(?:\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4})'  # US
    r'|\+\d{1,3}[\s\-]?\d{7,12})',                         # generic international
    _re.VERBOSE,
)
_EMAIL_RE = _re.compile(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}')
_LINKEDIN_RE = _re.compile(
    r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-_%\.]+/?',
    _re.IGNORECASE,
)

def _mask_text(text: str) -> str:
    text = _EMAIL_RE.sub('[email hidden]', text)
    text = _PHONE_RE.sub('[phone hidden]', text)
    text = _LINKEDIN_RE.sub('[linkedin hidden]', text)
    return text

def _mask_pdf_bytes(pdf_bytes: bytes) -> bytes:
    """Redact emails, phone numbers, and LinkedIn URLs from a PDF using PyMuPDF."""
    if not _FITZ_OK:
        raise RuntimeError("PyMuPDF not installed — cannot mask PDF")
    doc = _fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        for pattern in (_EMAIL_RE, _PHONE_RE, _LINKEDIN_RE):
            for match in pattern.finditer(page.get_text()):
                rects = page.search_for(match.group())
                for r in rects:
                    page.add_redact_annot(r, fill=(0, 0, 0))
        page.apply_redactions()
    return doc.tobytes()

def _mask_docx_bytes(docx_bytes: bytes) -> bytes:
    """Replace emails and phone numbers in a DOCX."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    def _replace_in_para(para):
        for run in para.runs:
            run.text = _mask_text(run.text)
    for para in doc.paragraphs:
        _replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@api_router.get("/candidates/{candidate_id}/resume/masked")
async def download_masked_resume(candidate_id: str, request: Request):
    """Download the candidate's resume with phone and email redacted (no AI)."""
    from fastapi.responses import Response as FResponse
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    row = await safe_single(lambda: sb("candidates")
        .select("id,full_name,resume_url").eq("id", candidate_id).single().execute())
    if not row:
        raise HTTPException(404, "Candidate not found")
    url = (row.get("resume_url") or "").strip()
    if not url:
        raise HTTPException(400, "This candidate has no resume on file")
    if download_resume is None:
        raise HTTPException(503, "Google Drive not configured")

    metadata = await run(lambda: get_file_metadata(url))
    mime_type = (metadata.get("mimeType") or "").strip().lower()
    file_name = (metadata.get("name") or "").strip().lower()
    raw = await run(lambda: download_resume(url))

    # Detect type from Drive metadata, not from the preview URL.
    if mime_type == "application/pdf" or file_name.endswith(".pdf"):
        try:
            masked = _mask_pdf_bytes(raw)
        except Exception:
            masked = raw   # return original if masking fails
        ctype = "application/pdf"
        ext   = "pdf"
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith(".docx"):
        try:
            masked = _mask_docx_bytes(raw)
        except Exception:
            masked = raw
        ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext   = "docx"
    elif mime_type == "application/msword" or file_name.endswith(".doc"):
        # Legacy .doc files cannot be safely redacted with python-docx.
        # Return the original binary with the correct MIME type so Word can open it.
        masked = raw
        ctype = "application/msword"
        ext   = "doc"
    else:
        # Unknown Drive MIME type. Preserve the original bytes rather than
        # rewriting them as DOCX, which would corrupt non-DOCX resumes.
        masked = raw
        ctype = mime_type or "application/octet-stream"
        ext   = "bin"

    safe_name = (row.get("full_name") or "candidate").replace(" ", "_")
    asyncio.create_task(_audit(
        "resume_download", user=user,
        entity_type="candidate", entity_id=candidate_id,
        entity_name=row.get("full_name"),
        new_value={"type": "masked_resume", "format": ext},
        ip=_get_ip(request), ua=request.headers.get("user-agent", ""),
    ))
    return FResponse(
        content=masked,
        media_type=ctype,
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_masked.{ext}"'},
    )



@api_router.delete("/candidates/{candidate_id}")
async def delete_candidate(candidate_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")

    # Fetch resume_url before deleting so we can clean up Drive
    row = await safe_single(
        lambda: sb("candidates").select("id,full_name,resume_url").eq("id", candidate_id).single().execute()
    )
    if not row:
        raise HTTPException(404, "Candidate not found.")

    # Delete from DB first
    await run(lambda: sb("candidates").delete().eq("id", candidate_id).execute())

    # Then remove resume from Google Drive (non-fatal if it fails)
    resume_url = (row or {}).get("resume_url") or ""
    if resume_url and "drive.google.com" in resume_url and delete_resume is not None:
        try:
            await run(lambda: delete_resume(resume_url))
            logger.info(f"[delete_candidate] Drive resume deleted for {candidate_id}")
        except Exception as exc:
            logger.warning(f"[delete_candidate] Drive delete failed for {candidate_id}: {exc}")

    asyncio.create_task(_audit(
        action="delete", user=user,
        entity_type="candidate", entity_id=candidate_id,
        entity_name=row.get("full_name") or candidate_id,
    ))

    return {"success": True, "message": "Candidate and resume deleted."}


@api_router.post("/candidates/{candidate_id}/resume")
async def upload_candidate_resume(
    candidate_id: str,
    request: Request,
    file: UploadFile = File(...),
):
    """
    Upload a resume (PDF / DOC / DOCX) to Google Drive and store the
    preview URL in the candidate's resume_url column in Supabase.
 
    The stored URL is a Google Drive /preview link so the frontend
    can embed it directly in an <iframe> without opening a new tab.
    """
    user = await get_current_user(request)
    _require_module(user, "recruitment")

    # ── Validate MIME type ───────────────────────────────────
    if file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            400,
            f"Unsupported file type '{file.content_type}'. "
            "Only PDF, DOC, and DOCX are accepted."
        )
 
    # ── Read bytes ──────────────────────────────────────────
    contents = await file.read()
    if len(contents) > MAX_FILE_BYTES:
        raise HTTPException(400, "File too large. Maximum allowed size is 10 MB.")
 
    # ── Fetch candidate (need name + existing resume_url) ────
    existing = await safe_single(
        lambda: sb("candidates")
        .select("id,full_name,resume_url")
        .eq("id", candidate_id)
        .single()
        .execute()
    )
    if not existing:
        raise HTTPException(404, "Candidate not found")
 
    # ── Delete old resume from Drive if one exists ───────────
    old_url = existing.get("resume_url") or ""
    if "drive.google.com" in old_url:
        await run(lambda: delete_resume(old_url))
 
    # ── Build a clean filename: Resume_John_Doe_abc12345.pdf ─
    ext       = ALLOWED_MIME_TYPES[file.content_type]
    safe_name = (existing.get("full_name") or "Candidate").replace(" ", "_")
    short_id  = candidate_id.replace("-", "")[:8]
    filename  = f"Resume_{safe_name}_{short_id}.{ext}"
 
    # ── Upload to Google Drive ───────────────────────────────
    if upload_resume is None:
        raise HTTPException(503, "Google Drive integration not configured. "
                                 "Ensure google_drive.py and service account credentials are present.")
    try:
        result = await run(
            lambda: upload_resume(contents, filename, file.content_type)
        )
    except RuntimeError as exc:
        raise HTTPException(503, str(exc))
    except Exception as exc:
        logger.exception("Unexpected Google Drive upload failure")
        raise HTTPException(500, "Unexpected Google Drive upload failure.")
 
    preview_url = result["preview_url"]

    # ── Extract text + call LLM for tech_stack / experience ──
    db_patch: dict = {"resume_url": preview_url}
    tech_stack_extracted: list = []
    experience_extracted = None

    if LLM_AVAILABLE:
        try:
            resume_text = _extract_resume_text(contents, file.content_type)
            if resume_text:
                insights = await extract_resume_insights(resume_text)
                tech_stack_extracted = insights.get("tech_stack") or []
                experience_extracted = insights.get("experience_years")
                if tech_stack_extracted:
                    db_patch["tech_stack"] = tech_stack_extracted
                if experience_extracted is not None:
                    db_patch["experience_years"] = experience_extracted
        except Exception as llm_exc:
            logger.warning(f"[LLM] Resume insights failed (non-fatal): {llm_exc}")

    # ── Persist URL + LLM fields in Supabase ─────────────────
    await run(
        lambda: sb("candidates")
        .update(db_patch)
        .eq("id", candidate_id)
        .execute()
    )

    await _log_activity(
        candidate_id=candidate_id,
        user=user,
        atype="note",
        desc=f"Resume uploaded: {filename}"
              + (f" | {len(tech_stack_extracted)} skills extracted" if tech_stack_extracted else ""),
    )

    asyncio.create_task(_audit("resume_upload", user=user, entity_type="candidate",
                                entity_id=candidate_id, entity_name=filename,
                                new_value={"filename": filename,
                                           "skills_extracted": len(tech_stack_extracted or []),
                                           "experience_years": experience_extracted},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {
        "resume_url":        preview_url,
        "view_url":          result["view_url"],
        "file_id":           result["file_id"],
        "filename":          filename,
        "tech_stack":        tech_stack_extracted,
        "experience_years":  experience_extracted,
    }
 
 
@api_router.delete("/candidates/{candidate_id}/resume")
async def delete_candidate_resume(candidate_id: str, request: Request):
    """
    Remove the candidate's resume from Google Drive and clear resume_url
    in Supabase.
    """
    user = await get_current_user(request)
    _require_module(user, "recruitment")
 
    existing = await safe_single(
        lambda: sb("candidates")
        .select("id,resume_url")
        .eq("id", candidate_id)
        .single()
        .execute()
    )
    if not existing:
        raise HTTPException(404, "Candidate not found")
 
    url = existing.get("resume_url") or ""
    if "drive.google.com" in url:
        await run(lambda: delete_resume(url))
 
    await run(
        lambda: sb("candidates")
        .update({"resume_url": None})
        .eq("id", candidate_id)
        .execute()
    )
 
    await _log_activity(
        candidate_id=candidate_id,
        user=user,
        atype="note",
        desc="Resume removed",
    )
 
    asyncio.create_task(_audit("resume_delete", user=user, entity_type="candidate",
                                entity_id=candidate_id,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Resume deleted successfully"}


# ============================================================
# ATS MATCH — JD → Top-10 candidates
# ============================================================
class ATSMatchRequest(BaseModel):
    jd_text:        str
    candidate_type: str = "domestic"   # domestic | international


@api_router.post("/candidates/ats-match")
async def ats_match(body: ATSMatchRequest, request: Request):
    """
    Given a job description, find the top-10 best-fit candidates using:
      1. LLM (Gemini Flash, 1 call total) — parse JD into required/nice skills
      2. Pure-code scoring — tech_stack column overlap, zero per-candidate LLM calls

    Scoring formula (0-100):
      required_match% * 75  +  nice_to_have_match% * 25
    matched_skills / missing_skills computed locally from set intersection.
    """
    user = await get_current_user(request)
    if user.get("role") not in ("admin", "sales", "viewer"):
        raise HTTPException(403, "ATS matching is not available for your role.")

    if not LLM_AVAILABLE:
        raise HTTPException(503, "LLM service not configured. Set LLM_PROVIDER and API key.")

    jd_text = (body.jd_text or "").strip()
    if len(jd_text) < 30:
        raise HTTPException(400, "Job description is too short. Please provide more detail.")

    # ── Step 1: 1 LLM call — parse JD into structured skills ─
    jd_meta  = await extract_jd_keywords(jd_text)
    required = set(s.lower() for s in jd_meta.get("required_skills")     or [])
    nice     = set(s.lower() for s in jd_meta.get("nice_to_have_skills") or [])

    # ── Step 2: fetch all candidates of this type ─────────────
    res = await run(
        lambda: sb("candidates")
        .select("id,full_name,candidate_role,experience_years,tech_stack,"
                "location,visa_status,candidate_type,total_experience,status")
        .eq("candidate_type", body.candidate_type)
        .execute()
    )
    all_candidates = res.data or []

    if not all_candidates:
        return {"matches": [], "jd_meta": jd_meta, "total_scanned": 0, "pre_filtered": 0}

    # ── Step 3: pure-code scoring (no LLM per candidate) ──────
    def _score_candidate(cand: dict) -> dict:
        cand_set = set(s.lower() for s in (cand.get("tech_stack") or []))

        # Set intersections for matched / missing
        req_hits  = cand_set & required
        nice_hits = cand_set & nice

        # ATS score: required skills worth 75 pts, nice-to-have worth 25 pts
        req_pct  = len(req_hits)  / max(len(required), 1)
        nice_pct = len(nice_hits) / max(len(nice), 1) if nice else 0.0
        if not required and not nice:
            ats_score = 0
        elif not required:
            ats_score = round(nice_pct * 60)   # no required skills defined → cap at 60
        else:
            ats_score = round(req_pct * 75 + nice_pct * 25)

        # Keep original-case skill names for display
        cand_stack_orig = {s.lower(): s for s in (cand.get("tech_stack") or [])}
        jd_required_orig = {s.lower(): s for s in jd_meta.get("required_skills") or []}
        jd_nice_orig     = {s.lower(): s for s in jd_meta.get("nice_to_have_skills") or []}

        matched  = sorted(jd_required_orig.get(k, k) for k in req_hits)
        matched += sorted(jd_nice_orig.get(k, k) for k in nice_hits)
        missing  = sorted(jd_required_orig.get(k, k) for k in (required - cand_set))

        # Templated fit summary (no LLM needed)
        nr, nt = len(required), len(req_hits)
        if ats_score >= 75:
            summary = f"Strong match: {nt}/{nr} required skills covered."
        elif ats_score >= 50:
            summary = f"Good match: {nt}/{nr} required skills found. Gaps are bridgeable."
        elif ats_score >= 25:
            summary = f"Partial match: only {nt}/{nr} required skills present."
        else:
            summary = f"Limited overlap: {nt}/{nr} required skills found in tech stack."

        return {
            "id":               cand["id"],
            "full_name":        cand.get("full_name"),
            "candidate_role":   cand.get("candidate_role"),
            "experience_years": cand.get("experience_years"),
            "total_experience": cand.get("total_experience"),
            "tech_stack":       cand.get("tech_stack") or [],
            "location":         cand.get("location"),
            "visa_status":      cand.get("visa_status"),
            "status":           cand.get("status"),
            "ats_score":        ats_score,
            "matched_skills":   matched,
            "missing_skills":   missing,
            "fit_summary":      summary,
        }

    # Score every candidate, keep those with at least 1 skill match for pre_filtered count
    all_scored   = [_score_candidate(c) for c in all_candidates]
    pre_filtered = [s for s in all_scored if s["ats_score"] > 0]

    # Sort by score desc, return top 10
    top10 = sorted(all_scored, key=lambda x: x["ats_score"], reverse=True)[:10]

    return {
        "matches":       top10,
        "jd_meta":       jd_meta,
        "total_scanned": len(all_candidates),
        "pre_filtered":  len(pre_filtered),
    }


# ============================================================
# LINKEDIN POSTING
# ============================================================

async def _post_job_to_linkedin(job: dict, apply_url: str) -> dict:
    """
    Post a job opening to JM Data Talent LinkedIn company page.
    Returns {"success": True/False, "post_id": "...", "error": "..."}
    
    Required HuggingFace env vars:
      LINKEDIN_ACCESS_TOKEN    — OAuth 2.0 access token (pages:read + w_organization_social)
      LINKEDIN_ORGANIZATION_ID — numeric ID from linkedin.com/company/<name>/admin/ URL
    """
    token  = os.environ.get("LINKEDIN_ACCESS_TOKEN", "").strip()
    # org_id = os.environ.get("LINKEDIN_ORGANIZATION_ID", "").strip() # having issues 
    member_id = os.environ.get("LINKEDIN_MEMBER_ID", "").strip()

    # if not token or not org_id:
    if not token or not member_id:
        return {"success": False, "error": "LinkedIn credentials not configured — set LINKEDIN_ACCESS_TOKEN and LINKEDIN_ORGANIZATION_ID in HuggingFace Spaces secrets."}

    # ── Build post text via LLM (falls back to template if LLM unavailable) ──
    post_text = ""
    if LLM_AVAILABLE:
        try:
            from llm_utils import generate_linkedin_post as _gen_li_post
            post_text = await _gen_li_post(job, apply_url)
        except Exception as _li_llm_err:
            logger.warning(f"[LinkedIn] LLM post failed: {_li_llm_err}")

    if not post_text:
        # Template fallback
        skills_lines = ""
        if job.get("skills"):
            skills_lines = "\n" + "\n".join(f"• {s}" for s in (job["skills"] or [])[:6])

        desc = (job.get("description") or "").strip()
        short_desc = desc[:300] + ("…" if len(desc) > 300 else "")

        location_line = ""
        if job.get("location"):
            location_line = f"📍 {job['location']}"
        if job.get("employment_type"):
            location_line += f" | {job['employment_type']}" if location_line else f"💼 {job['employment_type']}"

        post_text = f"""🚀 We're Hiring: {job['title']}

{location_line}

{short_desc}{skills_lines}

👉 Apply here → {apply_url}

#Hiring #ITJobs #Dublin #Ireland #JMDataTalent #TechJobs #Recruitment"""

    payload = {
        # "author": f"urn:li:organization:{org_id}",
        "author": f"urn:li:person:{member_id}",
        "lifecycleState": "PUBLISHED",
        "specificContent": {
            "com.linkedin.ugc.ShareContent": {
                "shareCommentary": {"text": post_text},
                "shareMediaCategory": "ARTICLE",
                "media": [{
                    "status": "READY",
                    "description": {"text": short_desc or f"{job['title']} — Apply at JM Data Talent"},
                    "originalUrl": apply_url,
                    "title": {"text": f"Apply: {job['title']} at JM Data Talent"},
                }],
            }
        },
        "visibility": {"com.linkedin.ugc.MemberNetworkVisibility": "PUBLIC"},
    }

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type":  "application/json",
        "X-Restli-Protocol-Version": "2.0.0",
        "LinkedIn-Version": "202401",
    }

    try:
        import httpx
        async with httpx.AsyncClient(timeout=15) as client:
            resp = await client.post("https://api.linkedin.com/v2/ugcPosts", json=payload, headers=headers)
        if resp.status_code in (200, 201):
            return {"success": True, "post_id": resp.headers.get("x-linkedin-id", "posted")}
        return {"success": False, "error": f"LinkedIn API returned {resp.status_code}: {resp.text[:200]}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# BULK RESUME ZIP UPLOAD
# ============================================================

import uuid as _uuid
import zipfile as _zipfile
import io as _io

# In-memory job tracking (cleared on restart — that is fine for short jobs)
_bulk_jobs: dict = {}


async def _process_bulk_zip(job_id: str, zip_bytes: bytes, user: dict):
    """Background task: extract ZIP → parse each resume → upload Drive → insert candidate."""
    state = _bulk_jobs.get(job_id)
    if not state:
        logging.getLogger(__name__).error(f"[bulk] job_id {job_id} not found in _bulk_jobs")
        return
    try:
        from llm_utils import extract_resume_full_profile
    except ImportError as ie:
        state["status"] = "error"
        state["error"]  = f"LLM import failed: {ie}"
        logging.getLogger(__name__).error(f"[bulk] llm_utils import failed: {ie}")
        return

    ALLOWED = {
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    EXT_MIME = {".pdf": "application/pdf", ".doc": "application/msword", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}

    try:
        zf = _zipfile.ZipFile(_io.BytesIO(zip_bytes))
    except Exception as e:
        state["status"] = "error"
        state["error"]  = f"Could not open ZIP: {e}"
        return

    # Filter to resume files only (skip __MACOSX junk and directories)
    entries = [
        n for n in zf.namelist()
        if not n.startswith("__MACOSX") and not n.endswith("/")
        and any(n.lower().endswith(ext) for ext in EXT_MIME)
    ]

    state["total"] = len(entries)
    if not entries:
        state["status"] = "done"
        return

    for i, name in enumerate(entries):
        fname = name.split("/")[-1]  # strip any folder prefix inside zip
        ext   = "." + fname.rsplit(".", 1)[-1].lower()
        mime  = EXT_MIME.get(ext, "application/pdf")

        result_entry = {"filename": fname, "status": "processing", "candidate_id": None, "error": None}
        state["results"].append(result_entry)

        try:
            file_bytes = zf.read(name)

            # 1. Extract text
            resume_text = _extract_resume_text(file_bytes, mime)
            if not resume_text.strip():
                result_entry["status"] = "skipped"
                result_entry["error"]  = "Could not extract text"
                state["processed"] += 1
                continue

            # 2. LLM — full profile (1 call per resume)
            profile = await extract_resume_full_profile(resume_text)

            # 3. Upload to Google Drive — upload_resume returns dict {file_id, preview_url, view_url}
            drive_url = None
            if upload_resume is not None:
                try:
                    _loop = asyncio.get_running_loop()
                    drive_result = await _loop.run_in_executor(
                        None, lambda b=file_bytes, fn=fname, m=mime: upload_resume(b, fn, m)
                    )
                    if isinstance(drive_result, dict):
                        drive_url = drive_result.get("preview_url") or drive_result.get("view_url")
                    elif isinstance(drive_result, str):
                        drive_url = drive_result
                except Exception as de:
                    logging.getLogger(__name__).warning(f"[bulk] Drive upload failed for {fname}: {de}")

            # 4. Insert candidate (upsert by email if available)
            row = {
                "full_name":        profile.get("full_name") or fname.rsplit(".", 1)[0].replace("_"," ").replace("-"," ").title(),
                "email":            profile.get("email"),
                "phone":            profile.get("phone"),
                "current_company":  profile.get("current_company"),
                "candidate_role":   profile.get("candidate_role"),
                "experience_years": profile.get("experience_years"),
                "location":         profile.get("location"),
                "tech_stack":       profile.get("tech_stack") or [],
                "skills":           profile.get("tech_stack") or [],
                "resume_url":       drive_url,
                "source":           "bulk_upload",
                "status":           "sourced",
                "created_by":       user.get("id"),
            }

            sb_client = get_supabase()
            # Upsert by email if available, else plain insert
            if row["email"]:
                existing = sb_client.table("candidates").select("id").eq("email", row["email"]).execute()
                if existing.data:
                    cid = existing.data[0]["id"]
                    sb_client.table("candidates").update({k: v for k, v in row.items() if v is not None}).eq("id", cid).execute()
                    result_entry["candidate_id"] = cid
                    result_entry["action"] = "updated"
                else:
                    ins = sb_client.table("candidates").insert(row).execute()
                    result_entry["candidate_id"] = ins.data[0]["id"] if ins.data else None
                    result_entry["action"] = "created"
            else:
                ins = sb_client.table("candidates").insert(row).execute()
                result_entry["candidate_id"] = ins.data[0]["id"] if ins.data else None
                result_entry["action"] = "created"

            result_entry["status"]    = "done"
            result_entry["full_name"] = row["full_name"]
            result_entry["email"]     = row["email"]
            result_entry["candidate_role"] = row["candidate_role"]
            state["created"] += 1

        except Exception as exc:
            result_entry["status"] = "error"
            result_entry["error"]  = str(exc)[:200]
            state["errors"] += 1

        state["processed"] += 1

        # Sleep between calls to respect Groq free-tier rate limits
        if i < len(entries) - 1:
            await asyncio.sleep(4)

    state["status"] = "done"
    zf.close()


@api_router.post("/candidates/bulk-upload-zip")
async def bulk_upload_zip(
    request: Request,
    background_tasks: BackgroundTasks,
    zip_file: UploadFile = File(...),
):
    """Accept a ZIP of resumes, process in background, return a job_id to poll."""
    user = await get_current_user(request)
    _require_module(user, "recruitment")

    if not zip_file.filename.lower().endswith(".zip"):
        raise HTTPException(400, "Please upload a .zip file.")

    raw = await zip_file.read()
    if len(raw) > 100 * 1024 * 1024:  # 100 MB limit
        raise HTTPException(400, "ZIP file too large (max 100 MB).")

    job_id = str(_uuid.uuid4())
    _bulk_jobs[job_id] = {
        "status": "running", "total": 0, "processed": 0,
        "created": 0, "errors": 0, "results": [],
    }

    background_tasks.add_task(_process_bulk_zip, job_id, raw, user)
    asyncio.create_task(_audit("import", user=user, entity_type="candidate",
                                entity_name=zip_file.filename,
                                new_value={"filename": zip_file.filename, "job_id": job_id, "type": "bulk_zip"},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"job_id": job_id}


@api_router.get("/candidates/bulk-upload-status/{job_id}")
async def bulk_upload_status(job_id: str, request: Request):
    """Poll for bulk upload progress."""
    await get_current_user(request)
    if job_id not in _bulk_jobs:
        raise HTTPException(404, "Job not found. Jobs are cleared on server restart.")
    return _bulk_jobs[job_id]


# ─────────────────────────────────────────────────────────────
# ATS RESUME UPLOAD SCORING  (free-tier safe: 1 LLM call/resume)
# ─────────────────────────────────────────────────────────────

class BackfillCandidateRoleRequest(BaseModel):
    limit: Optional[int] = 100


@api_router.post("/candidates/backfill-roles")
async def backfill_candidate_roles(request: Request, body: BackfillCandidateRoleRequest = None):
    """
    Backfill missing candidate_role values from already stored Drive resumes.
    Only updates rows where candidate_role is empty and resume_url is present.
    """
    user = await get_current_user(request)
    if user.get("role") != "admin":
        raise HTTPException(403, "Only admins can run role backfills.")

    if download_resume is None:
        raise HTTPException(503, "Google Drive download is not configured.")

    limit = int(getattr(body, "limit", 100) or 100)
    limit = max(1, min(limit, 500))

    from llm_utils import extract_resume_full_profile

    rows_res = await run(
        lambda: sb("candidates")
        .select("id,full_name,candidate_role,resume_url")
        .execute()
    )
    rows = getattr(rows_res, "data", []) or []

    updated = 0
    scanned = 0
    skipped = 0
    errors = 0

    for cand in rows:
        if scanned >= limit:
            break
        if (cand.get("candidate_role") or "").strip():
            continue
        resume_url = (cand.get("resume_url") or "").strip()
        if not resume_url:
            continue

        scanned += 1
        try:
            file_bytes = await run(lambda url=resume_url: download_resume(url))

            resume_text = _extract_resume_text(file_bytes, "application/pdf")
            if not resume_text.strip():
                resume_text = _extract_resume_text(file_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if not resume_text.strip():
                resume_text = _extract_resume_text(file_bytes, "application/msword")

            if not resume_text.strip():
                skipped += 1
                continue

            profile = await extract_resume_full_profile(resume_text)
            candidate_role = (profile.get("candidate_role") or "").strip()
            if not candidate_role:
                skipped += 1
                continue

            await run(
                lambda cid=cand["id"], role=candidate_role: sb("candidates")
                .update({"candidate_role": role})
                .eq("id", cid)
                .execute()
            )
            updated += 1
        except Exception as exc:
            errors += 1
            logger.warning(f"[backfill-roles] failed for {cand.get('id')}: {exc}")

    return {
        "scanned": scanned,
        "updated": updated,
        "skipped": skipped,
        "errors": errors,
        "limit": limit,
    }


class ParseJDRequest(BaseModel):
    jd_text: str

@api_router.post("/candidates/parse-jd-for-scoring")
async def parse_jd_for_scoring(body: ParseJDRequest, request: Request):
    """Parse JD into structured skills (1 LLM call). Call once before scoring resumes."""
    await get_current_user(request)
    if not LLM_AVAILABLE:
        raise HTTPException(503, "LLM service not configured. Set LLM_PROVIDER and API key.")
    jd_text = (body.jd_text or "").strip()
    if len(jd_text) < 30:
        raise HTTPException(400, "Job description is too short (minimum 30 characters).")
    jd_meta = await extract_jd_keywords(jd_text)
    return jd_meta


@api_router.post("/candidates/score-resume-upload")
async def score_resume_upload(
    request: Request,
    resume: UploadFile = File(...),
    jd_skills: str = Form(...),
):
    """
    Score a single uploaded resume against pre-parsed JD skills.
    1 LLM call per resume. Caller should sleep 2s between consecutive calls (Groq free tier).
    """
    await get_current_user(request)

    allowed_ct = {
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    ct = (resume.content_type or "").split(";")[0].strip()
    if ct not in allowed_ct:
        raise HTTPException(400, f"Unsupported file: {resume.filename}. Upload PDF or DOCX only.")

    contents = await resume.read()
    if len(contents) > 10 * 1024 * 1024:
        raise HTTPException(400, "Resume too large (max 10 MB).")

    resume_text = _extract_resume_text(contents, ct)
    if not resume_text.strip():
        return {
            "filename": resume.filename,
            "ats_score": 0,
            "matched_skills": [],
            "missing_skills": [],
            "tech_stack": [],
            "experience_years": None,
            "fit_summary": "Could not extract text from this file.",
            "error": "text_extraction_failed",
        }

    insights = await extract_resume_insights(resume_text)
    tech_stack = insights.get("tech_stack") or []
    experience_years = insights.get("experience_years")

    try:
        jd_meta = json.loads(jd_skills)
    except Exception:
        raise HTTPException(400, "Invalid jd_skills JSON.")

    required = set(s.lower() for s in jd_meta.get("required_skills") or [])
    nice     = set(s.lower() for s in jd_meta.get("nice_to_have_skills") or [])
    cand_set = set(s.lower() for s in tech_stack)

    req_hits  = cand_set & required
    nice_hits = cand_set & nice
    req_pct   = len(req_hits) / max(len(required), 1)
    nice_pct  = len(nice_hits) / max(len(nice), 1) if nice else 0.0

    if not required and not nice:
        ats_score = 0
    elif not required:
        ats_score = round(nice_pct * 60)
    else:
        ats_score = round(req_pct * 75 + nice_pct * 25)

    jd_req_orig  = {s.lower(): s for s in jd_meta.get("required_skills") or []}
    jd_nice_orig = {s.lower(): s for s in jd_meta.get("nice_to_have_skills") or []}

    matched  = sorted(jd_req_orig.get(k, k)  for k in req_hits)
    matched += sorted(jd_nice_orig.get(k, k) for k in nice_hits)
    missing  = sorted(jd_req_orig.get(k, k)  for k in (required - cand_set))

    nr, nt = len(required), len(req_hits)
    if ats_score >= 75:
        summary = f"Strong match: {nt}/{nr} required skills covered."
    elif ats_score >= 50:
        summary = f"Good match: {nt}/{nr} required skills found. Gaps are bridgeable."
    elif ats_score >= 25:
        summary = f"Partial match: only {nt}/{nr} required skills present."
    else:
        summary = f"Limited overlap: {nt}/{nr} required skills found."

    return {
        "filename": resume.filename,
        "ats_score": ats_score,
        "matched_skills": matched,
        "missing_skills": missing,
        "tech_stack": tech_stack,
        "experience_years": experience_years,
        "fit_summary": summary,
    }


# ============================================================

# INTERVIEWS
# ============================================================
@api_router.post("/interviews")
async def create_interview(interview: InterviewCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    res = await run(lambda: sb("interviews").insert({
        "candidate_id":  interview.candidate_id,
        "job_id":        interview.job_id,
        "scheduled_at":  interview.scheduled_at,
        "interview_type":interview.interview_type,
        "interviewers":  interview.interviewers,
        "notes":         interview.notes,
        "created_by":    user["id"],
    }).execute())

    # Move candidate to interview_scheduled
    await run(lambda: sb("candidates").update({"status": "interview_scheduled"}).eq("id", interview.candidate_id).execute())
    await _log_activity(candidate_id=interview.candidate_id, user=user, atype="interview",
                        desc=f"{interview.interview_type} scheduled for {interview.scheduled_at}")
    asyncio.create_task(_audit("create", user=user, entity_type="interview",
                                entity_id=res.data[0].get("id"),
                                entity_name=f"{interview.interview_type} — {interview.scheduled_at[:10]}",
                                new_value={"type": interview.interview_type, "scheduled_at": interview.scheduled_at,
                                           "candidate_id": interview.candidate_id},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return res.data[0]


@api_router.get("/interviews")
async def get_interviews(
    request:      Request,
    candidate_id: Optional[str]  = None,
    upcoming:     Optional[bool] = None,
):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    q = sb("interviews").select(
        "*, candidate:candidate_id(full_name), job:job_id(title)"
    ).order("scheduled_at")
    if candidate_id: q = q.eq("candidate_id", candidate_id)
    if upcoming:
        now = datetime.now(timezone.utc).isoformat()
        q = q.gte("scheduled_at", now).eq("completed", False)
    res = await run(lambda: q.execute())
    return res.data or []


@api_router.put("/interviews/{interview_id}")
async def update_interview(interview_id: str, interview: InterviewUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    patch = {k: v for k, v in interview.model_dump().items() if v is not None}
    if interview.completed:
        patch["completed_at"] = datetime.now(timezone.utc).isoformat()

    await run(lambda: sb("interviews").update(patch).eq("id", interview_id).execute())

    # If marked complete, advance candidate to interviewed
    if interview.completed:
        iv_cand = await safe_single(lambda: sb("interviews").select("candidate_id").eq("id", interview_id).single().execute())
        if iv_cand:
            cand_id = iv_cand["candidate_id"]
            await run(lambda: sb("candidates").update({"status": "interviewed"}).eq("id", cand_id).execute())
            await _log_activity(candidate_id=cand_id, user=user, atype="interview",
                                 desc="Interview completed")
    asyncio.create_task(_audit("update", user=user, entity_type="interview",
                                entity_id=interview_id, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Interview updated"}




# ============================================================
# RECRUITMENT DASHBOARD
# ============================================================
@api_router.get("/dashboard/recruitment")
async def recruitment_dashboard(request: Request):
    user  = await get_current_user(request)
    _require_module(user, "recruitment")
    today = datetime.now(timezone.utc).date().isoformat()

    # Fire all independent queries in parallel
    cands_all_res, active_jobs, upcoming_ivs, tasks_today, recent_cands = await asyncio.gather(
        run(lambda: sb("candidates").select("id,status").execute()),
        run(lambda: sb("jobs").select("id", count="exact").eq("is_active", True).execute()),
        run(lambda: sb("interviews").select("*, candidate:candidate_id(full_name), job:job_id(title)").gte("scheduled_at", today).eq("completed", False).order("scheduled_at").limit(10).execute()),
        run(lambda uid=user["id"]: sb("tasks").select("*").eq("assigned_to", uid).eq("due_date", today).eq("completed", False).execute()),
        run(lambda: sb("candidates").select("*").order("created_at", desc=True).limit(10).execute()),
    )
    # Build cand_stats from already-fetched data (no extra queries)
    cand_stats = {s.value: 0 for s in CandidateStatus}
    for c in (cands_all_res.data or []):
        st = c.get("status")
        if st in cand_stats:
            cand_stats[st] += 1
    total_cands_count = len(cands_all_res.data or [])

    return {
        "candidate_stats":    cand_stats,
        "total_candidates":   total_cands_count,
        "active_jobs":        active_jobs.count or 0,
        "upcoming_interviews":upcoming_ivs.data or [],
        "today_tasks":        tasks_today.data or [],
        "recent_candidates":  recent_cands.data or [],
    }



# CANDIDATE SUBMISSIONS (Lead ↔ Candidate link)
# ============================================================
@api_router.post("/submissions")
async def create_submission(body: SubmissionCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    doc = {
        "lead_id":      body.lead_id,
        "candidate_id": body.candidate_id,
        "status":       body.status,
        "notes":        body.notes,
        "created_by":   user["id"],
    }
    res = await run(lambda: sb("candidate_submissions").upsert(doc, on_conflict="lead_id,candidate_id").execute())
    await _log_activity(lead_id=body.lead_id, user=user, atype="note",
                        desc=f"Candidate submission created/updated")
    row = res.data[0] if res.data else {}
    asyncio.create_task(_audit("create", user=user, entity_type="submission",
                                entity_id=row.get("id"),
                                entity_name=f"Submission — lead:{body.lead_id}",
                                new_value={"lead_id": body.lead_id, "candidate_id": body.candidate_id, "status": body.status},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return row

@api_router.get("/submissions")
async def get_submissions(request: Request, lead_id: Optional[str] = None, candidate_id: Optional[str] = None):
    user = await get_current_user(request)
    _require_module(user, "sales")
    q = sb("candidate_submissions").select(
        "*, candidate:candidate_id(id,full_name,candidate_role,status,email), lead:lead_id(id,full_name,company,status)"
    ).order("created_at", desc=True)
    if lead_id:      q = q.eq("lead_id", lead_id)
    if candidate_id: q = q.eq("candidate_id", candidate_id)
    res = await run(lambda: q.execute())
    return res.data or []

@api_router.put("/submissions/{submission_id}")
async def update_submission(submission_id: str, body: SubmissionUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    patch = {k: v for k, v in body.model_dump().items() if v is not None}
    await run(lambda: sb("candidate_submissions").update(patch).eq("id", submission_id).execute())
    asyncio.create_task(_audit("update", user=user, entity_type="submission",
                                entity_id=submission_id, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Updated"}

@api_router.delete("/submissions/{submission_id}")
async def delete_submission(submission_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    await run(lambda: sb("candidate_submissions").delete().eq("id", submission_id).execute())
    asyncio.create_task(_audit("delete", user=user, entity_type="submission",
                                entity_id=submission_id,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Deleted"}


# ============================================================

# ============================================================
# PUBLIC JOB API  (no authentication required — secured via
#                  per-job apply_key validated server-side)
#
# Security model:
#   • apply_key is a random URL-safe string stored in the jobs table.
#   • It is generated server-side at job creation (secrets.token_urlsafe).
#   • Even if a candidate tampers the URL, a wrong key returns 404 —
#     no CRM data is exposed and no credentials touch the browser.
#   • job_title is NEVER taken from user input; it is always fetched
#     from the database using the validated key.
#   • Rate-limited: 5 submissions per minute per IP (slowapi).
# ============================================================
 
 
@api_router.get("/public/jobs")
async def public_list_jobs(
    search:          Optional[str] = None,
    department:      Optional[str] = None,
    employment_type: Optional[str] = None,
):
    """
    PUBLIC — No authentication required.
    Returns all active job listings for display on jmdatatalent.com/jobs.
 
    Query params (all optional):
      search          — Full-text search in title, department, description
      department      — Filter by department name (exact, case-insensitive)
      employment_type — Filter by type: Full-time | Part-time | Contract | Internship
    """
    result = await run(
        lambda: sb("jobs")
        .select("id,title,department,location,employment_type,description,"
                "requirements,salary_range,skills,is_urgent,created_at,updated_at")
        .eq("is_active", True)
        .order("is_urgent", desc=True)
        .order("created_at", desc=True)
        .execute()
    )
    jobs = result.data or []
 
    if search:
        s = search.lower()
        jobs = [
            j for j in jobs
            if s in (j.get("title") or "").lower()
            or s in (j.get("department") or "").lower()
            or s in (j.get("description") or "").lower()
            or s in " ".join(j.get("skills") or []).lower()
        ]
    if department:
        jobs = [j for j in jobs if (j.get("department") or "").lower() == department.lower()]
    if employment_type:
        jobs = [
            j for j in jobs
            if (j.get("employment_type") or "").lower() == employment_type.lower()
        ]
 
    return {
        "success":    True,
        "count":      len(jobs),
        "jobs":       jobs,
        "generated":  datetime.now(timezone.utc).isoformat(),
    }
 
 
@api_router.get("/public/jobs/{job_id}")
async def public_get_job(job_id: str):
    """
    PUBLIC — No authentication required.
    Returns full details of a single active job.
    Returns 404 if the job does not exist or is not active.
    """
    job = await safe_single(
        lambda: sb("jobs")
        .select("id,title,department,location,employment_type,description,"
                "requirements,salary_range,skills,is_urgent,created_at,updated_at")
        .eq("id", job_id)
        .eq("is_active", True)
        .single()
        .execute()
    )
    if not job:
        raise HTTPException(
            404,
            detail={
                "error":   "JOB_NOT_FOUND",
                "message": "This job posting does not exist or is no longer active.",
                "job_id":  job_id,
            },
        )
    return {"success": True, "job": job}


@api_router.get("/public/jobs/by-key/{apply_key}")
async def public_get_job_by_key(apply_key: str):
    """
    PUBLIC — No authentication required.
    Validates the apply_key and returns safe job details for the application form.
    Returns 404 if the key is wrong, the job is closed, or the job doesn't exist.

    Only returns data the candidate is allowed to see (title, dept, location, type).
    Never exposes internal IDs, credentials, or CRM data.
    """
    apply_key = apply_key.strip()
    if not apply_key or len(apply_key) > 64:
        raise HTTPException(404, detail={"error": "JOB_NOT_FOUND", "message": "Job not found."})

    job = await safe_single(
        lambda: sb("jobs")
        .select("id,title,department,location,employment_type,is_urgent")
        .eq("apply_key", apply_key)
        .eq("is_active", True)
        .single()
        .execute()
    )
    if not job:
        raise HTTPException(
            404,
            detail={
                "error":   "JOB_NOT_FOUND",
                "message": "This job is no longer available or the link is invalid.",
            },
        )
    # Return only what the form needs — internal job UUID is intentionally excluded
    return {
        "success":         True,
        "title":           job["title"],
        "department":      job.get("department") or "",
        "location":        job.get("location") or "",
        "employment_type": job.get("employment_type") or "",
        "is_urgent":       job.get("is_urgent") or False,
    }


@api_router.post("/public/apply")
@limiter.limit("5/minute")          # 5 submissions per IP per minute — prevents spam bots
async def public_apply_job(
    request:          Request,
    # Required fields
    first_name:       str            = Form(..., description="Applicant's first name"),
    last_name:        str            = Form(..., description="Applicant's last name"),
    email:            str            = Form(..., description="Applicant's email address"),
    phone:            str            = Form(..., description="Phone with country code, e.g. +91-9876543210"),
    apply_key:        str            = Form(..., description="Unique job key from the URL (?key=…)"),
    resume:           UploadFile     = File(...,  description="Resume file — PDF, DOC, or DOCX, max 10 MB"),
    # Optional fields
    current_company:  Optional[str]  = Form(None, description="Current employer"),
    candidate_role:   Optional[str]  = Form(None, description="Current job title / role"),
    experience_years: Optional[str]  = Form(None, description="Total years of experience (numeric)"),
    linkedin_url:     Optional[str]  = Form(None, description="LinkedIn profile URL"),
    portfolio_url:    Optional[str]  = Form(None, description="Portfolio, GitHub, or personal site URL"),
):
    """
    PUBLIC — No authentication header required.

    Security: job_id and job_title are NEVER taken from user input.
    The apply_key is validated against the database; the real job data
    is fetched server-side. A tampered key returns 404 — no CRM data exposed.

    Returns 409 if the same email has already applied for the same job.
    Returns 429 if the IP exceeds 5 submissions per minute.
    """

    # ── 0. Compute full name ─────────────────────────────────────
    first_name = first_name.strip()
    last_name  = last_name.strip()
    if not first_name:
        raise HTTPException(422, detail={"error": "MISSING_FIRST_NAME", "message": "First name is required.", "field": "first_name"})
    if not last_name:
        raise HTTPException(422, detail={"error": "MISSING_LAST_NAME", "message": "Last name is required.", "field": "last_name"})
    full_name = f"{first_name} {last_name}"

    # ── 1. Validate apply_key → fetch real job data server-side ──
    #   This is the core security step. job_title is NEVER sourced from user input.
    apply_key = apply_key.strip()
    if not apply_key or len(apply_key) > 64:
        raise HTTPException(404, detail={"error": "JOB_NOT_FOUND", "message": "Invalid application link."})

    job_row = await safe_single(
        lambda: sb("jobs")
        .select("id,title,department")
        .eq("apply_key", apply_key)
        .eq("is_active", True)
        .single()
        .execute()
    )
    if not job_row:
        raise HTTPException(
            404,
            detail={
                "error":   "JOB_NOT_FOUND",
                "message": "This job is no longer available. The position may have been filled.",
            },
        )
    # Real values from the DB — not from the user
    real_job_id    = job_row["id"]
    real_job_title = job_row["title"]

    # ── 2. Validate required text fields ─────────────────────────
    # (No job_id / job_title validation needed — they come from DB)

    # ── 3. Validate email ─────────────────────────────────────────
    email = email.strip().lower()
    if not _re.match(r"^[^\s@]+@[^\s@]+\.[^\s@]+$", email):
        raise HTTPException(
            422,
            detail={
                "error":   "INVALID_EMAIL",
                "message": "Please provide a valid email address (e.g. name@example.com).",
                "field":   "email",
            },
        )

    # ── 4. Validate phone ─────────────────────────────────────────
    try:
        phone_clean = _validate_phone(phone)
    except ValueError as exc:
        raise HTTPException(
            422,
            detail={
                "error":   "INVALID_PHONE",
                "message": str(exc),
                "field":   "phone",
                "example": "+91-9876543210 or +353851234567",
            },
        )

    # ── 5. Validate resume file ───────────────────────────────────
    if resume.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            422,
            detail={
                "error":    "INVALID_FILE_TYPE",
                "message":  f"Resume must be PDF, DOC, or DOCX. Received: {resume.content_type}",
                "field":    "resume",
                "accepted": ["application/pdf",
                             "application/msword",
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
            },
        )
    resume_bytes = await resume.read()
    if len(resume_bytes) > MAX_FILE_BYTES:
        raise HTTPException(
            422,
            detail={
                "error":   "FILE_TOO_LARGE",
                "message": f"Resume must be under 10 MB. Received: {round(len(resume_bytes)/1024/1024, 1)} MB.",
                "field":   "resume",
                "max_mb":  10,
            },
        )

    # ── 6. Duplicate check — same email + real job UUID ──────────
    dup_check = await run(
        lambda: sb("candidates")
        .select("id")
        .eq("email", email)
        .eq("job_id", real_job_id)
        .eq("source", "website")
        .execute()
    )
    if dup_check.data:
        raise HTTPException(
            409,
            detail={
                "error":   "ALREADY_APPLIED",
                "message": "An application with this email already exists for this position. "
                           "Our team will review it and be in touch.",
                "email":   email,
            },
        )

    # ── 7. Upload resume to Google Drive ─────────────────────────
    ext       = ALLOWED_MIME_TYPES[resume.content_type]
    safe_name = _re.sub(r"[^\w\s-]", "", full_name.strip()).replace(" ", "_")[:30]
    short_id  = str(uuid.uuid4()).replace("-", "")[:8]
    filename  = f"WebApp_{safe_name}_{short_id}.{ext}"

    resume_url = None
    if upload_resume is not None:
        try:
            drive_result = await run(
                lambda: upload_resume(
                    resume_bytes,
                    filename,
                    resume.content_type,
                    folder_id=PUBLIC_RESUME_FOLDER_ID,
                )
            )
            resume_url = drive_result["preview_url"]
        except Exception as exc:
            logger.error(f"[public-apply] Drive upload failed: {exc}")
    else:
        logger.warning("[public-apply] Google Drive not configured — resume not uploaded")

    # ── 8. Parse optional numeric fields ─────────────────────────
    exp_years = None
    if experience_years:
        try:
            exp_years = int(float(experience_years.strip()))
            if exp_years < 0 or exp_years > 60:
                exp_years = None
        except (ValueError, AttributeError):
            pass

    # ── 9. Create candidate record ────────────────────────────────
    candidate_payload = {
        "full_name":        full_name,
        "email":            email,
        "phone":            phone_clean,
        "current_company":  (current_company  or "").strip() or None,
        "candidate_role":   (candidate_role   or "").strip() or None,
        "experience_years": exp_years,
        "linkedin_url":     (linkedin_url     or "").strip() or None,
        "portfolio_url":    (portfolio_url    or "").strip() or None,
        "source":           "website",
        "status":           "sourced",
        "job_id":           real_job_id,       # from DB — not from user
        "job_title":        real_job_title,    # from DB — not from user
        "notes":            f"Applied via website for: {real_job_title}",
        "resume_url":       resume_url,
    }

    result = await run(lambda: sb("candidates").insert(candidate_payload).execute())
    if not result.data:
        raise HTTPException(
            500,
            detail={
                "error":   "SERVER_ERROR",
                "message": "Failed to record your application. Please try again or contact us directly.",
            },
        )

    new_candidate = result.data[0]

    # ── 10. Audit log ─────────────────────────────────────────────
    asyncio.create_task(_audit(
        action="create",
        user={"id": None, "email": email, "name": full_name},
        entity_type="candidate",
        entity_id=new_candidate["id"],
        entity_name=f"{full_name} → {real_job_title} (website)",
    ))

    logger.info(
        f"[public-apply] New application: {full_name} <{email}> "
        f"→ {real_job_title} (job_id={real_job_id}) candidate_id={new_candidate['id']}"
    )

    return {
        "success":        True,
        "application_id": new_candidate["id"],
        "message":        "Your application has been submitted successfully. "
                          "Our team will review it and be in touch.",
    }




# ============================================================
# HEALTH  (UptimeRobot pings this every 5 min to keep Space warm)
# ============================================================
@app.api_route("/health", methods=["GET", "HEAD"])
async def root_health():
    return {
        "status":  "ok",
        "service": "Nexus CRM — Recruit & Onboarding Service",
        "version": "2.0.0",
    }


# ============================================================
# WIRE UP
# ============================================================
app.include_router(api_router)
app.add_middleware(GZipMiddleware, minimum_size=512)
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)
