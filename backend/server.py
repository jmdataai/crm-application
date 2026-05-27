from dotenv import load_dotenv
from pathlib import Path

# Load both the local backend .env and the default search path so the app works
# whether it is started from the repo root or from the backend directory.
load_dotenv(Path(__file__).with_name(".env"))
load_dotenv()

from fastapi import FastAPI, APIRouter, BackgroundTasks, HTTPException, Request, Response, File, UploadFile, Form
from starlette.middleware.cors import CORSMiddleware
from starlette.middleware.gzip import GZipMiddleware
from supabase import create_client, Client
import inspect
try:
    from supabase import ClientOptions
except Exception:
    ClientOptions = None
import os
import logging
import bcrypt
import jwt
import asyncio
import httpx
import resend
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron  import CronTrigger
import pandas as pd
import io
import uuid
from datetime import datetime, timezone, timedelta
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from typing import List, Optional
from pydantic import BaseModel, EmailStr, field_validator, model_validator
from enum import Enum
import re as _re
import json
import secrets                        # apply_key generation
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ── PyMuPDF — resume masking ─────────────────────────────────
try:
    import fitz as _fitz
    _FITZ_OK = True
except ImportError:
    _fitz    = None
    _FITZ_OK = False

# ── MSAL — Microsoft Graph email ─────────────────────────────
try:
    import msal as _msal
    _MSAL_OK = True
except ImportError:
    _msal    = None
    _MSAL_OK = False

MS_TENANT_ID = os.environ.get("MS_TENANT_ID", "").strip()
MS_CLIENT_ID_ENV = os.environ.get("MS_CLIENT_ID", "").strip()
MS_CLIENT_SECRET = (
    os.environ.get("MS_CLIENT_SECRET_VALUE", "").strip()
    or os.environ.get("MS_CLIENT_SECRET", "").strip()
)


def _looks_like_uuid(value: str) -> bool:
    return bool(_re.fullmatch(r"[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}", value or ""))


def _validate_graph_config() -> None:
    missing = [
        name
        for name, value in {
            "MS_TENANT_ID": MS_TENANT_ID,
            "MS_CLIENT_ID": MS_CLIENT_ID_ENV,
            "MS_CLIENT_SECRET_VALUE": MS_CLIENT_SECRET,
        }.items()
        if not value
    ]
    if missing:
        raise RuntimeError(
            "Microsoft Graph not configured. Set MS_TENANT_ID, MS_CLIENT_ID, and "
            "MS_CLIENT_SECRET_VALUE in the environment. "
            "Use the client secret VALUE from Azure App Registration -> Certificates & secrets, "
            "not the secret ID."
        )
    if _looks_like_uuid(MS_CLIENT_SECRET):
        raise RuntimeError(
            "Microsoft Graph client secret looks like a Secret ID, not a secret VALUE. "
            "Open Azure App Registration -> Certificates & secrets and copy the secret VALUE "
            "into MS_CLIENT_SECRET_VALUE. This fixes both test email and real send."
        )

def _get_graph_token() -> str:
    if not _MSAL_OK:
        raise RuntimeError("Microsoft Graph not configured because the msal package is unavailable.")
    _validate_graph_config()
    authority = f"https://login.microsoftonline.com/{MS_TENANT_ID}"
    app = _msal.ConfidentialClientApplication(
        MS_CLIENT_ID_ENV, authority=authority, client_credential=MS_CLIENT_SECRET,
    )
    result = app.acquire_token_for_client(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" not in result:
        raise RuntimeError(f"MSAL error: {result.get('error_description', str(result))}")
    return result["access_token"]

async def _graph_send_mail(from_email: str, to_emails: list, subject: str, html_body: str, save_to_sent: bool = True):
    """Send email via Microsoft Graph, saving to Outlook Sent Items."""
    token = await asyncio.get_event_loop().run_in_executor(None, _get_graph_token)
    url   = f"https://graph.microsoft.com/v1.0/users/{from_email}/sendMail"
    payload = {
        "message": {
            "subject": subject,
            "body": {"contentType": "HTML", "content": html_body},
            "toRecipients": [{"emailAddress": {"address": e}} for e in to_emails],
        },
        "saveToSentItems": save_to_sent,
    }
    async with httpx.AsyncClient(timeout=20) as client:
        resp = await client.post(url, json=payload,
            headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"})
    if resp.status_code not in (200, 202):
        raise RuntimeError(f"Graph sendMail failed {resp.status_code}: {resp.text[:400]}")

# ── Google Drive helpers ──────────────────────────────────────
try:
    from google_drive import upload_resume, delete_resume, download_resume, get_file_metadata, ALLOWED_MIME_TYPES, MAX_FILE_BYTES
    _GDRIVE_OK = True
except Exception as _gdrive_err:
    import logging as _l
    _l.getLogger(__name__).warning(f"[google_drive] import failed — resume upload disabled: {_gdrive_err}")
    upload_resume = delete_resume = download_resume = None
    _GDRIVE_OK = False
    ALLOWED_MIME_TYPES = {
        "application/pdf": "pdf",
        "application/msword": "doc",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    }
    MAX_FILE_BYTES = 10 * 1024 * 1024

# ── LLM helpers ──────────────────────────────────────────────
try:
    from llm_utils import extract_resume_insights, extract_jd_keywords
    LLM_AVAILABLE = True
except Exception as _llm_err:
    import logging as _l
    _l.getLogger(__name__).warning(f"[llm_utils] import failed — LLM disabled: {_llm_err}")
    LLM_AVAILABLE = False

# ── Logging ──────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# ── Supabase client ───────────────────────────────────────────
SUPABASE_URL: str = os.environ["SUPABASE_URL"]
SUPABASE_KEY: str = os.environ["SUPABASE_SERVICE_ROLE_KEY"]   # service role — bypasses RLS
if ClientOptions:
    _supabase_httpx = httpx.Client(http2=False)
    options = None
    try:
        init_sig = inspect.signature(ClientOptions.__init__)
        if "http_client" in init_sig.parameters:
            options = ClientOptions(http_client=_supabase_httpx)
        else:
            options = ClientOptions()
    except Exception:
        options = None

    if options is not None:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY, options=options)
    else:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
else:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# ── Resend email ──────────────────────────────────────────────
resend.api_key  = os.environ.get("RESEND_API_KEY", "")
SENDER_EMAIL    = os.environ.get("SENDER_EMAIL", "onboarding@resend.dev")
APIFY_API_KEY   = os.environ.get("APIFY_API_KEY", "")
APIFY_ACTOR_ID  = "dev_fusion~linkedin-profile-scraper"
scheduler       = AsyncIOScheduler()

# ── JWT ───────────────────────────────────────────────────────
JWT_SECRET    = os.environ.get("JWT_SECRET", "change_me_in_production")
JWT_ALGORITHM = "HS256"

# ── App ───────────────────────────────────────────────────────
app        = FastAPI(title="Nexus CRM + ATS")
api_router = APIRouter(prefix="/api")

# ── Rate limiter (protects public /apply from spam bots) ──────
limiter = Limiter(key_func=get_remote_address, default_limits=[])
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Startup diagnostics — visible in HuggingFace / container logs
logger.info("=== Nexus CRM startup ===")
logger.info(f"  Google Drive : {'✓ loaded' if _GDRIVE_OK else '✗ DISABLED (check logs above)'}")
logger.info(f"  LLM (Gemini) : {'✓ loaded' if LLM_AVAILABLE else '✗ DISABLED (check logs above)'}")


# ============================================================
# ENUMS
# ============================================================
class LeadStatus(str, Enum):
    NEW        = "new"
    CONTACTED  = "contacted"
    CALLED     = "called"
    INTERESTED = "interested"
    CLOSED     = "closed"
    COMPLETED  = "completed"
    REJECTED   = "rejected"
    LOST       = "lost"
    FOLLOW_UP  = "follow_up_needed"

class CandidateStatus(str, Enum):
    SOURCED              = "sourced"
    SCREENED             = "screened"
    SHORTLISTED          = "shortlisted"
    INTERVIEW_SCHEDULED  = "interview_scheduled"
    INTERVIEWED          = "interviewed"
    SELECTED             = "selected"
    REJECTED             = "rejected"
    ONBOARDED            = "onboarded"

class TaskPriority(str, Enum):
    LOW    = "low"
    MEDIUM = "medium"
    HIGH   = "high"

class ActivityType(str, Enum):
    CALL          = "call"
    EMAIL         = "email"
    MEETING       = "meeting"
    NOTE          = "note"
    STATUS_CHANGE = "status_change"
    INTERVIEW     = "interview"


# ============================================================
# PYDANTIC MODELS
# ============================================================
class UserCreate(BaseModel):
    email:    EmailStr
    password: str
    name:     str
    role:     str = "sales"   # default role for new users

class UserLogin(BaseModel):
    email:    EmailStr
    password: str

class LeadCreate(BaseModel):
    # Company (primary entity)
    company:            str
    company_type:       Optional[str]   = None
    company_linkedin:   Optional[str]   = None
    hq_location:        Optional[str]   = None
    india_office:       Optional[str]   = None
    segment:            Optional[str]   = None
    domain_focus:       Optional[str]   = None
    website:            Optional[str]   = None
    source:             Optional[str]   = None
    status:             Optional[str]   = "new"
    notes:              Optional[str]   = None
    next_follow_up:     Optional[str]   = None
    assigned_owner_id:  Optional[str]   = None
    deal_value:         Optional[float] = None
    # Legacy / outreach tracking
    source_file:                    Optional[str]   = None
    industry:                       Optional[str]   = None
    business_type:                  Optional[str]   = None
    address:                        Optional[str]   = None
    country:                        Optional[str]   = None
    turnover_headcount:             Optional[str]   = None
    intro_sent:                     Optional[str]   = None
    linkedin_invite_sent:           Optional[bool]  = None
    linkedin_invite_accepted:       Optional[bool]  = None
    lead_share_date:                Optional[str]   = None
    solution_skills:                Optional[str]   = None
    # Contact person 1
    full_name:                      Optional[str]   = None
    job_title:                      Optional[str]   = None
    email:                          Optional[str]   = None
    phone:                          Optional[str]   = None
    linkedin_url:                   Optional[str]   = None
    # Contact person 2
    contact_person_2_name:          Optional[str]   = None
    contact_person_2_designation:   Optional[str]   = None
    contact_person_2_email:         Optional[str]   = None
    contact_person_2_phone:         Optional[str]   = None
    contact_person_2_linkedin:      Optional[str]   = None
    # Contact person 3
    contact_person_3_name:          Optional[str]   = None
    contact_person_3_designation:   Optional[str]   = None
    contact_person_3_email:         Optional[str]   = None
    contact_person_3_phone:         Optional[str]   = None
    contact_person_3_linkedin:      Optional[str]   = None

    @field_validator('status', mode='before')
    @classmethod
    def coerce_status(cls, v):
        valid = {"new","contacted","called","interested","closed","completed","rejected","lost","follow_up_needed"}
        if not v or str(v).strip() not in valid:
            return "new"
        return str(v).strip()

    @field_validator('email', mode='before')
    @classmethod
    def clean_email(cls, v):
        if not v or not str(v).strip():
            return None
        val = str(v).strip()
        if "@" not in val or "." not in val.split("@")[-1]:
            return None
        return val

    @model_validator(mode='after')
    def ensure_full_name_from_company(self):
        # full_name DB column is NOT NULL - default to company for company-only entries
        if not self.full_name or not self.full_name.strip():
            self.full_name = self.company or "Unknown"
        # Back-fill aliases so both old and new column names work
        if not self.hq_location and self.address:
            self.hq_location = self.address
        if not self.domain_focus and self.industry:
            self.domain_focus = self.industry
        return self

class LeadUpdate(BaseModel):
    # Company
    company:            Optional[str]   = None
    company_type:       Optional[str]   = None
    company_linkedin:   Optional[str]   = None
    hq_location:        Optional[str]   = None
    india_office:       Optional[str]   = None
    segment:            Optional[str]   = None
    domain_focus:       Optional[str]   = None
    website:            Optional[str]   = None
    source:             Optional[str]   = None
    status:             Optional[str]   = None
    notes:              Optional[str]   = None
    next_follow_up:     Optional[str]   = None
    assigned_owner_id:  Optional[str]   = None
    deal_value:         Optional[float] = None
    # Legacy
    source_file:                    Optional[str]   = None
    industry:                       Optional[str]   = None
    business_type:                  Optional[str]   = None
    address:                        Optional[str]   = None
    country:                        Optional[str]   = None
    turnover_headcount:             Optional[str]   = None
    intro_sent:                     Optional[str]   = None
    linkedin_invite_sent:           Optional[bool]  = None
    linkedin_invite_accepted:       Optional[bool]  = None
    lead_share_date:                Optional[str]   = None
    solution_skills:                Optional[str]   = None
    # Contact person 1
    full_name:                      Optional[str]   = None
    job_title:                      Optional[str]   = None
    email:                          Optional[str]   = None
    phone:                          Optional[str]   = None
    linkedin_url:                   Optional[str]   = None
    # Contact person 2
    contact_person_2_name:          Optional[str]   = None
    contact_person_2_designation:   Optional[str]   = None
    contact_person_2_email:         Optional[str]   = None
    contact_person_2_phone:         Optional[str]   = None
    contact_person_2_linkedin:      Optional[str]   = None
    # Contact person 3
    contact_person_3_name:          Optional[str]   = None
    contact_person_3_designation:   Optional[str]   = None
    contact_person_3_email:         Optional[str]   = None
    contact_person_3_phone:         Optional[str]   = None
    contact_person_3_linkedin:      Optional[str]   = None

class ActivityCreate(BaseModel):
    lead_id:       Optional[str]  = None
    candidate_id:  Optional[str]  = None
    activity_type: ActivityType
    description:   str

class TaskCreate(BaseModel):
    title:        str
    description:  Optional[str]      = None
    task_type:    Optional[str]      = "note"
    due_date:     str
    due_time:     Optional[str]      = None
    priority:     TaskPriority       = TaskPriority.MEDIUM
    lead_id:      Optional[str]      = None
    candidate_id: Optional[str]      = None
    job_id:       Optional[str]      = None

class TaskUpdate(BaseModel):
    title:       Optional[str]          = None
    description: Optional[str]          = None
    due_date:    Optional[str]          = None
    priority:    Optional[TaskPriority] = None
    completed:   Optional[bool]         = None

class ReminderCreate(BaseModel):
    title:        str
    note:         Optional[str] = None
    due_date:     str
    due_time:     Optional[str] = None
    repeat_type:  str           = "none"
    email_alert:  bool          = False
    lead_id:      Optional[str] = None
    candidate_id: Optional[str] = None

class JobCreate(BaseModel):
    title:             str
    department:        Optional[str]  = None
    location:          Optional[str]  = None
    employment_type:   Optional[str]  = None
    description:       Optional[str]  = None
    requirements:      Optional[str]  = None
    salary_range:      Optional[str]  = None
    skills:            List[str]      = []
    is_active:         bool           = True
    is_urgent:         bool           = False
    post_to_linkedin:  bool           = False   # trigger LinkedIn company page post

class JobUpdate(BaseModel):
    title:           Optional[str]  = None
    department:      Optional[str]  = None
    location:        Optional[str]  = None
    employment_type: Optional[str]  = None
    description:     Optional[str]  = None
    requirements:    Optional[str]  = None
    salary_range:    Optional[str]  = None
    skills:          Optional[List[str]] = None
    is_active:       Optional[bool] = None
    is_urgent:       Optional[bool] = None

class CandidateCreate(BaseModel):
    full_name:        str
    email:            Optional[str]  = None
    phone:            Optional[str]  = None
    current_company:  Optional[str]  = None
    candidate_role:   Optional[str]  = None
    experience_years: Optional[int]  = None
    source:           Optional[str]  = None
    job_id:           Optional[str]  = None
    status:           Optional[str]  = "sourced"
    notes:            Optional[str]  = None
    resume_url:       Optional[str]  = None
    linkedin_url:     Optional[str]  = None
    portfolio_url:    Optional[str]  = None
    skills:           List[str]      = []
    # Extended fields v3
    candidate_type:       Optional[str]  = "domestic"  # domestic | international
    visa_status:          Optional[str]  = None
    total_experience:     Optional[str]  = None
    relevant_experience:  Optional[str]  = None
    location:             Optional[str]  = None
    relocation:           Optional[str]  = None
    # v4 fields
    work_mode:            List[str]      = []   # e.g. ["Remote","Hybrid"]
    current_ctc:          Optional[float] = None  # in LPA / numeric
    expected_ctc:         Optional[float] = None
    # LLM-extracted fields (populated on resume upload)
    tech_stack:           List[str]      = []

    @field_validator("status", mode="before")
    @classmethod
    def coerce_cand_status(cls, v):
        valid = {"sourced","screened","shortlisted","interview_scheduled","interviewed","selected","rejected","onboarded"}
        if not v or str(v).strip() not in valid:
            return "sourced"
        return str(v).strip()

    @field_validator("email", mode="before")
    @classmethod
    def clean_cand_email(cls, v):
        if not v or not str(v).strip(): return None
        val = str(v).strip()
        if "@" not in val or "." not in val.split("@")[-1]: return None
        return val

class CandidateUpdate(BaseModel):
    full_name:            Optional[str]          = None
    email:                Optional[str]          = None
    phone:                Optional[str]          = None
    current_company:      Optional[str]          = None
    candidate_role:       Optional[str]          = None
    experience_years:     Optional[int]          = None
    source:               Optional[str]          = None
    job_id:               Optional[str]          = None
    status:               Optional[str]          = None
    notes:                Optional[str]          = None
    resume_url:           Optional[str]          = None
    linkedin_url:         Optional[str]          = None
    portfolio_url:        Optional[str]          = None
    skills:               Optional[List[str]]    = None
    assigned_recruiter_id:Optional[str]          = None
    # Extended fields v3
    candidate_type:       Optional[str]  = None
    visa_status:          Optional[str]  = None
    total_experience:     Optional[str]  = None
    relevant_experience:  Optional[str]  = None
    location:             Optional[str]  = None
    relocation:           Optional[str]  = None
    # v4 fields
    work_mode:            Optional[List[str]]  = None
    current_ctc:          Optional[float]      = None
    expected_ctc:         Optional[float]      = None
    # LLM-extracted fields
    tech_stack:           Optional[List[str]] = None

class SubmissionCreate(BaseModel):
    lead_id:      str
    candidate_id: str
    status:       str = "submitted"
    notes:        Optional[str] = None

class SubmissionUpdate(BaseModel):
    status: Optional[str] = None
    notes:  Optional[str] = None

class InterviewCreate(BaseModel):
    candidate_id:   str
    job_id:         str
    scheduled_at:   str           # ISO datetime
    interview_type: str
    interviewers:   List[str] = []
    notes:          Optional[str] = None

class InterviewUpdate(BaseModel):
    scheduled_at:   Optional[str]      = None
    interview_type: Optional[str]      = None
    interviewers:   Optional[List[str]]= None
    notes:          Optional[str]      = None
    feedback:       Optional[str]      = None
    rating:         Optional[int]      = None
    completed:      Optional[bool]     = None


# ============================================================
# HELPERS
# ============================================================
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(plain: str, hashed: str) -> bool:
    return bcrypt.checkpw(plain.encode(), hashed.encode())

def create_access_token(user_id: str, email: str) -> str:
    return jwt.encode(
        {"sub": user_id, "email": email,
         "exp": datetime.now(timezone.utc) + timedelta(hours=24),
         "type": "access"},
        JWT_SECRET, algorithm=JWT_ALGORITHM
    )

def create_refresh_token(user_id: str) -> str:
    return jwt.encode(
        {"sub": user_id,
         "exp": datetime.now(timezone.utc) + timedelta(days=7),
         "type": "refresh"},
        JWT_SECRET, algorithm=JWT_ALGORITHM
    )

async def get_current_user(request: Request) -> dict:
    token = request.cookies.get("access_token")
    if not token:
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
    if not token:
        raise HTTPException(401, "Not authenticated")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        if payload.get("type") != "access":
            raise HTTPException(401, "Invalid token type")
    except jwt.ExpiredSignatureError:
        raise HTTPException(401, "Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(401, "Invalid token")

    user = await safe_single(
        lambda: supabase.table("users")
            .select("id,email,name,role,avatar_url,created_at")
            .eq("id", payload["sub"])
            .single()
            .execute()
    )
    if not user:
        raise HTTPException(401, "User not found")
    return user

# ── Module-level role guard ───────────────────────────────────
# Call after get_current_user() on any sales or recruitment endpoint.
# Workers are timesheet-only; this blocks them from all other data.
_MODULE_ROLES: dict = {
    "sales":       {"admin", "sales", "viewer"},
    "recruitment": {"admin", "sales", "viewer"},
}

def _require_module(user: dict, module: str) -> None:
    """Raise 403 if user's role is not permitted to access this module."""
    allowed = _MODULE_ROLES.get(module, set())
    role = user.get("role", "")
    if role not in allowed:
        raise HTTPException(
            status_code=403,
            detail=f"Your role '{role}' does not have access to the {module} module.",
        )


async def send_email(to: str, subject: str, html: str):
    if not resend.api_key:
        logger.info(f"[email-sim] To: {to} | {subject}")
        return {"status": "simulated"}
    try:
        result = await asyncio.to_thread(
            resend.Emails.send, {"from": SENDER_EMAIL, "to": [to], "subject": subject, "html": html}
        )
        return {"status": "success", "id": result.get("id")}
    except Exception as e:
        logger.error(f"Email failed: {e}")
        return {"status": "failed", "error": str(e)}

def sb(table: str):
    """Shorthand for supabase.table()"""
    return supabase.table(table)

def get_supabase() -> Client:
    """Return the shared Supabase client used by route handlers."""
    return supabase

async def run(fn):
    """Run a synchronous supabase call in a thread pool."""
    return await asyncio.to_thread(fn)

async def safe_single(fn):
    """Run a .single() query safely — returns None on 0 rows (PGRST116) instead of crashing."""
    try:
        res = await asyncio.to_thread(fn)
        return res.data
    except Exception as e:
        err = str(e)
        if "PGRST116" in err or "0 rows" in err or "406" in err:
            return None
        raise



# ============================================================
# AUDIT LOG HELPERS
# ============================================================
def _get_ip(request: Request) -> str:
    xff = request.headers.get("x-forwarded-for", "")
    if xff:
        return xff.split(",")[0].strip()
    return request.client.host if request.client else "unknown"

async def _audit(
    action: str,
    user: dict = None,
    entity_type: str = None,
    entity_id:   str = None,
    entity_name: str = None,
    old_value:   dict = None,
    new_value:   dict = None,
    ip:          str = None,
    ua:          str = None,
):
    """Fire-and-forget audit log entry. Never raises."""
    try:
        await run(lambda: sb("audit_logs").insert({
            "user_id":     user["id"]    if user else None,
            "user_email":  user["email"] if user else None,
            "user_name":   user["name"]  if user else None,
            "action":      action,
            "entity_type": entity_type,
            "entity_id":   str(entity_id) if entity_id else None,
            "entity_name": entity_name,
            "old_value":   old_value,
            "new_value":   new_value,
            "ip_address":  ip,
            "user_agent":  ua,
        }).execute())
    except Exception as e:
        logger.warning(f"Audit log write failed: {e}")


# ============================================================
# AUTH
# ============================================================
@api_router.post("/auth/register")
async def register(data: UserCreate, request: Request):
    # Only logged-in admins can create new users
    caller = await get_current_user(request)
    if caller.get("role") != "admin":
        raise HTTPException(403, "Only admins can create new user accounts")

    # Validate role value
    allowed_roles = {"admin", "sales", "viewer", "worker"}
    role = data.role if data.role in allowed_roles else "sales"

    existing = await run(lambda: sb("users").select("id").eq("email", data.email.lower()).execute())
    if existing.data:
        raise HTTPException(400, "Email already registered")

    user_id = str(uuid.uuid4())
    await run(lambda: sb("users").insert({
        "id":            user_id,
        "email":         data.email.lower(),
        "password_hash": hash_password(data.password),
        "name":          data.name,
        "role":          role,
    }).execute())

    return {"id": user_id, "email": data.email, "name": data.name, "role": role}


@api_router.post("/auth/login")
async def login(data: UserLogin, response: Response, request: Request):
    user = await safe_single(lambda: sb("users").select("*").eq("email", data.email.lower()).single().execute())
    if not user:
        await _audit("login_failed", entity_name=data.email,
                     ip=_get_ip(request),
                     ua=request.headers.get("user-agent", ""))
        raise HTTPException(401, "Access not authorized. This email is not registered.")
    if not verify_password(data.password, user["password_hash"]):
        await _audit("login_failed", entity_name=user["email"],
                     ip=_get_ip(request),
                     ua=request.headers.get("user-agent", ""))
        raise HTTPException(401, "Incorrect password. Please try again.")

    access  = create_access_token(user["id"], user["email"])
    refresh = create_refresh_token(user["id"])
    _set_cookies(response, access, refresh)
    await _audit("login", user=user,
                 ip=_get_ip(request),
                 ua=request.headers.get("user-agent", ""))
    return {
        "id":    user["id"],
        "email": user["email"],
        "name":  user["name"],
        "role":  user.get("role", "viewer"),
    }


@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    try:
        user = await get_current_user(request)
        asyncio.create_task(_audit("logout", user=user,
                                    ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    except Exception:
        pass
    response.delete_cookie("access_token",  path="/")
    response.delete_cookie("refresh_token", path="/")
    return {"message": "Logged out"}


@api_router.get("/auth/me")
async def get_me(request: Request):
    return await get_current_user(request)


def _set_cookies(response: Response, access: str, refresh: str):
    kw = dict(httponly=True, secure=True, samesite="none", path="/")
    response.set_cookie("access_token",  access,  max_age=86400,   **kw)
    response.set_cookie("refresh_token", refresh, max_age=604800,  **kw)


# ============================================================
# USERS
# ============================================================
@api_router.get("/users")
async def get_users(request: Request):
    await get_current_user(request)
    res = await run(lambda: sb("users").select("id,email,name,role,avatar_url,created_at").execute())
    return res.data or []


@api_router.put("/users/{user_id}/role")
async def update_user_role(user_id: str, body: dict, request: Request):
    caller = await get_current_user(request)
    if caller.get("role") != "admin":
        raise HTTPException(403, "Only admins can change user roles")
    allowed_roles = {"admin", "sales", "viewer", "worker"}
    role = body.get("role")
    if role not in allowed_roles:
        raise HTTPException(400, f"Role must be one of: {', '.join(allowed_roles)}")
    target = await run(lambda: sb("users").select("name,email,role").eq("id", user_id).execute())
    old_role = (target.data or [{}])[0].get("role")
    target_name = (target.data or [{}])[0].get("name") or user_id
    await run(lambda: sb("users").update({"role": role}).eq("id", user_id).execute())
    asyncio.create_task(_audit("update", user=caller, entity_type="user", entity_id=user_id,
                                entity_name=target_name,
                                old_value={"role": old_role}, new_value={"role": role},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"id": user_id, "role": role}


@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, request: Request):
    caller = await get_current_user(request)
    if caller.get("role") != "admin":
        raise HTTPException(403, "Only admins can delete users")
    if caller["id"] == user_id:
        raise HTTPException(400, "You cannot delete your own account")
    target = await run(lambda: sb("users").select("name,email").eq("id", user_id).execute())
    target_name = (target.data or [{}])[0].get("name") or user_id
    await run(lambda: sb("users").delete().eq("id", user_id).execute())
    asyncio.create_task(_audit("delete", user=caller, entity_type="user", entity_id=user_id,
                                entity_name=target_name,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "User deleted"}


# ============================================================
# LEADS
# ============================================================
@api_router.post("/leads")
async def create_lead(lead: LeadCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    doc = {
        # Company
        "company":           lead.company,
        "company_type":      lead.company_type,
        "company_linkedin":  lead.company_linkedin,
        "hq_location":       lead.hq_location or lead.address,
        "india_office":      lead.india_office,
        "segment":           lead.segment,
        "domain_focus":      lead.domain_focus or lead.industry,
        "website":           lead.website,
        "source":            lead.source,
        "status":            lead.status or "new",
        "notes":             lead.notes,
        "next_follow_up":    lead.next_follow_up,
        "deal_value":        lead.deal_value,
        # Legacy fields (kept for DB compat)
        "source_file":                  lead.source_file,
        "industry":                     lead.industry or lead.domain_focus,
        "business_type":                lead.business_type,
        "address":                      lead.address or lead.hq_location,
        "country":                      lead.country,
        "turnover_headcount":           lead.turnover_headcount,
        "intro_sent":                   lead.intro_sent,
        "linkedin_invite_sent":         lead.linkedin_invite_sent,
        "linkedin_invite_accepted":     lead.linkedin_invite_accepted,
        "lead_share_date":              lead.lead_share_date,
        "solution_skills":              lead.solution_skills,
        # Contact person 1
        "full_name":         lead.full_name or lead.company,
        "job_title":         lead.job_title,
        "email":             lead.email,
        "phone":             lead.phone,
        "linkedin_url":      lead.linkedin_url,
        # Contact person 2
        "contact_person_2_name":        lead.contact_person_2_name,
        "contact_person_2_designation": lead.contact_person_2_designation,
        "contact_person_2_email":       lead.contact_person_2_email,
        "contact_person_2_phone":       lead.contact_person_2_phone,
        "contact_person_2_linkedin":    lead.contact_person_2_linkedin,
        # Contact person 3
        "contact_person_3_name":        lead.contact_person_3_name,
        "contact_person_3_designation": lead.contact_person_3_designation,
        "contact_person_3_email":       lead.contact_person_3_email,
        "contact_person_3_phone":       lead.contact_person_3_phone,
        "contact_person_3_linkedin":    lead.contact_person_3_linkedin,
        # Audit
        "assigned_owner_id": lead.assigned_owner_id or user["id"],
        "created_by":        user["id"],
    }
    # Strip None values - avoids inserting NULL for columns that may not exist yet
    doc = {k: v for k, v in doc.items() if v is not None}

    try:
        res = await run(lambda: sb("leads").insert(doc).execute())
    except Exception as e:
        err_str = str(e)
        # PGRST204 = column not found in schema cache (migration not yet applied)
        # Retry after removing the offending column so other fields still save
        if "PGRST204" in err_str:
            import re as _re
            col_match = _re.search(r"Could not find the '(\w+)' column", err_str)
            if col_match:
                bad_col = col_match.group(1)
                logger.warning(f"[create_lead] Column '{bad_col}' missing in DB — skipping it. Run add_features_v3.sql to add it.")
                doc.pop(bad_col, None)
                res = await run(lambda: sb("leads").insert(doc).execute())
            else:
                raise
        else:
            raise

    lead_id = res.data[0]["id"]
    await _log_activity(lead_id=lead_id, user=user, atype="note", desc=f"Lead created by {user['name']}")
    asyncio.create_task(_audit("create", user=user, entity_type="lead", entity_id=lead_id,
                                entity_name=res.data[0].get("full_name") or res.data[0].get("company"),
                                new_value={"company": res.data[0].get("company"), "status": res.data[0].get("status")},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return res.data[0]


@api_router.get("/leads")
async def get_leads(
    request:     Request,
    status:      Optional[str] = None,
    source:      Optional[str] = None,
    search:      Optional[str] = None,
    assigned_to: Optional[str] = None,
    skip:        int = 0,
    limit:       int = 50,
):
    user = await get_current_user(request)
    _require_module(user, "sales")
    q = sb("leads").select(
        "*, assigned_owner:assigned_owner_id(name)",
        count="exact"
    ).order("created_at", desc=True).range(skip, skip + limit - 1)

    if status:      q = q.eq("status", status)
    if source:      q = q.eq("source", source)
    if assigned_to: q = q.eq("assigned_owner_id", assigned_to)
    if search:
        q = q.or_(
            f"full_name.ilike.%{search}%,"
            f"email.ilike.%{search}%,"
            f"company.ilike.%{search}%,"
            f"phone.ilike.%{search}%"
        )

    res = await run(lambda: q.execute())
    return {"leads": res.data or [], "total": res.count or 0}


@api_router.get("/leads/{lead_id}")
async def get_lead(lead_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    lead = await safe_single(lambda: sb("leads").select("*, assigned_owner:assigned_owner_id(name), submissions:candidate_submissions(*, candidate:candidate_id(id,full_name,candidate_role,status))").eq("id", lead_id).single().execute())
    if not lead:
        raise HTTPException(404, "Lead not found")

    acts = await run(lambda: sb("activities").select("*").eq("lead_id", lead_id).order("created_at", desc=True).execute())
    lead["activities"] = acts.data or []

    hist = await run(lambda: sb("lead_status_history").select("*").eq("lead_id", lead_id).order("created_at", desc=True).execute())
    lead["status_history"] = hist.data or []

    asyncio.create_task(_audit("view", user=user, entity_type="lead", entity_id=lead_id,
                                entity_name=lead.get("full_name"), ip=_get_ip(request),
                                ua=request.headers.get("user-agent","")))
    return lead


@api_router.put("/leads/{lead_id}")
async def update_lead(lead_id: str, lead: LeadUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")

    existing = await safe_single(lambda: sb("leads").select("status").eq("id", lead_id).single().execute())
    if not existing:
        raise HTTPException(404, "Lead not found")

    patch = {k: v for k, v in lead.model_dump().items() if v is not None}
    if "status" in patch and isinstance(patch["status"], LeadStatus):
        patch["status"] = patch["status"].value

    if lead.status and lead.status.value != existing.get("status"):
        old_status = existing.get("status")
        new_status = lead.status.value
        await run(lambda: sb("lead_status_history").insert({
            "lead_id":        lead_id,
            "old_status":     old_status,
            "new_status":     new_status,
            "changed_by":     user["id"],
            "changed_by_name":user["name"],
        }).execute())
        await _log_activity(lead_id=lead_id, user=user, atype="status_change",
                            desc=f"Status changed from {old_status} to {new_status}")

    old = await safe_single(lambda: sb("leads").select("*").eq("id", lead_id).single().execute())
    await run(lambda: sb("leads").update(patch).eq("id", lead_id).execute())
    asyncio.create_task(_audit("update", user=user, entity_type="lead", entity_id=lead_id,
                                entity_name=old.get("full_name") if old else None,
                                old_value={k:old.get(k) for k in patch if old} if old else None,
                                new_value=patch, ip=_get_ip(request),
                                ua=request.headers.get("user-agent","")))
    return {"message": "Lead updated"}


@api_router.delete("/leads/{lead_id}")
async def delete_lead(lead_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    lead = await safe_single(lambda: sb("leads").select("full_name").eq("id", lead_id).single().execute())
    await run(lambda: sb("leads").delete().eq("id", lead_id).execute())
    asyncio.create_task(_audit("delete", user=user, entity_type="lead", entity_id=lead_id,
                                entity_name=lead.get("full_name") if lead else None,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Lead deleted"}


# ============================================================
# CSV / EXCEL IMPORT
# ============================================================
@api_router.post("/leads/import")
async def import_leads(file: UploadFile, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    if not file.filename:
        raise HTTPException(400, "No file provided")
    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ("csv", "xlsx", "xls"):
        raise HTTPException(400, "Only CSV / Excel files are supported")

    content = await file.read()
    try:
        if ext == "csv":
            df = pd.read_csv(io.BytesIO(content))
        else:
            df = pd.read_excel(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(400, f"Error reading file: {e}")

    # 1. Normalize column names
    # Strip whitespace, lowercase, then apply mapping.
    # Pandas auto-renames duplicate column names with .1, .2 suffixes.
    raw_cols = list(df.columns)
    df.columns = df.columns.str.lower().str.strip()

    col_map = {
        # Company fields
        "company name":     "company",
        "company":          "company",
        "organisation":     "company",
        "organization":     "company",
        "account":          "company",
        # Company meta
        "type":             "company_type",
        "company type":     "company_type",
        "segment":          "segment",
        "location":         "hq_location",
        "hq location":      "hq_location",
        "irish hq":         "hq_location",
        "india office(s)":  "india_office",
        "india offices":    "india_office",
        "domain focus":     "domain_focus",
        "core services":    "domain_focus",
        "industry":         "domain_focus",
        "website":          "website",
        # Company LinkedIn - the company page (not a person)
        "linkedn":          "company_linkedin",
        "linkedin":         "company_linkedin",
        "linkedin url":     "company_linkedin",
        "company linkedin": "company_linkedin",
        # Status / tracking
        "status":           "status",
        "lead status":      "status",
        "remark":           "notes",
        "remarks":          "notes",
        "remark ":          "notes",
        "follow up":        "next_follow_up",
        "follow-up":        "next_follow_up",
        "next f date":      "next_follow_up",
        "next follow date": "next_follow_up",
        "intro sent":       "intro_sent",
        "f- date":          "intro_sent",
        "source":           "source",
        "lead come from":   "source",
        "data form":        "source_file",
        "turnover/ headcount": "turnover_headcount",
        "turnover/headcount":  "turnover_headcount",
        "solution using/looking skills": "solution_skills",
        "linkedin invite sent":     "linkedin_invite_sent",
        "linkedin invite accepted": "linkedin_invite_accepted",
        "in crm":           "_skip",
        "sr. no":           "_skip",
        "no.":              "_skip",
        "sl. no.":          "_skip",

        # Contact person 1
        "contact person-1": "full_name",
        "contact person 1": "full_name",
        "name":             "full_name",
        "full name":        "full_name",
        "fullname":         "full_name",
        # Designation / title (first occurrence)
        "designtaion":      "job_title",
        "designation":      "job_title",
        "title":            "job_title",
        "job title":        "job_title",
        # Email (first occurrence)
        "e-mail":           "email",
        "email":            "email",
        "email address":    "email",
        # Phone (first occurrence)
        "mobile number":    "phone",
        "mobile":           "phone",
        "phone":            "phone",
        "contact":          "phone",
        "contact no. ":     "phone",
        # LinkedIn (first occurrence) - person LinkedIn
        "linkden":          "linkedin_url",
        "link":             "linkedin_url",
        "person linkedin url": "linkedin_url",

        # Contact person 2 (pandas renames duplicate cols with .1 suffix)
        "contact person-2":     "contact_person_2_name",
        "contact person 2":     "contact_person_2_name",
        "contact person 1.1":   "contact_person_2_name",
        "designtaion.1":        "contact_person_2_designation",
        "designation.1":        "contact_person_2_designation",
        "e-mail.1":             "contact_person_2_email",
        "email.1":              "contact_person_2_email",
        "email 1":              "contact_person_2_email",
        "mobile number.1":      "contact_person_2_phone",
        "mobile.1":             "contact_person_2_phone",
        "contact person 1 mobile": "contact_person_2_phone",
        "linkden.1":            "contact_person_2_linkedin",
        "link.1":               "contact_person_2_linkedin",

        # Contact person 3
        "contact person-3":     "contact_person_3_name",
        "contact person 3":     "contact_person_3_name",
        "contact person 1.2":   "contact_person_3_name",
        "designtaion.2":        "contact_person_3_designation",
        "designation.2":        "contact_person_3_designation",
        "e-mail.2":             "contact_person_3_email",
        "email.2":              "contact_person_3_email",
        "email 2":              "contact_person_3_email",
        "mobile number.2":      "contact_person_3_phone",
        "mobile.2":             "contact_person_3_phone",
        "contact person 2 mobile": "contact_person_3_phone",
        "linkden.2":            "contact_person_3_linkedin",
        "link.2":               "contact_person_3_linkedin",

        # Apollo / LinkedIn CSV format
        "first name":           "_first_name",
        "firstname":            "_first_name",
        "last name":            "_last_name",
        "lastname":             "_last_name",
        "work email":           "email",
        "work direct phone":    "phone",
        "corporate phone":      "_phone2",
        "person linkedin url":  "linkedin_url",
        "company name for emails": "_company2",
        "stage":                "_stage",
        "seniority":            "_seniority",
    }

    df = df.rename(columns=col_map)

    # Remove columns we want to skip
    df = df.drop(columns=[c for c in df.columns if c == "_skip"], errors="ignore")

    # 2. Handle Apollo first/last name split
    if "full_name" not in df.columns and "_first_name" in df.columns:
        first = df.get("_first_name", pd.Series([""] * len(df))).fillna("")
        last  = df.get("_last_name",  pd.Series([""] * len(df))).fillna("")
        df["full_name"] = (first + " " + last).str.strip()

    # 3. Phone fallback
    if "phone" in df.columns:
        for fallback in ["_phone2", "_phone3"]:
            if fallback in df.columns:
                df["phone"] = df["phone"].fillna(df[fallback]).replace("", None)

    # 4. Company fallback
    if "_company2" in df.columns:
        if "company" not in df.columns:
            df["company"] = df["_company2"]
        else:
            df["company"] = df["company"].fillna(df["_company2"])

    # 5. Detect sheet segment from filename
    fname_lower = file.filename.lower()
    auto_segment = None
    if "staffing" in fname_lower:
        auto_segment = "staffing_partner"
    elif "end client" in fname_lower or "end_client" in fname_lower:
        auto_segment = "end_client"
    elif "ireland" in fname_lower:
        auto_segment = "ireland_company"
    elif "spoc" in fname_lower:
        auto_segment = "general"

    import_id  = str(uuid.uuid4())
    successful = 0
    errors:    list = []
    batch:     list = []

    for idx, row in df.iterrows():
        company = str(row.get("company", "")).strip()
        full_name = str(row.get("full_name", "")).strip()

        # Skip empty rows
        if (not company or company.lower() == "nan") and            (not full_name or full_name.lower() == "nan"):
            errors.append({"row": idx + 2, "error": "Missing company name"})
            continue

        # For company-centric import, company is primary
        if not company or company.lower() == "nan":
            company = full_name
        if not full_name or full_name.lower() == "nan":
            full_name = company

        def c(key):
            # Clean a cell value to None or stripped string.
            val = str(row.get(key, "")).strip()
            return None if val.lower() in ("nan", "", "none") else val

        def cb(key):
            # Clean boolean cell.
            val = str(row.get(key, "")).strip().lower()
            if val in ("yes", "true", "1", "y"):
                return True
            if val in ("no", "false", "0", "n"):
                return False
            return None

        def cd(key):
            # Clean date cell.
            val = str(row.get(key, "")).strip()
            if val.lower() in ("nan", "", "none"):
                return None
            try:
                import re
                # Accept YYYY-MM-DD or DD/MM/YYYY etc.
                val = re.sub(r"[/\\]", "-", val)
                parts = val.split("-")
                if len(parts) == 3 and len(parts[0]) == 2:
                    # Likely DD-MM-YYYY
                    val = f"{parts[2]}-{parts[1]}-{parts[0]}"
                return val[:10]
            except Exception:
                return None

        segment = c("segment") or auto_segment
        hq = c("hq_location") or c("address") or c("country")

        batch.append({
            # Company
            "company":           company,
            "company_type":      c("company_type"),
            "company_linkedin":  c("company_linkedin"),
            "hq_location":       hq,
            "india_office":      c("india_office"),
            "segment":           segment,
            "domain_focus":      c("domain_focus") or c("industry"),
            "website":           c("website"),
            "source":            c("source") or "Import",
            "source_file":       f"{file.filename}",
            "status":            "new",
            "notes":             c("notes"),
            "next_follow_up":    cd("next_follow_up"),
            # Legacy aliases kept for DB completeness
            "address":           hq,
            "country":           c("country"),
            "industry":          c("domain_focus") or c("industry"),
            "business_type":     c("business_type"),
            "turnover_headcount":c("turnover_headcount"),
            "solution_skills":   c("solution_skills"),
            "intro_sent":        cd("intro_sent"),
            "linkedin_invite_sent":     cb("linkedin_invite_sent"),
            "linkedin_invite_accepted": cb("linkedin_invite_accepted"),
            "lead_share_date":   cd("lead_share_date"),
            # CP1
            "full_name":         full_name,
            "job_title":         c("job_title"),
            "email":             c("email"),
            "phone":             c("phone"),
            "linkedin_url":      c("linkedin_url"),
            # CP2
            "contact_person_2_name":        c("contact_person_2_name"),
            "contact_person_2_designation": c("contact_person_2_designation"),
            "contact_person_2_email":       c("contact_person_2_email"),
            "contact_person_2_phone":       c("contact_person_2_phone"),
            "contact_person_2_linkedin":    c("contact_person_2_linkedin"),
            # CP3
            "contact_person_3_name":        c("contact_person_3_name"),
            "contact_person_3_designation": c("contact_person_3_designation"),
            "contact_person_3_email":       c("contact_person_3_email"),
            "contact_person_3_phone":       c("contact_person_3_phone"),
            "contact_person_3_linkedin":    c("contact_person_3_linkedin"),
            # Audit
            "assigned_owner_id": user["id"],
            "created_by":        user["id"],
            "import_id":         import_id,
        })

    # 6. Bulk insert with per-row error capture
    SKIP_KEYS_ON_MISSING = {
        "company_type", "company_linkedin", "hq_location", "india_office",
        "segment", "domain_focus", "contact_person_2_linkedin",
        "contact_person_3_linkedin",
    }

    for record in batch:
        try:
            await run(lambda r=record: sb("leads").insert(r).execute())
            successful += 1
        except Exception as e:
            err_str = str(e)
            # If new columns don't exist yet (migration not run), retry without them
            if "column" in err_str.lower() and "does not exist" in err_str.lower():
                fallback = {k: v for k, v in record.items() if k not in SKIP_KEYS_ON_MISSING}
                try:
                    await run(lambda r=fallback: sb("leads").insert(r).execute())
                    successful += 1
                    logger.warning("[import_leads] New columns missing - run add_sales_v4.sql")
                    continue
                except Exception as e2:
                    errors.append({"row": "?", "error": str(e2)})
            else:
                errors.append({"row": "?", "error": err_str})

    await run(lambda: sb("imports").insert({
        "id":         import_id,
        "filename":   file.filename,
        "user_id":    user["id"],
        "total_rows": len(df),
        "successful": successful,
        "failed":     len(errors),
        "errors":     errors[:50],
    }).execute())

    asyncio.create_task(_audit("import", user=user, entity_type="lead",
                                entity_name=file.filename,
                                new_value={"total_rows": len(df), "successful": successful, "failed": len(errors)},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {
        "import_id":  import_id,
        "total_rows": len(df),
        "successful": successful,
        "failed":     len(errors),
        "errors":     errors[:20],
    }

def _clean(row, key):
    val = str(row.get(key, "")).strip()
    return None if val.lower() in ("nan", "") else val


@api_router.get("/imports")
async def get_imports(request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    res = await run(lambda: sb("imports").select("*").order("created_at", desc=True).limit(100).execute())
    return res.data or []


# ============================================================
# ACTIVITIES
# ============================================================
async def _log_activity(*, lead_id=None, candidate_id=None, user: dict, atype: str, desc: str):
    doc = {
        "lead_id":       lead_id,
        "candidate_id":  candidate_id,
        "user_id":       user["id"],
        "user_name":     user["name"],
        "activity_type": atype,
        "description":   desc,
    }
    await run(lambda: sb("activities").insert(doc).execute())


@api_router.post("/activities")
async def create_activity(activity: ActivityCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    await _log_activity(
        lead_id=activity.lead_id, candidate_id=activity.candidate_id,
        user=user, atype=activity.activity_type.value, desc=activity.description
    )
    return {"message": "Activity logged"}


@api_router.get("/activities")
async def get_activities(request: Request, lead_id: Optional[str] = None, candidate_id: Optional[str] = None):
    user = await get_current_user(request)
    _require_module(user, "sales")
    q = sb("activities").select("*").order("created_at", desc=True)
    if lead_id:      q = q.eq("lead_id", lead_id)
    if candidate_id: q = q.eq("candidate_id", candidate_id)
    res = await run(lambda: q.execute())
    return res.data or []


# ============================================================
# TASKS
# ============================================================
@api_router.post("/tasks")
async def create_task(task: TaskCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    res = await run(lambda: sb("tasks").insert({
        "title":           task.title,
        "description":     task.description,
        "task_type":       task.task_type,
        "due_date":        task.due_date,
        "due_time":        task.due_time,
        "priority":        task.priority.value,
        "lead_id":         task.lead_id,
        "candidate_id":    task.candidate_id,
        "job_id":          task.job_id,
        "assigned_to":     user["id"],
        "assigned_to_name":user["name"],
        "created_by":      user["id"],
    }).execute())
    return res.data[0]


@api_router.get("/tasks")
async def get_tasks(
    request:      Request,
    completed:    Optional[bool] = None,
    due_today:    Optional[bool] = None,
    overdue:      Optional[bool] = None,
    lead_id:      Optional[str]  = None,
    candidate_id: Optional[str]  = None,
):
    user = await get_current_user(request)
    today = datetime.now(timezone.utc).date().isoformat()

    q = sb("tasks").select("*").eq("assigned_to", user["id"]).order("due_date")
    if completed is not None: q = q.eq("completed", completed)
    if lead_id:               q = q.eq("lead_id", lead_id)
    if candidate_id:          q = q.eq("candidate_id", candidate_id)
    if due_today:             q = q.eq("due_date", today)
    if overdue:               q = q.lt("due_date", today).eq("completed", False)

    res = await run(lambda: q.execute())
    return res.data or []


@api_router.put("/tasks/{task_id}")
async def update_task(task_id: str, task: TaskUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    patch = {k: v for k, v in task.model_dump().items() if v is not None}
    if "priority" in patch and isinstance(patch["priority"], TaskPriority):
        patch["priority"] = patch["priority"].value
    if task.completed:
        patch["completed_at"] = datetime.now(timezone.utc).isoformat()
    await run(lambda: sb("tasks").update(patch).eq("id", task_id).execute())
    return {"message": "Task updated"}


@api_router.delete("/tasks/{task_id}")
async def delete_task(task_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    await run(lambda: sb("tasks").delete().eq("id", task_id).execute())
    return {"message": "Task deleted"}


# ============================================================
# REMINDERS
# ============================================================
@api_router.post("/reminders")
async def create_reminder(reminder: ReminderCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    res = await run(lambda: sb("reminders").insert({
        "title":        reminder.title,
        "note":         reminder.note,
        "due_date":     reminder.due_date,
        "due_time":     reminder.due_time,
        "repeat_type":  reminder.repeat_type,
        "email_alert":  reminder.email_alert,
        "lead_id":      reminder.lead_id,
        "candidate_id": reminder.candidate_id,
        "user_id":      user["id"],
        "user_email":   user["email"],
    }).execute())
    return res.data[0]


@api_router.get("/reminders")
async def get_reminders(request: Request, upcoming: Optional[bool] = None):
    user = await get_current_user(request)
    _require_module(user, "sales")
    q = sb("reminders").select("*").eq("user_id", user["id"]).order("due_date")
    if upcoming:
        q = q.gte("due_date", datetime.now(timezone.utc).date().isoformat())
    res = await run(lambda: q.execute())
    return res.data or []


@api_router.put("/reminders/{reminder_id}/dismiss")
async def dismiss_reminder(reminder_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    await run(lambda: sb("reminders").update({
        "dismissed": True,
        "dismissed_at": datetime.now(timezone.utc).isoformat()
    }).eq("id", reminder_id).execute())
    return {"message": "Reminder dismissed"}


@api_router.delete("/reminders/{reminder_id}")
async def delete_reminder(reminder_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    await run(lambda: sb("reminders").delete().eq("id", reminder_id).execute())
    return {"message": "Reminder deleted"}


# ============================================================
# JOBS
# ============================================================
@api_router.post("/jobs")
async def create_job(job: JobCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    apply_key = secrets.token_urlsafe(8)
    res = await run(lambda: sb("jobs").insert({
        "title":           job.title,
        "department":      job.department,
        "location":        job.location,
        "employment_type": job.employment_type,
        "description":     job.description,
        "requirements":    job.requirements,
        "salary_range":    job.salary_range,
        "skills":          job.skills,
        "is_active":       job.is_active,
        "is_urgent":       job.is_urgent,
        "created_by":      user["id"],
        # URL-safe random key — used by the public /apply?key=… form.
        # Generated here (not in the DB) so the value is visible in the response immediately.
        "apply_key":       apply_key,
    }).execute())
    created_job = res.data[0]

    # Optionally post to LinkedIn (non-blocking — job creation succeeds even if LinkedIn fails)
    linkedin_result = {"success": False, "error": "not_requested"}
    if job.post_to_linkedin:
        # Apply URL uses the frontend origin — stored as env var or derived from request
        frontend_origin = os.environ.get("FRONTEND_URL", "https://jmdata-crm-application.jmdatatalent.com")
        apply_url = f"{frontend_origin}/apply?key={apply_key}"
        linkedin_result = await _post_job_to_linkedin(created_job, apply_url)

    asyncio.create_task(_audit("create", user=user, entity_type="job", entity_id=created_job.get("id"),
                                entity_name=created_job.get("title"),
                                new_value={"title": created_job.get("title"), "department": created_job.get("department")},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {**created_job, "linkedin_post": linkedin_result}


@api_router.get("/jobs")
async def get_jobs(request: Request, is_active: Optional[bool] = None, search: Optional[str] = None):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    q = sb("jobs").select("*").order("created_at", desc=True)
    if is_active is not None: q = q.eq("is_active", is_active)
    if search: q = q.or_(f"title.ilike.%{search}%,department.ilike.%{search}%")
    res = await run(lambda: q.execute())
    jobs = res.data or []

    # Attach candidate count
    for job in jobs:
        cnt = await run(lambda jid=job["id"]: sb("candidates").select("id", count="exact").eq("job_id", jid).execute())
        job["candidate_count"] = cnt.count or 0
    return jobs


@api_router.get("/jobs/{job_id}")
async def get_job(job_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    job = await safe_single(lambda: sb("jobs").select("*").eq("id", job_id).single().execute())
    if not job:
        raise HTTPException(404, "Job not found")
    cands = await run(lambda: sb("candidates").select("*").eq("job_id", job_id).execute())
    job["candidates"] = cands.data or []
    return job


@api_router.put("/jobs/{job_id}")
async def update_job(job_id: str, job: JobUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    patch = {k: v for k, v in job.model_dump().items() if v is not None}
    old_job = await run(lambda: sb("jobs").select("title").eq("id", job_id).execute())
    job_name = (old_job.data or [{}])[0].get("title") or job_id
    await run(lambda: sb("jobs").update(patch).eq("id", job_id).execute())
    asyncio.create_task(_audit("update", user=user, entity_type="job", entity_id=job_id,
                                entity_name=job_name, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Job updated"}


@api_router.delete("/jobs/{job_id}")
async def delete_job(job_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    old_job = await run(lambda: sb("jobs").select("title").eq("id", job_id).execute())
    job_name = (old_job.data or [{}])[0].get("title") or job_id
    await run(lambda: sb("jobs").delete().eq("id", job_id).execute())
    asyncio.create_task(_audit("delete", user=user, entity_type="job", entity_id=job_id,
                                entity_name=job_name,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Job deleted"}


# ============================================================
# RESUME TEXT EXTRACTION HELPER
# ============================================================
def _extract_resume_text(file_bytes: bytes, content_type: str) -> str:
    """Extract plain text from PDF or DOCX bytes for LLM analysis."""
    try:
        if content_type == "application/pdf":
            import pdfplumber, io as _io
            with pdfplumber.open(_io.BytesIO(file_bytes)) as pdf:
                return "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                ).strip()
        elif content_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            import docx, io as _io
            doc = docx.Document(_io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs).strip()
        elif content_type == "application/msword":
            # Legacy .doc — best-effort UTF-8 decode
            return file_bytes.decode("utf-8", errors="ignore")
    except Exception as exc:
        logger.warning(f"[resume-extract] text extraction failed: {exc}")
    return ""


# ============================================================
# CANDIDATES
# ============================================================
@api_router.post("/candidates")
async def create_candidate(candidate: CandidateCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    doc_c = {
        "full_name":            candidate.full_name,
        "email":                candidate.email,
        "phone":                candidate.phone,
        "current_company":      candidate.current_company,
        "candidate_role":       candidate.candidate_role,
        "experience_years":     candidate.experience_years,
        "source":               candidate.source,
        "job_id":               candidate.job_id,
        "status":               candidate.status or "sourced",
        "notes":                candidate.notes,
        "resume_url":           candidate.resume_url,
        "linkedin_url":         candidate.linkedin_url,
        "portfolio_url":        candidate.portfolio_url,
        "skills":               candidate.skills,
        "assigned_recruiter_id":user["id"],
        "created_by":           user["id"],
        # Extended v3 fields
        "candidate_type":       candidate.candidate_type or "domestic",
        "visa_status":          candidate.visa_status,
        "total_experience":     candidate.total_experience,
        "relevant_experience":  candidate.relevant_experience,
        "location":             candidate.location,
        "relocation":           candidate.relocation,
        # v4 fields
        "work_mode":            candidate.work_mode or [],
        "current_ctc":          candidate.current_ctc,
        "expected_ctc":         candidate.expected_ctc,
        # LLM-extracted fields
        "tech_stack":           candidate.tech_stack or [],
    }
    doc_c = {k: v for k, v in doc_c.items() if v is not None and v != [] }
    if candidate.skills: doc_c["skills"] = candidate.skills
    if candidate.tech_stack is not None: doc_c["tech_stack"] = candidate.tech_stack
    try:
        res = await run(lambda: sb("candidates").insert(doc_c).execute())
    except Exception as e:
        err_str = str(e)
        if "PGRST204" in err_str:
            import re as _re
            col_match = _re.search(r"Could not find the '(\w+)' column", err_str)
            if col_match:
                bad_col = col_match.group(1)
                logger.warning(f"[create_candidate] Column '{bad_col}' missing — skipping. Run add_features_v3.sql.")
                doc_c.pop(bad_col, None)
                res = await run(lambda: sb("candidates").insert(doc_c).execute())
            else:
                raise
        else:
            raise
    candidate_id = res.data[0]["id"]
    await _log_activity(candidate_id=candidate_id, user=user, atype="note",
                        desc=f"Candidate added by {user['name']}")
    asyncio.create_task(_audit("create", user=user, entity_type="candidate", entity_id=candidate_id,
                                entity_name=res.data[0].get("full_name"),
                                new_value={"role": res.data[0].get("candidate_role"), "source": res.data[0].get("source"), "status": res.data[0].get("status")},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return res.data[0]


@api_router.get("/candidates")
async def get_candidates(
    request:      Request,
    job_id:       Optional[str] = None,
    status:       Optional[str] = None,
    source:       Optional[str] = None,
    search:       Optional[str] = None,
    skip:         int = 0,
    limit:        int = 50,
):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    q = sb("candidates").select(
        "*",
        count="exact"
    ).order("created_at", desc=True).range(skip, skip + limit - 1)

    if job_id:  q = q.eq("job_id", job_id)
    if status:  q = q.eq("status", status)
    if source:  q = q.eq("source", source)
    if search:
        q = q.or_(
            f"full_name.ilike.%{search}%,"
            f"email.ilike.%{search}%,"
            f"current_company.ilike.%{search}%,"
            f"candidate_role.ilike.%{search}%"
        )

    res = await run(lambda: q.execute())
    return {"candidates": res.data or [], "total": res.count or 0}


@api_router.get("/candidates/pipeline")
async def get_pipeline(request: Request, job_id: Optional[str] = None):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    pipeline = {}
    for s in CandidateStatus:
        q = sb("candidates").select("*").eq("status", s.value)
        if job_id: q = q.eq("job_id", job_id)
        res = await run(lambda qq=q: qq.execute())
        pipeline[s.value] = res.data or []
    return pipeline


@api_router.get("/candidates/{candidate_id}")
async def get_candidate(candidate_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    candidate_data = await safe_single(lambda: sb("candidates").select("*").eq("id", candidate_id).single().execute())
    if not candidate_data:
        raise HTTPException(404, "Candidate not found")
    cand = candidate_data

    acts = await run(lambda: sb("activities").select("*").eq("candidate_id", candidate_id).order("created_at", desc=True).execute())
    cand["activities"] = acts.data or []

    ivs = await run(lambda: sb("interviews").select("*").eq("candidate_id", candidate_id).order("scheduled_at", desc=True).execute())
    cand["interviews"] = ivs.data or []

    hist = await run(lambda: sb("candidate_status_history").select("*").eq("candidate_id", candidate_id).order("created_at", desc=True).execute())
    cand["status_history"] = hist.data or []
    return cand


@api_router.put("/candidates/{candidate_id}")
async def update_candidate(candidate_id: str, candidate: CandidateUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    existing = await safe_single(lambda: sb("candidates").select("status").eq("id", candidate_id).single().execute())
    if not existing:
        raise HTTPException(404, "Candidate not found")

    patch = {k: v for k, v in candidate.model_dump().items() if v is not None}
    # Explicitly allow empty lists for array fields (model_dump filters None but keeps [])
    for arr_field in ("tech_stack", "work_mode", "skills"):
        val = candidate.model_dump().get(arr_field)
        if val is not None:   # includes []
            patch[arr_field] = val
    if "status" in patch and isinstance(patch["status"], CandidateStatus):
        old_status = existing["status"]
        new_status = patch["status"].value
        patch["status"] = new_status
        if old_status != new_status:
            await run(lambda: sb("candidate_status_history").insert({
                "candidate_id":   candidate_id,
                "old_status":     old_status,
                "new_status":     new_status,
                "changed_by":     user["id"],
                "changed_by_name":user["name"],
            }).execute())
            await _log_activity(candidate_id=candidate_id, user=user, atype="status_change",
                                 desc=f"Stage moved from {old_status} to {new_status}")

    old_cand = await run(lambda: sb("candidates").select("full_name").eq("id", candidate_id).execute())
    cand_name = (old_cand.data or [{}])[0].get("full_name") or candidate_id
    await run(lambda: sb("candidates").update(patch).eq("id", candidate_id).execute())
    asyncio.create_task(_audit("update", user=user, entity_type="candidate", entity_id=candidate_id,
                                entity_name=cand_name, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Candidate updated"}


# ── Resume: Masked download ────────────────────────────────────────────────────
_PHONE_RE = _re.compile(
    r'(\+?(?:91[\s\-]?)?(?:[6-9]\d{9})'                   # Indian mobile
    r'|\+?(?:1[\s\-]?)?(?:\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4})'  # US
    r'|\+\d{1,3}[\s\-]?\d{7,12})',                         # generic international
    _re.VERBOSE,
)
_EMAIL_RE = _re.compile(r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}')

def _mask_text(text: str) -> str:
    text = _EMAIL_RE.sub('[email hidden]', text)
    text = _PHONE_RE.sub('[phone hidden]', text)
    return text

def _mask_pdf_bytes(pdf_bytes: bytes) -> bytes:
    """Redact emails and phone numbers from a PDF using PyMuPDF."""
    if not _FITZ_OK:
        raise RuntimeError("PyMuPDF not installed — cannot mask PDF")
    doc = _fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in doc:
        for pattern in (_EMAIL_RE, _PHONE_RE):
            for match in pattern.finditer(page.get_text()):
                rects = page.search_for(match.group())
                for r in rects:
                    page.add_redact_annot(r, fill=(0, 0, 0))
        page.apply_redactions()
    return doc.tobytes()

def _mask_docx_bytes(docx_bytes: bytes) -> bytes:
    """Replace emails and phone numbers in a DOCX."""
    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    def _replace_in_para(para):
        for run in para.runs:
            run.text = _mask_text(run.text)
    for para in doc.paragraphs:
        _replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@api_router.get("/candidates/{candidate_id}/resume/masked")
async def download_masked_resume(candidate_id: str, request: Request):
    """Download the candidate's resume with phone and email redacted (no AI)."""
    from fastapi.responses import Response as FResponse
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    row = await safe_single(lambda: sb("candidates")
        .select("id,full_name,resume_url").eq("id", candidate_id).single().execute())
    if not row:
        raise HTTPException(404, "Candidate not found")
    url = (row.get("resume_url") or "").strip()
    if not url:
        raise HTTPException(400, "This candidate has no resume on file")
    if download_resume is None:
        raise HTTPException(503, "Google Drive not configured")

    metadata = await run(lambda: get_file_metadata(url))
    mime_type = (metadata.get("mimeType") or "").strip().lower()
    file_name = (metadata.get("name") or "").strip().lower()
    raw = await run(lambda: download_resume(url))

    # Detect type from Drive metadata, not from the preview URL.
    if mime_type == "application/pdf" or file_name.endswith(".pdf"):
        try:
            masked = _mask_pdf_bytes(raw)
        except Exception:
            masked = raw   # return original if masking fails
        ctype = "application/pdf"
        ext   = "pdf"
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.endswith(".docx"):
        try:
            masked = _mask_docx_bytes(raw)
        except Exception:
            masked = raw
        ctype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext   = "docx"
    elif mime_type == "application/msword" or file_name.endswith(".doc"):
        # Legacy .doc files cannot be safely redacted with python-docx.
        # Return the original binary with the correct MIME type so Word can open it.
        masked = raw
        ctype = "application/msword"
        ext   = "doc"
    else:
        # Unknown Drive MIME type. Preserve the original bytes rather than
        # rewriting them as DOCX, which would corrupt non-DOCX resumes.
        masked = raw
        ctype = mime_type or "application/octet-stream"
        ext   = "bin"

    safe_name = (row.get("full_name") or "candidate").replace(" ", "_")
    asyncio.create_task(_audit(
        "resume_download", user=user,
        entity_type="candidate", entity_id=candidate_id,
        entity_name=row.get("full_name"),
        new_value={"type": "masked_resume", "format": ext},
        ip=_get_ip(request), ua=request.headers.get("user-agent", ""),
    ))
    return FResponse(
        content=masked,
        media_type=ctype,
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_masked.{ext}"'},
    )


# ── Single email compose (Sales & Recruit) ────────────────────────────────────
class ComposeEmailRequest(BaseModel):
    to_email:   str
    to_name:    Optional[str] = None
    subject:    str
    html_body:  str
    lead_id:    Optional[str] = None
    candidate_id: Optional[str] = None

@api_router.post("/email/compose")
async def compose_and_send_email(body: ComposeEmailRequest, request: Request):
    """Send a single email from the logged-in user's Outlook mailbox."""
    user = await get_current_user(request)
    from_email = user["email"]
    try:
        await _graph_send_mail(from_email, [body.to_email], body.subject, body.html_body)
    except Exception as exc:
        raise HTTPException(503, f"Email send failed: {exc}")

    # Log activity on lead/candidate
    if body.lead_id:
        await _log_activity(lead_id=body.lead_id, user=user, atype="email",
                            desc=f"Email sent to {body.to_email} — \"{body.subject}\"")
    if body.candidate_id:
        await _log_activity(candidate_id=body.candidate_id, user=user, atype="email",
                            desc=f"Email sent to {body.to_email} — \"{body.subject}\"")
    asyncio.create_task(_audit("email_sent", user=user, entity_type="email",
                                entity_name=body.to_email,
                                new_value={"subject": body.subject, "to": body.to_email, "lead_id": body.lead_id, "candidate_id": body.candidate_id},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"success": True, "sent_from": from_email, "sent_to": body.to_email}


# ── Bulk Welcome Email ────────────────────────────────────────────────────────
class BulkEmailSendRequest(BaseModel):
    subject:        str
    html_body:      str
    extra_emails:   List[str] = []
    excluded_emails: List[str] = []   # emails to skip this send
    target_emails:  Optional[List[str]] = None  # if set, send ONLY to these (subset send)

class BulkEmailTestSendRequest(BaseModel):
    subject:     str
    html_body:   str
    test_emails: List[EmailStr] = []

@api_router.get("/bulk-email/recipients")
async def get_bulk_email_recipients(request: Request):
    """Return all lead contact emails annotated with whether welcome has been sent."""
    user = await get_current_user(request)
    _require_module(user, "sales")

    leads_res = await run(lambda: sb("leads").select(
        "id,company,full_name,email,hq_location,"
        "contact_person_1_name,contact_person_1_email,"
        "contact_person_2_name,contact_person_2_email,"
        "contact_person_3_name,contact_person_3_email"
    ).execute())
    leads = leads_res.data or []

    # Collect all unique non-empty addresses, including hq_location
    recipients: dict = {}  # email -> {name, company, lead_id, hq_location}
    for l in leads:
        loc = l.get("hq_location") or ""
        def _add(email, name, company, lead_id, location=loc):
            if email and email.strip() and "@" in email:
                e = email.strip().lower()
                if e not in recipients:
                    recipients[e] = {"email": e, "name": name or "", "company": company or "", "lead_id": lead_id, "hq_location": location}
        _add(l.get("contact_person_1_email") or l.get("email"), l.get("contact_person_1_name") or l.get("full_name"), l.get("company"), l["id"])
        _add(l.get("contact_person_2_email"), l.get("contact_person_2_name"), l.get("company"), l["id"])
        _add(l.get("contact_person_3_email"), l.get("contact_person_3_name"), l.get("company"), l["id"])

    # Check which have already received a welcome email
    sent_res = await run(lambda: sb("email_sends").select("sent_to").eq("email_type", "welcome").execute())
    sent_set = set()
    for row in (sent_res.data or []):
        sent_to = row.get("sent_to") if isinstance(row, dict) else None
        if isinstance(sent_to, str) and sent_to.strip():
            sent_set.add(sent_to.strip().lower())

    result = []
    for e, meta in recipients.items():
        result.append({**meta, "already_sent": e in sent_set})
    result.sort(key=lambda x: (x["already_sent"], x["company"]))
    return {"recipients": result, "total": len(result), "pending": sum(1 for r in result if not r["already_sent"])}


@api_router.post("/bulk-email/send")
async def send_bulk_email(body: BulkEmailSendRequest, request: Request):
    """Send welcome emails via Microsoft Graph from logged-in user's mailbox."""
    user = await get_current_user(request)
    _require_module(user, "sales")
    from_email = user["email"]

    # Build recipient list from leads + extra
    recip_resp = await get_bulk_email_recipients(request)
    all_recips  = recip_resp["recipients"]

    # Build email -> metadata map for personalization
    email_meta: dict = {r["email"]: r for r in all_recips}

    # Normalize exclusion sets
    excluded = {e.strip().lower() for e in body.excluded_emails if e.strip()}

    # Determine base candidate list
    if body.target_emails is not None:
        # Explicit subset send — send only to target_emails (minus excluded)
        candidates = [e.strip().lower() for e in body.target_emails if e.strip() and "@" in e]
    else:
        # Default: all pending (not already sent)
        candidates = [r["email"] for r in all_recips if not r["already_sent"]]
        # Merge extra emails
        for e in body.extra_emails:
            e = e.strip().lower()
            if e and "@" in e and e not in candidates:
                exists = await run(lambda ee=e: sb("email_sends").select("sent_to")
                    .eq("sent_to", ee).eq("email_type", "welcome").execute())
                if not (exists.data):
                    candidates.append(e)

    # Apply exclusions
    to_send = [e for e in candidates if e not in excluded]

    if not to_send:
        return {"success": True, "sent_count": 0, "skipped_count": len(excluded), "message": "No recipients after exclusions."}

    sent_ok: list[str] = []
    failed:  list[str] = []
    for addr in to_send:
        try:
            # Personalize [Client Name] with company name if available
            meta    = email_meta.get(addr, {})
            company = meta.get("company") or meta.get("name") or "there"
            subject = body.subject
            html    = body.html_body.replace("[Client Name]", company).replace("[client name]", company)
            await _graph_send_mail(from_email, [addr], subject, html)
            sent_ok.append(addr)
        except Exception as exc:
            logger.warning(f"[bulk-email] failed {addr}: {exc}")
            failed.append(addr)

    # Record successful sends in DB
    if sent_ok:
        rows = [{"sent_to": a, "sent_by_id": user["id"], "sent_by_email": from_email,
                 "sent_by_name": user["name"], "email_type": "welcome", "subject": body.subject}
                for a in sent_ok]
        try:
            await run(lambda: sb("email_sends").upsert(rows, on_conflict="sent_to,email_type").execute())
        except Exception as exc:
            logger.warning(f"[bulk-email] upsert failed: {exc}")

    asyncio.create_task(_audit("bulk_email_sent", user=user, entity_type="bulk_email",
                                entity_name=body.subject,
                                new_value={"sent_count": len(sent_ok), "failed_count": len(failed), "subject": body.subject},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {
        "success":       True,
        "sent_count":    len(sent_ok),
        "failed_count":  len(failed),
        "skipped_count": len(excluded),
        "failed":        failed[:20],
        "sent_from":     from_email,
    }


@api_router.post("/bulk-email/test-send")
async def send_bulk_email_test(body: BulkEmailTestSendRequest, request: Request):
    """Send a draft bulk-email preview to manually entered test addresses."""
    user = await get_current_user(request)
    _require_module(user, "sales")
    from_email = user["email"]

    test_emails = []
    for addr in body.test_emails:
        email = str(addr).strip().lower()
        if email and email not in test_emails:
            test_emails.append(email)

    if not test_emails:
        raise HTTPException(400, "Add at least one test email address.")

    try:
        await _graph_send_mail(from_email, test_emails, body.subject, body.html_body)
    except Exception as exc:
        raise HTTPException(503, f"Test email send failed: {exc}")

    return {
        "success": True,
        "sent_count": len(test_emails),
        "sent_from": from_email,
        "sent_to": test_emails,
    }


@api_router.get("/bulk-email/sent")
async def get_sent_bulk_emails(request: Request):
    """Return history of welcome emails sent."""
    user = await get_current_user(request)
    _require_module(user, "sales")
    res = await run(lambda: sb("email_sends")
        .select("*").eq("email_type", "welcome").order("sent_at", desc=True).limit(500).execute())
    return {"sent": res.data or []}


@api_router.delete("/candidates/{candidate_id}")
async def delete_candidate(candidate_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")

    # Fetch resume_url before deleting so we can clean up Drive
    row = await safe_single(
        lambda: sb("candidates").select("id,full_name,resume_url").eq("id", candidate_id).single().execute()
    )
    if not row:
        raise HTTPException(404, "Candidate not found.")

    # Delete from DB first
    await run(lambda: sb("candidates").delete().eq("id", candidate_id).execute())

    # Then remove resume from Google Drive (non-fatal if it fails)
    resume_url = (row or {}).get("resume_url") or ""
    if resume_url and "drive.google.com" in resume_url and delete_resume is not None:
        try:
            await run(lambda: delete_resume(resume_url))
            logger.info(f"[delete_candidate] Drive resume deleted for {candidate_id}")
        except Exception as exc:
            logger.warning(f"[delete_candidate] Drive delete failed for {candidate_id}: {exc}")

    asyncio.create_task(_audit(
        action="delete", user=user,
        entity_type="candidate", entity_id=candidate_id,
        entity_name=row.get("full_name") or candidate_id,
    ))

    return {"success": True, "message": "Candidate and resume deleted."}


@api_router.post("/candidates/{candidate_id}/resume")
async def upload_candidate_resume(
    candidate_id: str,
    request: Request,
    file: UploadFile = File(...),
):
    """
    Upload a resume (PDF / DOC / DOCX) to Google Drive and store the
    preview URL in the candidate's resume_url column in Supabase.
 
    The stored URL is a Google Drive /preview link so the frontend
    can embed it directly in an <iframe> without opening a new tab.
    """
    user = await get_current_user(request)
    _require_module(user, "recruitment")

    # ── Validate MIME type ───────────────────────────────────
    if file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            400,
            f"Unsupported file type '{file.content_type}'. "
            "Only PDF, DOC, and DOCX are accepted."
        )
 
    # ── Read bytes ──────────────────────────────────────────
    contents = await file.read()
    if len(contents) > MAX_FILE_BYTES:
        raise HTTPException(400, "File too large. Maximum allowed size is 10 MB.")
 
    # ── Fetch candidate (need name + existing resume_url) ────
    existing = await safe_single(
        lambda: sb("candidates")
        .select("id,full_name,resume_url")
        .eq("id", candidate_id)
        .single()
        .execute()
    )
    if not existing:
        raise HTTPException(404, "Candidate not found")
 
    # ── Delete old resume from Drive if one exists ───────────
    old_url = existing.get("resume_url") or ""
    if "drive.google.com" in old_url:
        await run(lambda: delete_resume(old_url))
 
    # ── Build a clean filename: Resume_John_Doe_abc12345.pdf ─
    ext       = ALLOWED_MIME_TYPES[file.content_type]
    safe_name = (existing.get("full_name") or "Candidate").replace(" ", "_")
    short_id  = candidate_id.replace("-", "")[:8]
    filename  = f"Resume_{safe_name}_{short_id}.{ext}"
 
    # ── Upload to Google Drive ───────────────────────────────
    if upload_resume is None:
        raise HTTPException(503, "Google Drive integration not configured. "
                                 "Ensure google_drive.py and service account credentials are present.")
    try:
        result = await run(
            lambda: upload_resume(contents, filename, file.content_type)
        )
    except RuntimeError as exc:
        raise HTTPException(503, str(exc))
    except Exception as exc:
        logger.exception("Unexpected Google Drive upload failure")
        raise HTTPException(500, "Unexpected Google Drive upload failure.")
 
    preview_url = result["preview_url"]

    # ── Extract text + call LLM for tech_stack / experience ──
    db_patch: dict = {"resume_url": preview_url}
    tech_stack_extracted: list = []
    experience_extracted = None

    if LLM_AVAILABLE:
        try:
            resume_text = _extract_resume_text(contents, file.content_type)
            if resume_text:
                insights = await extract_resume_insights(resume_text)
                tech_stack_extracted = insights.get("tech_stack") or []
                experience_extracted = insights.get("experience_years")
                if tech_stack_extracted:
                    db_patch["tech_stack"] = tech_stack_extracted
                if experience_extracted is not None:
                    db_patch["experience_years"] = experience_extracted
        except Exception as llm_exc:
            logger.warning(f"[LLM] Resume insights failed (non-fatal): {llm_exc}")

    # ── Persist URL + LLM fields in Supabase ─────────────────
    await run(
        lambda: sb("candidates")
        .update(db_patch)
        .eq("id", candidate_id)
        .execute()
    )

    await _log_activity(
        candidate_id=candidate_id,
        user=user,
        atype="note",
        desc=f"Resume uploaded: {filename}"
              + (f" | {len(tech_stack_extracted)} skills extracted" if tech_stack_extracted else ""),
    )

    asyncio.create_task(_audit("resume_upload", user=user, entity_type="candidate",
                                entity_id=candidate_id, entity_name=filename,
                                new_value={"filename": filename,
                                           "skills_extracted": len(tech_stack_extracted or []),
                                           "experience_years": experience_extracted},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {
        "resume_url":        preview_url,
        "view_url":          result["view_url"],
        "file_id":           result["file_id"],
        "filename":          filename,
        "tech_stack":        tech_stack_extracted,
        "experience_years":  experience_extracted,
    }
 
 
@api_router.delete("/candidates/{candidate_id}/resume")
async def delete_candidate_resume(candidate_id: str, request: Request):
    """
    Remove the candidate's resume from Google Drive and clear resume_url
    in Supabase.
    """
    user = await get_current_user(request)
    _require_module(user, "recruitment")
 
    existing = await safe_single(
        lambda: sb("candidates")
        .select("id,resume_url")
        .eq("id", candidate_id)
        .single()
        .execute()
    )
    if not existing:
        raise HTTPException(404, "Candidate not found")
 
    url = existing.get("resume_url") or ""
    if "drive.google.com" in url:
        await run(lambda: delete_resume(url))
 
    await run(
        lambda: sb("candidates")
        .update({"resume_url": None})
        .eq("id", candidate_id)
        .execute()
    )
 
    await _log_activity(
        candidate_id=candidate_id,
        user=user,
        atype="note",
        desc="Resume removed",
    )
 
    asyncio.create_task(_audit("resume_delete", user=user, entity_type="candidate",
                                entity_id=candidate_id,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Resume deleted successfully"}


# ============================================================
# ATS MATCH — JD → Top-10 candidates
# ============================================================
class ATSMatchRequest(BaseModel):
    jd_text:        str
    candidate_type: str = "domestic"   # domestic | international


@api_router.post("/candidates/ats-match")
async def ats_match(body: ATSMatchRequest, request: Request):
    """
    Given a job description, find the top-10 best-fit candidates using:
      1. LLM (Gemini Flash, 1 call total) — parse JD into required/nice skills
      2. Pure-code scoring — tech_stack column overlap, zero per-candidate LLM calls

    Scoring formula (0-100):
      required_match% * 75  +  nice_to_have_match% * 25
    matched_skills / missing_skills computed locally from set intersection.
    """
    user = await get_current_user(request)
    if user.get("role") not in ("admin", "sales", "viewer"):
        raise HTTPException(403, "ATS matching is not available for your role.")

    if not LLM_AVAILABLE:
        raise HTTPException(503, "LLM service not configured. Set LLM_PROVIDER and API key.")

    jd_text = (body.jd_text or "").strip()
    if len(jd_text) < 30:
        raise HTTPException(400, "Job description is too short. Please provide more detail.")

    # ── Step 1: 1 LLM call — parse JD into structured skills ─
    jd_meta  = await extract_jd_keywords(jd_text)
    required = set(s.lower() for s in jd_meta.get("required_skills")     or [])
    nice     = set(s.lower() for s in jd_meta.get("nice_to_have_skills") or [])

    # ── Step 2: fetch all candidates of this type ─────────────
    res = await run(
        lambda: sb("candidates")
        .select("id,full_name,candidate_role,experience_years,tech_stack,"
                "location,visa_status,candidate_type,total_experience,status")
        .eq("candidate_type", body.candidate_type)
        .execute()
    )
    all_candidates = res.data or []

    if not all_candidates:
        return {"matches": [], "jd_meta": jd_meta, "total_scanned": 0, "pre_filtered": 0}

    # ── Step 3: pure-code scoring (no LLM per candidate) ──────
    def _score_candidate(cand: dict) -> dict:
        cand_set = set(s.lower() for s in (cand.get("tech_stack") or []))

        # Set intersections for matched / missing
        req_hits  = cand_set & required
        nice_hits = cand_set & nice

        # ATS score: required skills worth 75 pts, nice-to-have worth 25 pts
        req_pct  = len(req_hits)  / max(len(required), 1)
        nice_pct = len(nice_hits) / max(len(nice), 1) if nice else 0.0
        if not required and not nice:
            ats_score = 0
        elif not required:
            ats_score = round(nice_pct * 60)   # no required skills defined → cap at 60
        else:
            ats_score = round(req_pct * 75 + nice_pct * 25)

        # Keep original-case skill names for display
        cand_stack_orig = {s.lower(): s for s in (cand.get("tech_stack") or [])}
        jd_required_orig = {s.lower(): s for s in jd_meta.get("required_skills") or []}
        jd_nice_orig     = {s.lower(): s for s in jd_meta.get("nice_to_have_skills") or []}

        matched  = sorted(jd_required_orig.get(k, k) for k in req_hits)
        matched += sorted(jd_nice_orig.get(k, k) for k in nice_hits)
        missing  = sorted(jd_required_orig.get(k, k) for k in (required - cand_set))

        # Templated fit summary (no LLM needed)
        nr, nt = len(required), len(req_hits)
        if ats_score >= 75:
            summary = f"Strong match: {nt}/{nr} required skills covered."
        elif ats_score >= 50:
            summary = f"Good match: {nt}/{nr} required skills found. Gaps are bridgeable."
        elif ats_score >= 25:
            summary = f"Partial match: only {nt}/{nr} required skills present."
        else:
            summary = f"Limited overlap: {nt}/{nr} required skills found in tech stack."

        return {
            "id":               cand["id"],
            "full_name":        cand.get("full_name"),
            "candidate_role":   cand.get("candidate_role"),
            "experience_years": cand.get("experience_years"),
            "total_experience": cand.get("total_experience"),
            "tech_stack":       cand.get("tech_stack") or [],
            "location":         cand.get("location"),
            "visa_status":      cand.get("visa_status"),
            "status":           cand.get("status"),
            "ats_score":        ats_score,
            "matched_skills":   matched,
            "missing_skills":   missing,
            "fit_summary":      summary,
        }

    # Score every candidate, keep those with at least 1 skill match for pre_filtered count
    all_scored   = [_score_candidate(c) for c in all_candidates]
    pre_filtered = [s for s in all_scored if s["ats_score"] > 0]

    # Sort by score desc, return top 10
    top10 = sorted(all_scored, key=lambda x: x["ats_score"], reverse=True)[:10]

    return {
        "matches":       top10,
        "jd_meta":       jd_meta,
        "total_scanned": len(all_candidates),
        "pre_filtered":  len(pre_filtered),
    }


# ============================================================
# LINKEDIN POSTING
# ============================================================

async def _post_job_to_linkedin(job: dict, apply_url: str) -> dict:
    """
    Post a job opening to JM Data Talent LinkedIn company page.
    Returns {"success": True/False, "post_id": "...", "error": "..."}
    
    Required HuggingFace env vars:
      LINKEDIN_ACCESS_TOKEN    — OAuth 2.0 access token (pages:read + w_organization_social)
      LINKEDIN_ORGANIZATION_ID — numeric ID from linkedin.com/company/<name>/admin/ URL
    """
    token  = os.environ.get("LINKEDIN_ACCESS_TOKEN", "").strip()
    # org_id = os.environ.get("LINKEDIN_ORGANIZATION_ID", "").strip() # having issues 
    member_id = os.environ.get("LINKEDIN_MEMBER_ID", "").strip()

    # if not token or not org_id:
    if not token or not member_id:
        return {"success": False, "error": "LinkedIn credentials not configured — set LINKEDIN_ACCESS_TOKEN and LINKEDIN_ORGANIZATION_ID in HuggingFace Spaces secrets."}

    # Build post text
    skills_lines = ""
    if job.get("skills"):
        skills_lines = "\n" + "\n".join(f"• {s}" for s in (job["skills"] or [])[:6])

    desc = (job.get("description") or "").strip()
    short_desc = desc[:300] + ("…" if len(desc) > 300 else "")

    location_line = ""
    if job.get("location"):
        location_line = f"📍 {job['location']}"
    if job.get("employment_type"):
        location_line += f" | {job['employment_type']}" if location_line else f"💼 {job['employment_type']}"

    post_text = f"""🚀 We're Hiring: {job['title']}

{location_line}

{short_desc}{skills_lines}

👉 Apply here → {apply_url}

#Hiring #ITJobs #Dublin #Ireland #JMDataTalent #TechJobs #Recruitment"""

    payload = {
        # "author": f"urn:li:organization:{org_id}",
        "author": f"urn:li:person:{member_id}",
        "lifecycleState": "PUBLISHED",
        "specificContent": {
            "com.linkedin.ugc.ShareContent": {
                "shareCommentary": {"text": post_text},
                "shareMediaCategory": "ARTICLE",
                "media": [{
                    "status": "READY",
                    "description": {"text": short_desc or f"{job['title']} — Apply at JM Data Talent"},
                    "originalUrl": apply_url,
                    "title": {"text": f"Apply: {job['title']} at JM Data Talent"},
                }],
            }
        },
        "visibility": {"com.linkedin.ugc.MemberNetworkVisibility": "PUBLIC"},
    }

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type":  "application/json",
        "X-Restli-Protocol-Version": "2.0.0",
        "LinkedIn-Version": "202401",
    }

    try:
        import httpx
        async with httpx.AsyncClient(timeout=15) as client:
            resp = await client.post("https://api.linkedin.com/v2/ugcPosts", json=payload, headers=headers)
        if resp.status_code in (200, 201):
            return {"success": True, "post_id": resp.headers.get("x-linkedin-id", "posted")}
        return {"success": False, "error": f"LinkedIn API returned {resp.status_code}: {resp.text[:200]}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ============================================================
# BULK RESUME ZIP UPLOAD
# ============================================================

import uuid as _uuid
import zipfile as _zipfile
import io as _io

# In-memory job tracking (cleared on restart — that is fine for short jobs)
_bulk_jobs: dict = {}


async def _process_bulk_zip(job_id: str, zip_bytes: bytes, user: dict):
    """Background task: extract ZIP → parse each resume → upload Drive → insert candidate."""
    state = _bulk_jobs.get(job_id)
    if not state:
        logging.getLogger(__name__).error(f"[bulk] job_id {job_id} not found in _bulk_jobs")
        return
    try:
        from llm_utils import extract_resume_full_profile
    except ImportError as ie:
        state["status"] = "error"
        state["error"]  = f"LLM import failed: {ie}"
        logging.getLogger(__name__).error(f"[bulk] llm_utils import failed: {ie}")
        return

    ALLOWED = {
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    EXT_MIME = {".pdf": "application/pdf", ".doc": "application/msword", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}

    try:
        zf = _zipfile.ZipFile(_io.BytesIO(zip_bytes))
    except Exception as e:
        state["status"] = "error"
        state["error"]  = f"Could not open ZIP: {e}"
        return

    # Filter to resume files only (skip __MACOSX junk and directories)
    entries = [
        n for n in zf.namelist()
        if not n.startswith("__MACOSX") and not n.endswith("/")
        and any(n.lower().endswith(ext) for ext in EXT_MIME)
    ]

    state["total"] = len(entries)
    if not entries:
        state["status"] = "done"
        return

    for i, name in enumerate(entries):
        fname = name.split("/")[-1]  # strip any folder prefix inside zip
        ext   = "." + fname.rsplit(".", 1)[-1].lower()
        mime  = EXT_MIME.get(ext, "application/pdf")

        result_entry = {"filename": fname, "status": "processing", "candidate_id": None, "error": None}
        state["results"].append(result_entry)

        try:
            file_bytes = zf.read(name)

            # 1. Extract text
            resume_text = _extract_resume_text(file_bytes, mime)
            if not resume_text.strip():
                result_entry["status"] = "skipped"
                result_entry["error"]  = "Could not extract text"
                state["processed"] += 1
                continue

            # 2. LLM — full profile (1 call per resume)
            profile = await extract_resume_full_profile(resume_text)

            # 3. Upload to Google Drive — upload_resume returns dict {file_id, preview_url, view_url}
            drive_url = None
            if upload_resume is not None:
                try:
                    _loop = asyncio.get_running_loop()
                    drive_result = await _loop.run_in_executor(
                        None, lambda b=file_bytes, fn=fname, m=mime: upload_resume(b, fn, m)
                    )
                    if isinstance(drive_result, dict):
                        drive_url = drive_result.get("preview_url") or drive_result.get("view_url")
                    elif isinstance(drive_result, str):
                        drive_url = drive_result
                except Exception as de:
                    logging.getLogger(__name__).warning(f"[bulk] Drive upload failed for {fname}: {de}")

            # 4. Insert candidate (upsert by email if available)
            row = {
                "full_name":        profile.get("full_name") or fname.rsplit(".", 1)[0].replace("_"," ").replace("-"," ").title(),
                "email":            profile.get("email"),
                "phone":            profile.get("phone"),
                "current_company":  profile.get("current_company"),
                "candidate_role":   profile.get("candidate_role"),
                "experience_years": profile.get("experience_years"),
                "location":         profile.get("location"),
                "tech_stack":       profile.get("tech_stack") or [],
                "skills":           profile.get("tech_stack") or [],
                "resume_url":       drive_url,
                "source":           "bulk_upload",
                "status":           "sourced",
                "created_by":       user.get("id"),
            }

            sb_client = get_supabase()
            # Upsert by email if available, else plain insert
            if row["email"]:
                existing = sb_client.table("candidates").select("id").eq("email", row["email"]).execute()
                if existing.data:
                    cid = existing.data[0]["id"]
                    sb_client.table("candidates").update({k: v for k, v in row.items() if v is not None}).eq("id", cid).execute()
                    result_entry["candidate_id"] = cid
                    result_entry["action"] = "updated"
                else:
                    ins = sb_client.table("candidates").insert(row).execute()
                    result_entry["candidate_id"] = ins.data[0]["id"] if ins.data else None
                    result_entry["action"] = "created"
            else:
                ins = sb_client.table("candidates").insert(row).execute()
                result_entry["candidate_id"] = ins.data[0]["id"] if ins.data else None
                result_entry["action"] = "created"

            result_entry["status"]    = "done"
            result_entry["full_name"] = row["full_name"]
            result_entry["email"]     = row["email"]
            result_entry["candidate_role"] = row["candidate_role"]
            state["created"] += 1

        except Exception as exc:
            result_entry["status"] = "error"
            result_entry["error"]  = str(exc)[:200]
            state["errors"] += 1

        state["processed"] += 1

        # Sleep between calls to respect Groq free-tier rate limits
        if i < len(entries) - 1:
            await asyncio.sleep(4)

    state["status"] = "done"
    zf.close()


@api_router.post("/candidates/bulk-upload-zip")
async def bulk_upload_zip(
    request: Request,
    background_tasks: BackgroundTasks,
    zip_file: UploadFile = File(...),
):
    """Accept a ZIP of resumes, process in background, return a job_id to poll."""
    user = await get_current_user(request)
    _require_module(user, "recruitment")

    if not zip_file.filename.lower().endswith(".zip"):
        raise HTTPException(400, "Please upload a .zip file.")

    raw = await zip_file.read()
    if len(raw) > 100 * 1024 * 1024:  # 100 MB limit
        raise HTTPException(400, "ZIP file too large (max 100 MB).")

    job_id = str(_uuid.uuid4())
    _bulk_jobs[job_id] = {
        "status": "running", "total": 0, "processed": 0,
        "created": 0, "errors": 0, "results": [],
    }

    background_tasks.add_task(_process_bulk_zip, job_id, raw, user)
    asyncio.create_task(_audit("import", user=user, entity_type="candidate",
                                entity_name=zip_file.filename,
                                new_value={"filename": zip_file.filename, "job_id": job_id, "type": "bulk_zip"},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"job_id": job_id}


@api_router.get("/candidates/bulk-upload-status/{job_id}")
async def bulk_upload_status(job_id: str, request: Request):
    """Poll for bulk upload progress."""
    await get_current_user(request)
    if job_id not in _bulk_jobs:
        raise HTTPException(404, "Job not found. Jobs are cleared on server restart.")
    return _bulk_jobs[job_id]


# ─────────────────────────────────────────────────────────────
# ATS RESUME UPLOAD SCORING  (free-tier safe: 1 LLM call/resume)
# ─────────────────────────────────────────────────────────────

class BackfillCandidateRoleRequest(BaseModel):
    limit: Optional[int] = 100


@api_router.post("/candidates/backfill-roles")
async def backfill_candidate_roles(request: Request, body: BackfillCandidateRoleRequest = None):
    """
    Backfill missing candidate_role values from already stored Drive resumes.
    Only updates rows where candidate_role is empty and resume_url is present.
    """
    user = await get_current_user(request)
    if user.get("role") != "admin":
        raise HTTPException(403, "Only admins can run role backfills.")

    if download_resume is None:
        raise HTTPException(503, "Google Drive download is not configured.")

    limit = int(getattr(body, "limit", 100) or 100)
    limit = max(1, min(limit, 500))

    from llm_utils import extract_resume_full_profile

    rows_res = await run(
        lambda: sb("candidates")
        .select("id,full_name,candidate_role,resume_url")
        .execute()
    )
    rows = getattr(rows_res, "data", []) or []

    updated = 0
    scanned = 0
    skipped = 0
    errors = 0

    for cand in rows:
        if scanned >= limit:
            break
        if (cand.get("candidate_role") or "").strip():
            continue
        resume_url = (cand.get("resume_url") or "").strip()
        if not resume_url:
            continue

        scanned += 1
        try:
            file_bytes = await run(lambda url=resume_url: download_resume(url))

            resume_text = _extract_resume_text(file_bytes, "application/pdf")
            if not resume_text.strip():
                resume_text = _extract_resume_text(file_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if not resume_text.strip():
                resume_text = _extract_resume_text(file_bytes, "application/msword")

            if not resume_text.strip():
                skipped += 1
                continue

            profile = await extract_resume_full_profile(resume_text)
            candidate_role = (profile.get("candidate_role") or "").strip()
            if not candidate_role:
                skipped += 1
                continue

            await run(
                lambda cid=cand["id"], role=candidate_role: sb("candidates")
                .update({"candidate_role": role})
                .eq("id", cid)
                .execute()
            )
            updated += 1
        except Exception as exc:
            errors += 1
            logger.warning(f"[backfill-roles] failed for {cand.get('id')}: {exc}")

    return {
        "scanned": scanned,
        "updated": updated,
        "skipped": skipped,
        "errors": errors,
        "limit": limit,
    }


class ParseJDRequest(BaseModel):
    jd_text: str

@api_router.post("/candidates/parse-jd-for-scoring")
async def parse_jd_for_scoring(body: ParseJDRequest, request: Request):
    """Parse JD into structured skills (1 LLM call). Call once before scoring resumes."""
    await get_current_user(request)
    if not LLM_AVAILABLE:
        raise HTTPException(503, "LLM service not configured. Set LLM_PROVIDER and API key.")
    jd_text = (body.jd_text or "").strip()
    if len(jd_text) < 30:
        raise HTTPException(400, "Job description is too short (minimum 30 characters).")
    jd_meta = await extract_jd_keywords(jd_text)
    return jd_meta


@api_router.post("/candidates/score-resume-upload")
async def score_resume_upload(
    request: Request,
    resume: UploadFile = File(...),
    jd_skills: str = Form(...),
):
    """
    Score a single uploaded resume against pre-parsed JD skills.
    1 LLM call per resume. Caller should sleep 2s between consecutive calls (Groq free tier).
    """
    await get_current_user(request)

    allowed_ct = {
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    ct = (resume.content_type or "").split(";")[0].strip()
    if ct not in allowed_ct:
        raise HTTPException(400, f"Unsupported file: {resume.filename}. Upload PDF or DOCX only.")

    contents = await resume.read()
    if len(contents) > 10 * 1024 * 1024:
        raise HTTPException(400, "Resume too large (max 10 MB).")

    resume_text = _extract_resume_text(contents, ct)
    if not resume_text.strip():
        return {
            "filename": resume.filename,
            "ats_score": 0,
            "matched_skills": [],
            "missing_skills": [],
            "tech_stack": [],
            "experience_years": None,
            "fit_summary": "Could not extract text from this file.",
            "error": "text_extraction_failed",
        }

    insights = await extract_resume_insights(resume_text)
    tech_stack = insights.get("tech_stack") or []
    experience_years = insights.get("experience_years")

    try:
        jd_meta = json.loads(jd_skills)
    except Exception:
        raise HTTPException(400, "Invalid jd_skills JSON.")

    required = set(s.lower() for s in jd_meta.get("required_skills") or [])
    nice     = set(s.lower() for s in jd_meta.get("nice_to_have_skills") or [])
    cand_set = set(s.lower() for s in tech_stack)

    req_hits  = cand_set & required
    nice_hits = cand_set & nice
    req_pct   = len(req_hits) / max(len(required), 1)
    nice_pct  = len(nice_hits) / max(len(nice), 1) if nice else 0.0

    if not required and not nice:
        ats_score = 0
    elif not required:
        ats_score = round(nice_pct * 60)
    else:
        ats_score = round(req_pct * 75 + nice_pct * 25)

    jd_req_orig  = {s.lower(): s for s in jd_meta.get("required_skills") or []}
    jd_nice_orig = {s.lower(): s for s in jd_meta.get("nice_to_have_skills") or []}

    matched  = sorted(jd_req_orig.get(k, k)  for k in req_hits)
    matched += sorted(jd_nice_orig.get(k, k) for k in nice_hits)
    missing  = sorted(jd_req_orig.get(k, k)  for k in (required - cand_set))

    nr, nt = len(required), len(req_hits)
    if ats_score >= 75:
        summary = f"Strong match: {nt}/{nr} required skills covered."
    elif ats_score >= 50:
        summary = f"Good match: {nt}/{nr} required skills found. Gaps are bridgeable."
    elif ats_score >= 25:
        summary = f"Partial match: only {nt}/{nr} required skills present."
    else:
        summary = f"Limited overlap: {nt}/{nr} required skills found."

    return {
        "filename": resume.filename,
        "ats_score": ats_score,
        "matched_skills": matched,
        "missing_skills": missing,
        "tech_stack": tech_stack,
        "experience_years": experience_years,
        "fit_summary": summary,
    }


# ============================================================
# SALES TRACKER
# ============================================================

STAGE_PROBABILITY_MAP = {
    "Cold Outreach": 10,
    "Engaged / Replied": 20,
    "Discovery Call Booked": 30,
    "Discovery Done / Qualified": 40,
    "Proposal Sent": 60,
    "Negotiation": 75,
    "Closed-Won": 100,
    "Closed-Lost": 0,
}

SALES_TARGETS = {
    "weeklyEmails":   {"min": 50,    "max": 75},
    "weeklyLinkedin": {"min": 30,    "max": 50},
    "weeklyCalls":    {"min": 15,    "max": 25},
    "weeklyMeetings": {"min": 8,     "max": 12},
    "weeklyProposals":{"min": 4,     "max": 6},
    "dailyEmails":    {"min": 10,    "max": 15},
    "dailyLinkedin":  {"min": 8,     "max": 10},
    "dailyCalls":     {"min": 3,     "max": 5},
    "dailyFollowups": {"min": 3,     "max": 5},
    "monthlyClients": {"min": 2,     "max": 3},
    "monthlyValue":   {"min": 15000, "max": 30000},
    "monthlyProposals":{"min": 16,   "max": 24},
    "pipelineValue":  50000,
}

def _get_week_bounds(iso_week: int, year: int):
    from datetime import date, timedelta
    jan4 = date(year, 1, 4)
    delta = timedelta(weeks=iso_week - 1, days=-jan4.weekday())
    monday = jan4 + delta
    friday  = monday + timedelta(days=4)
    return monday, friday

def _current_iso_week():
    from datetime import date
    today = date.today()
    iso   = today.isocalendar()
    return iso[1], iso[0]   # week, year

def _sum_logs(rows):
    return {
        "emails":         sum(r.get("emails_sent")      or 0 for r in rows),
        "linkedin":       sum(r.get("linkedin_sent")    or 0 for r in rows),
        "calls":          sum(r.get("calls_made")       or 0 for r in rows),
        "replies":        sum(r.get("replies_received") or 0 for r in rows),
        "meetingsBooked": sum(r.get("meetings_booked")  or 0 for r in rows),
        "meetingsDone":   sum(r.get("meetings_done")    or 0 for r in rows),
        "proposals":      sum(r.get("proposals_sent")   or 0 for r in rows),
        "followups":      sum(r.get("followups_done")   or 0 for r in rows),
        "newLeads":       sum(r.get("new_leads_added")  or 0 for r in rows),
        "daysLogged":     len(rows),
    }

def _days_in_stage(stage_updated_date_str):
    from datetime import date, datetime
    if not stage_updated_date_str:
        return 0
    try:
        d = datetime.strptime(stage_updated_date_str, "%Y-%m-%d").date()
    except Exception:
        return 0
    return (date.today() - d).days


# ── Pydantic models ────────────────────────────────────────────

class SalesActivityLogCreate(BaseModel):
    log_date:          str
    emails_sent:       int = 0
    linkedin_sent:     int = 0
    calls_made:        int = 0
    replies_received:  int = 0
    meetings_booked:   int = 0
    meetings_done:     int = 0
    proposals_sent:    int = 0
    followups_done:    int = 0
    new_leads_added:   int = 0
    hours_worked:      float = 0.0
    mood:              Optional[int] = None
    biggest_win:       Optional[str] = None
    biggest_blocker:   Optional[str] = None

class SalesPipelineDealCreate(BaseModel):
    client_name:        str
    industry:           Optional[str] = None
    stage:              str = "Cold Outreach"
    deal_value:         float = 0.0
    next_action:        Optional[str] = None
    next_action_date:   Optional[str] = None
    owner:              Optional[str] = None
    notes:              Optional[str] = None
    stage_updated_date: Optional[str] = None

class SalesPipelineDealUpdate(BaseModel):
    client_name:        Optional[str] = None
    industry:           Optional[str] = None
    stage:              Optional[str] = None
    deal_value:         Optional[float] = None
    next_action:        Optional[str] = None
    next_action_date:   Optional[str] = None
    owner:              Optional[str] = None
    notes:              Optional[str] = None
    stage_updated_date: Optional[str] = None

class SalesWeeklyReviewCreate(BaseModel):
    week_number:     int
    year:            int
    date_range:      Optional[str] = None
    new_leads:       int = 0
    leads_qualified: int = 0
    deals_lost:      int = 0
    loss_reason:     Optional[str] = None
    clients_signed:  int = 0
    contract_value:  float = 0.0
    what_worked:     Optional[str] = None
    what_didnt:      Optional[str] = None
    what_to_change:  Optional[str] = None
    help_needed:     Optional[str] = None
    top_priorities:  Optional[str] = None

class SalesMonthlyRollupCreate(BaseModel):
    month:                str
    year:                 int
    clients_signed:       int = 0
    total_contract_value: float = 0.0
    avg_deal_size:        float = 0.0
    proposals_sent:       int = 0
    proposal_close_rate:  float = 0.0
    pipeline_value:       float = 0.0
    best_industry:        Optional[str] = None
    worst_industry:       Optional[str] = None
    top_objection:        Optional[str] = None
    best_channel:         Optional[str] = None
    top_fix:              Optional[str] = None
    pricing_feedback:     Optional[str] = None
    competitor_names:     Optional[str] = None


# ── Endpoints ──────────────────────────────────────────────────

@api_router.get("/sales/tracker/users")
async def get_sales_tracker_users(request: Request):
    """Return all users who have ever logged sales activity (for admin/viewer user picker)."""
    user = await get_current_user(request)
    _require_module(user, "sales")
    sb = get_supabase()
    # Get distinct logged_by users from the activity log
    res = await run(lambda: sb.table("sales_activity_log")
        .select("logged_by,logged_by_name").execute())
    rows = res.data or []
    seen = {}
    for r in rows:
        uid = r.get("logged_by")
        if uid and uid not in seen:
            seen[uid] = r.get("logged_by_name") or uid
    # Also pull from users table for completeness
    users_res = await run(lambda: sb.table("users")
        .select("id,name,role").in_("role", ["sales", "admin"]).execute())
    for u in (users_res.data or []):
        if u["id"] not in seen:
            seen[u["id"]] = u.get("name") or u["id"]
    return [{"id": uid, "name": name} for uid, name in seen.items()]


@api_router.get("/sales/tracker/dashboard")
async def get_sales_tracker_dashboard(
    request: Request,
    user_id: Optional[str] = None,   # filter to a specific user (admin/viewer can pass any; sales sees only own)
    week_offset: int = 0,            # 0 = current week, -1 = last week, etc.
):
    user = await get_current_user(request)
    _require_module(user, "sales")

    from datetime import date, timedelta
    from datetime import datetime as dt

    role = user.get("role", "")
    is_privileged = role in ("admin", "viewer")  # admin or CEO (viewer) can see everyone

    # Determine whose logs to show
    # - sales role: always their own logs only
    # - admin/viewer: can filter by user_id or see all
    filter_user_id = None
    if not is_privileged:
        filter_user_id = user["id"]   # sales: locked to own
    elif user_id:
        filter_user_id = user_id      # admin/viewer: optional filter

    today = date.today()
    week_num, year = _current_iso_week()

    # Apply week offset
    adjusted_week = week_num + week_offset
    adjusted_year = year
    while adjusted_week <= 0:
        adjusted_week += 52; adjusted_year -= 1
    while adjusted_week > 52:
        adjusted_week -= 52; adjusted_year += 1

    monday, friday = _get_week_bounds(adjusted_week, adjusted_year)
    prev_week = adjusted_week - 1
    prev_year = adjusted_year
    if prev_week <= 0:
        prev_week += 52; prev_year -= 1
    prev_monday, prev_friday = _get_week_bounds(prev_week, prev_year)

    sb = get_supabase()

    def _filter_by_user(q):
        if filter_user_id:
            q = q.eq("logged_by", filter_user_id)
        return q

    # ── This week logs ──
    q = sb.table("sales_activity_log").select("*").gte("log_date", monday.isoformat()).lte("log_date", friday.isoformat())
    this_week_rows = (await run(lambda qq=_filter_by_user(q): qq.execute())).data or []

    # ── Last week logs ──
    q2 = sb.table("sales_activity_log").select("*").gte("log_date", prev_monday.isoformat()).lte("log_date", prev_friday.isoformat())
    last_week_rows = (await run(lambda qq=_filter_by_user(q2): qq.execute())).data or []

    # ── Last 4 weeks for traffic light ──
    four_weeks_ago = monday - timedelta(weeks=3)
    q3 = sb.table("sales_activity_log").select("*").gte("log_date", four_weeks_ago.isoformat()).lte("log_date", friday.isoformat()).order("log_date")
    all_recent = (await run(lambda qq=_filter_by_user(q3): qq.execute())).data or []

    last4 = []
    for w in range(3, -1, -1):
        wk = adjusted_week - w
        yr = adjusted_year
        if wk <= 0:
            wk += 52; yr -= 1
        wmon, wfri = _get_week_bounds(wk, yr)
        wrows = [r for r in all_recent if wmon.isoformat() <= r.get("log_date","") <= wfri.isoformat()]
        s = _sum_logs(wrows)
        s["weekNumber"] = wk
        s["year"]       = yr
        s["dateRange"]  = f"{wmon.strftime('%d %b')} – {wfri.strftime('%d %b')}"
        last4.append(s)

    # ── Daily breakdown for current week ──
    # FIX: was `days_map[day_name] = r` which overwrote when >1 person logged same day.
    # Now we SUM all entries per day so multi-user logs are aggregated correctly.
    days_accumulator: dict = {}  # day_name -> list of rows
    for r in this_week_rows:
        d = r.get("log_date", "")
        if d:
            try:
                day_name = dt.strptime(d, "%Y-%m-%d").strftime("%A")
                days_accumulator.setdefault(day_name, []).append(r)
            except Exception:
                pass

    days_map = {}
    for day_name, rows in days_accumulator.items():
        summed = _sum_logs(rows)
        # Keep the most recent row's narrative fields (mood, biggest_win, etc.) for display
        latest = max(rows, key=lambda r: r.get("updated_at") or "")
        days_map[day_name] = {
            **summed,
            "mood":            latest.get("mood"),
            "biggest_win":     latest.get("biggest_win"),
            "biggest_blocker": latest.get("biggest_blocker"),
            "log_date":        latest.get("log_date"),
            "logged_by_name":  latest.get("logged_by_name"),
            "multiple_users":  len(rows) > 1,
        }

    # ── Per-user breakdown for this week (admin/viewer only, when showing all) ──
    per_user_breakdown = []
    if is_privileged and not filter_user_id:
        user_groups: dict = {}
        for r in this_week_rows:
            uid = r.get("logged_by", "unknown")
            uname = r.get("logged_by_name") or uid
            user_groups.setdefault(uid, {"name": uname, "rows": []})["rows"].append(r)
        for uid, info in user_groups.items():
            s = _sum_logs(info["rows"])
            s["userId"]   = uid
            s["userName"] = info["name"]
            per_user_breakdown.append(s)

    # ── Pipeline ──
    pipeline_res = await run(lambda: sb.table("sales_pipeline").select("*").eq("is_active", True).order("weighted_value", desc=True).execute())
    pipeline = pipeline_res.data or []
    for deal in pipeline:
        deal["daysInStage"] = _days_in_stage(deal.get("stage_updated_date"))

    total_value    = sum(float(d.get("deal_value") or 0) for d in pipeline)
    weighted_total = sum(float(d.get("weighted_value") or 0) for d in pipeline)
    stalled        = [d for d in pipeline if d.get("daysInStage", 0) > 7]
    open_deals     = [d for d in pipeline if d.get("stage") not in ("Closed-Won","Closed-Lost")]
    top_deals      = sorted(open_deals, key=lambda d: float(d.get("weighted_value") or 0), reverse=True)[:5]

    # ── Latest monthly rollup ──
    monthly_res = await run(lambda: sb.table("sales_monthly_rollup").select("*").order("year", desc=True).order("created_at", desc=True).limit(1).execute())
    monthly = (monthly_res.data or [None])[0]

    return {
        "thisWeek":          {**_sum_logs(this_week_rows),
                              "weekNumber": adjusted_week, "year": adjusted_year,
                              "dateRange": f"{monday.strftime('%d %b')} – {friday.strftime('%d %b')}",
                              "days": days_map},
        "lastWeek":          _sum_logs(last_week_rows),
        "last4Weeks":        last4,
        "pipeline":          pipeline,
        "pipelineStats":     {"totalDeals": len(pipeline), "totalValue": total_value,
                              "weightedValue": weighted_total, "stalledCount": len(stalled)},
        "topDeals":          top_deals,
        "stalledDeals":      stalled,
        "monthlyRollup":     monthly,
        "targets":           SALES_TARGETS,
        "viewingUserId":     filter_user_id,
        "isPrivileged":      is_privileged,
        "perUserBreakdown":  per_user_breakdown,
    }


@api_router.post("/sales/tracker/log")
async def submit_sales_log(body: SalesActivityLogCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    from datetime import datetime as dt, date as _date
    try:
        d = dt.strptime(body.log_date, "%d-%b-%Y").date()
    except Exception:
        try:
            d = _date.fromisoformat(body.log_date)
        except Exception:
            raise HTTPException(400, "Invalid date format. Expected YYYY-MM-DD.")
    day_name = d.strftime("%A")
    now_ts   = datetime.now(timezone.utc).isoformat()
    row = {
        "logged_by":        user.get("id"),
        "logged_by_name":   user.get("name", ""),
        "log_date":         d.isoformat(),
        "day_of_week":      day_name,
        "emails_sent":      body.emails_sent,
        "linkedin_sent":    body.linkedin_sent,
        "calls_made":       body.calls_made,
        "replies_received": body.replies_received,
        "meetings_booked":  body.meetings_booked,
        "meetings_done":    body.meetings_done,
        "proposals_sent":   body.proposals_sent,
        "followups_done":   body.followups_done,
        "new_leads_added":  body.new_leads_added,
        "hours_worked":     body.hours_worked,
        "mood":             body.mood,
        "biggest_win":      body.biggest_win,
        "biggest_blocker":  body.biggest_blocker,
        "updated_at":       now_ts,
    }
    res = await run(lambda: sb("sales_activity_log").upsert(row, on_conflict="log_date,logged_by").execute())
    return {"success": True, "data": (res.data or [{}])[0]}


@api_router.get("/sales/tracker/log")
async def get_sales_logs(
    request: Request,
    from_date: Optional[str] = None,
    to_date:   Optional[str] = None,
    user_id:   Optional[str] = None,
    limit:     int = 60,
):
    user = await get_current_user(request)
    _require_module(user, "sales")
    role = user.get("role", "")
    is_privileged = role in ("admin", "viewer")

    q = sb("sales_activity_log").select("*").order("log_date", desc=True).limit(limit)
    if from_date: q = q.gte("log_date", from_date)
    if to_date:   q = q.lte("log_date", to_date)
    # sales role: always locked to own logs
    if not is_privileged:
        q = q.eq("logged_by", user["id"])
    elif user_id:
        q = q.eq("logged_by", user_id)
    res  = await run(lambda qq=q: qq.execute())
    return res.data or []


@api_router.post("/sales/tracker/pipeline")
async def create_pipeline_deal(body: SalesPipelineDealCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    from datetime import date
    prob = STAGE_PROBABILITY_MAP.get(body.stage, 10)
    wv   = round(body.deal_value * prob / 100, 2)
    sud  = body.stage_updated_date or date.today().isoformat()
    next_act_date = None
    if body.next_action_date:
        try:
            from datetime import datetime as dt
            next_act_date = dt.strptime(body.next_action_date, "%d-%b-%Y").date().isoformat()
        except Exception:
            next_act_date = body.next_action_date
    row = {
        "client_name":        body.client_name,
        "industry":           body.industry,
        "stage":              body.stage,
        "deal_value":         body.deal_value,
        "probability":        prob,
        "weighted_value":     wv,
        "next_action":        body.next_action,
        "next_action_date":   next_act_date,
        "owner":              body.owner or user.get("name",""),
        "notes":              body.notes,
        "stage_updated_date": sud,
        "is_active":          True,
    }
    sb = get_supabase()
    res = sb.table("sales_pipeline").insert(row).execute()
    deal = (res.data or [{}])[0]
    deal["daysInStage"] = _days_in_stage(deal.get("stage_updated_date"))
    return {"success": True, "deal": deal}


@api_router.get("/sales/tracker/pipeline")
async def get_pipeline_deals(request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    sb = get_supabase()
    res = sb.table("sales_pipeline").select("*").eq("is_active", True)        .order("weighted_value", desc=True).execute()
    deals = res.data or []
    for d in deals:
        d["daysInStage"] = _days_in_stage(d.get("stage_updated_date"))
    return deals


@api_router.put("/sales/tracker/pipeline/{deal_id}")
async def update_pipeline_deal(deal_id: str, body: SalesPipelineDealUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    from datetime import date
    updates = {}
    if body.client_name is not None:      updates["client_name"]   = body.client_name
    if body.industry    is not None:      updates["industry"]      = body.industry
    if body.deal_value  is not None:      updates["deal_value"]    = body.deal_value
    if body.next_action is not None:      updates["next_action"]   = body.next_action
    if body.owner       is not None:      updates["owner"]         = body.owner
    if body.notes       is not None:      updates["notes"]         = body.notes
    if body.stage is not None:
        prob = STAGE_PROBABILITY_MAP.get(body.stage, 10)
        updates["stage"]       = body.stage
        updates["probability"] = prob
        if body.stage_updated_date:
            updates["stage_updated_date"] = body.stage_updated_date
        else:
            updates["stage_updated_date"] = date.today().isoformat()
    if body.deal_value is not None or "probability" in updates:
        sb = get_supabase()
        existing = sb.table("sales_pipeline").select("deal_value,probability").eq("id", deal_id).execute()
        ex = (existing.data or [{}])[0]
        dv  = updates.get("deal_value",  float(ex.get("deal_value")  or 0))
        pr  = updates.get("probability", float(ex.get("probability") or 10))
        updates["weighted_value"] = round(dv * pr / 100, 2)
    if body.next_action_date is not None:
        try:
            from datetime import datetime as dt
            updates["next_action_date"] = dt.strptime(body.next_action_date, "%d-%b-%Y").date().isoformat()
        except Exception:
            updates["next_action_date"] = body.next_action_date
    updates["updated_at"] = "now()"
    sb = get_supabase()
    res = sb.table("sales_pipeline").update(updates).eq("id", deal_id).execute()
    deal = (res.data or [{}])[0]
    deal["daysInStage"] = _days_in_stage(deal.get("stage_updated_date"))
    asyncio.create_task(_audit("update", user=user, entity_type="deal",
                                entity_id=deal_id, new_value=updates,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"success": True, "deal": deal}


@api_router.delete("/sales/tracker/pipeline/{deal_id}")
async def delete_pipeline_deal(deal_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    sb = get_supabase()
    sb.table("sales_pipeline").update({"is_active": False}).eq("id", deal_id).execute()
    asyncio.create_task(_audit("delete", user=user, entity_type="deal",
                                entity_id=deal_id,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"success": True}


@api_router.post("/sales/tracker/weekly-review")
async def submit_weekly_review(body: SalesWeeklyReviewCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    row = {
        "week_number":    body.week_number,
        "year":           body.year,
        "date_range":     body.date_range,
        "new_leads":      body.new_leads,
        "leads_qualified":body.leads_qualified,
        "deals_lost":     body.deals_lost,
        "loss_reason":    body.loss_reason,
        "clients_signed": body.clients_signed,
        "contract_value": body.contract_value,
        "what_worked":    body.what_worked,
        "what_didnt":     body.what_didnt,
        "what_to_change": body.what_to_change,
        "help_needed":    body.help_needed,
        "top_priorities": body.top_priorities,
    }
    sb = get_supabase()
    res = sb.table("sales_weekly_review").upsert(row, on_conflict="week_number,year").execute()
    return {"success": True, "data": (res.data or [{}])[0]}


@api_router.get("/sales/tracker/weekly-review")
async def get_weekly_reviews(request: Request, limit: int = 8):
    user = await get_current_user(request)
    _require_module(user, "sales")
    sb = get_supabase()
    res = sb.table("sales_weekly_review").select("*")        .order("year", desc=True).order("week_number", desc=True).limit(limit).execute()
    return res.data or []


@api_router.post("/sales/tracker/monthly-rollup")
async def submit_monthly_rollup(body: SalesMonthlyRollupCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    row = {
        "month":                body.month,
        "year":                 body.year,
        "clients_signed":       body.clients_signed,
        "total_contract_value": body.total_contract_value,
        "avg_deal_size":        body.avg_deal_size,
        "proposals_sent":       body.proposals_sent,
        "proposal_close_rate":  body.proposal_close_rate,
        "pipeline_value":       body.pipeline_value,
        "best_industry":        body.best_industry,
        "worst_industry":       body.worst_industry,
        "top_objection":        body.top_objection,
        "best_channel":         body.best_channel,
        "top_fix":              body.top_fix,
        "pricing_feedback":     body.pricing_feedback,
        "competitor_names":     body.competitor_names,
    }
    sb = get_supabase()
    res = sb.table("sales_monthly_rollup").upsert(row, on_conflict="month,year").execute()
    return {"success": True, "data": (res.data or [{}])[0]}


@api_router.get("/sales/tracker/monthly-rollup")
async def get_monthly_rollups(request: Request, limit: int = 6):
    user = await get_current_user(request)
    _require_module(user, "sales")
    sb = get_supabase()
    res = sb.table("sales_monthly_rollup").select("*")        .order("year", desc=True).order("created_at", desc=True).limit(limit).execute()
    return res.data or []



# ============================================================
# INTERVIEWS
# ============================================================
@api_router.post("/interviews")
async def create_interview(interview: InterviewCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    res = await run(lambda: sb("interviews").insert({
        "candidate_id":  interview.candidate_id,
        "job_id":        interview.job_id,
        "scheduled_at":  interview.scheduled_at,
        "interview_type":interview.interview_type,
        "interviewers":  interview.interviewers,
        "notes":         interview.notes,
        "created_by":    user["id"],
    }).execute())

    # Move candidate to interview_scheduled
    await run(lambda: sb("candidates").update({"status": "interview_scheduled"}).eq("id", interview.candidate_id).execute())
    await _log_activity(candidate_id=interview.candidate_id, user=user, atype="interview",
                        desc=f"{interview.interview_type} scheduled for {interview.scheduled_at}")
    asyncio.create_task(_audit("create", user=user, entity_type="interview",
                                entity_id=res.data[0].get("id"),
                                entity_name=f"{interview.interview_type} — {interview.scheduled_at[:10]}",
                                new_value={"type": interview.interview_type, "scheduled_at": interview.scheduled_at,
                                           "candidate_id": interview.candidate_id},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return res.data[0]


@api_router.get("/interviews")
async def get_interviews(
    request:      Request,
    candidate_id: Optional[str]  = None,
    upcoming:     Optional[bool] = None,
):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    q = sb("interviews").select(
        "*, candidate:candidate_id(full_name), job:job_id(title)"
    ).order("scheduled_at")
    if candidate_id: q = q.eq("candidate_id", candidate_id)
    if upcoming:
        now = datetime.now(timezone.utc).isoformat()
        q = q.gte("scheduled_at", now).eq("completed", False)
    res = await run(lambda: q.execute())
    return res.data or []


@api_router.put("/interviews/{interview_id}")
async def update_interview(interview_id: str, interview: InterviewUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "recruitment")
    patch = {k: v for k, v in interview.model_dump().items() if v is not None}
    if interview.completed:
        patch["completed_at"] = datetime.now(timezone.utc).isoformat()

    await run(lambda: sb("interviews").update(patch).eq("id", interview_id).execute())

    # If marked complete, advance candidate to interviewed
    if interview.completed:
        iv_cand = await safe_single(lambda: sb("interviews").select("candidate_id").eq("id", interview_id).single().execute())
        if iv_cand:
            cand_id = iv_cand["candidate_id"]
            await run(lambda: sb("candidates").update({"status": "interviewed"}).eq("id", cand_id).execute())
            await _log_activity(candidate_id=cand_id, user=user, atype="interview",
                                 desc="Interview completed")
    asyncio.create_task(_audit("update", user=user, entity_type="interview",
                                entity_id=interview_id, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Interview updated"}


# ============================================================
# DASHBOARD
# ============================================================
@api_router.get("/dashboard/sales")
async def sales_dashboard(request: Request):
    user  = await get_current_user(request)
    _require_module(user, "sales")
    today = datetime.now(timezone.utc).date().isoformat()

    # Fetch all leads once — derive stats + pipeline in Python (1 query instead of N+2)
    # Fire all independent queries in parallel
    leads_all_res, tasks_today, overdue_tasks, followups, recent_leads, reminders, await_feedback, subs_pending = await asyncio.gather(
        run(lambda: sb("leads").select("id,status,deal_value,next_follow_up,full_name,company,created_at").execute()),
        run(lambda uid=user["id"]: sb("tasks").select("*").eq("assigned_to", uid).eq("due_date", today).eq("completed", False).execute()),
        run(lambda uid=user["id"]: sb("tasks").select("*").eq("assigned_to", uid).lt("due_date", today).eq("completed", False).execute()),
        run(lambda: sb("leads").select("*").eq("next_follow_up", today).execute()),
        run(lambda: sb("leads").select("*").order("created_at", desc=True).limit(10).execute()),
        run(lambda uid=user["id"]: sb("reminders").select("*").eq("user_id", uid).gte("due_date", today).eq("dismissed", False).order("due_date").limit(10).execute()),
        run(lambda: sb("interviews").select("id,interview_type,scheduled_at,candidate:candidate_id(full_name),job:job_id(title)").eq("completed", True).is_("rating", "null").order("scheduled_at", desc=True).limit(10).execute()),
        run(lambda: sb("candidate_submissions").select("*, candidate:candidate_id(full_name), lead:lead_id(full_name,company)").eq("status", "submitted").order("created_at", desc=True).limit(10).execute()),
    )
    all_leads = leads_all_res.data or []
    # Build lead_stats from the already-fetched data (no extra queries)
    lead_stats = {s.value: 0 for s in LeadStatus}
    for l in all_leads:
        st = l.get("status")
        if st in lead_stats:
            lead_stats[st] += 1
    total_res_count = len(all_leads)
    closed_s    = ["closed","completed","rejected","lost"]
    pipeline_v  = sum(float(l.get("deal_value") or 0) for l in all_leads if l.get("status") not in closed_s)

    # Urgent follow-ups: overdue or due today, sorted by date
    urgent = [l for l in all_leads if l.get("next_follow_up") and l.get("next_follow_up") <= today
              and l.get("status") not in closed_s]
    urgent.sort(key=lambda l: l["next_follow_up"])

    return {
        "lead_stats":       lead_stats,
        "total_leads":      total_res_count,
        "pipeline_value":   pipeline_v,
        "today_tasks":      tasks_today.data or [],
        "overdue_tasks":    overdue_tasks.data or [],
        "today_followups":  followups.data or [],
        "urgent_followups": urgent[:3],
        "recent_leads":     recent_leads.data or [],
        "reminders":        reminders.data or [],
        "awaiting_feedback": await_feedback.data or [],
        "submissions_pending": subs_pending.data or [],
    }


@api_router.get("/dashboard/recruitment")
async def recruitment_dashboard(request: Request):
    user  = await get_current_user(request)
    _require_module(user, "recruitment")
    today = datetime.now(timezone.utc).date().isoformat()

    # Fire all independent queries in parallel
    cands_all_res, active_jobs, upcoming_ivs, tasks_today, recent_cands = await asyncio.gather(
        run(lambda: sb("candidates").select("id,status").execute()),
        run(lambda: sb("jobs").select("id", count="exact").eq("is_active", True).execute()),
        run(lambda: sb("interviews").select("*, candidate:candidate_id(full_name), job:job_id(title)").gte("scheduled_at", today).eq("completed", False).order("scheduled_at").limit(10).execute()),
        run(lambda uid=user["id"]: sb("tasks").select("*").eq("assigned_to", uid).eq("due_date", today).eq("completed", False).execute()),
        run(lambda: sb("candidates").select("*").order("created_at", desc=True).limit(10).execute()),
    )
    # Build cand_stats from already-fetched data (no extra queries)
    cand_stats = {s.value: 0 for s in CandidateStatus}
    for c in (cands_all_res.data or []):
        st = c.get("status")
        if st in cand_stats:
            cand_stats[st] += 1
    total_cands_count = len(cands_all_res.data or [])

    return {
        "candidate_stats":    cand_stats,
        "total_candidates":   total_cands_count,
        "active_jobs":        active_jobs.count or 0,
        "upcoming_interviews":upcoming_ivs.data or [],
        "today_tasks":        tasks_today.data or [],
        "recent_candidates":  recent_cands.data or [],
    }


# ============================================================
# EMAIL — send reminder
# ============================================================
@api_router.post("/reminders/{reminder_id}/send-email")
async def send_reminder_email(reminder_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    reminder_data = await safe_single(lambda: sb("reminders").select("*").eq("id", reminder_id).single().execute())
    if not reminder_data:
        raise HTTPException(404, "Reminder not found")
    r = reminder_data
    html = f"""
    <div style="font-family:Inter,sans-serif;padding:24px;max-width:480px">
      <h2 style="color:#131b2e;margin-bottom:8px">⏰ {r['title']}</h2>
      <p style="color:#434655">Due: <strong>{r['due_date']}</strong>{' at ' + r['due_time'] if r.get('due_time') else ''}</p>
      {f'<p style="color:#434655">{r["note"]}</p>' if r.get('note') else ''}
      <p style="color:#737686;font-size:13px;margin-top:24px">Nexus CRM Platform</p>
    </div>
    """
    result = await send_email(r.get("user_email", user["email"]), f"Reminder: {r['title']}", html)
    await run(lambda: sb("reminders").update({"email_sent": True}).eq("id", reminder_id).execute())
    return result



# ============================================================
# DAILY DIGEST EMAIL — sent every morning at 8 AM
# ============================================================
async def send_daily_digest():
    """Runs at 8 AM daily. Sends each user a digest of today's tasks + reminders."""
    today = datetime.now(timezone.utc).date().isoformat()
    logger.info(f"[digest] Running daily digest for {today}")

    try:
        users_res = await run(lambda: sb("users").select("id,email,name,role").execute())
        users_list = users_res.data or []
    except Exception as e:
        logger.error(f"[digest] Failed to fetch users: {e}")
        return

    for u in users_list:
        try:
            # Tasks due today for this user
            tasks_res = await run(lambda: sb("tasks")
                .select("title,task_type,priority,due_time")
                .eq("assigned_to", u["id"])
                .eq("due_date", today)
                .eq("completed", False)
                .order("due_time")
                .execute()
            )
            tasks_today = tasks_res.data or []

            # Reminders due today for this user
            reminders_res = await run(lambda: sb("reminders")
                .select("title,due_time,note")
                .eq("user_id", u["id"])
                .eq("due_date", today)
                .eq("dismissed", False)
                .order("due_time")
                .execute()
            )
            reminders_today = reminders_res.data or []

            # Skip if nothing due today
            if not tasks_today and not reminders_today:
                logger.info(f"[digest] Nothing due for {u['email']} — skipping")
                continue

            total = len(tasks_today) + len(reminders_today)

            # Build tasks rows
            tasks_rows = ""
            for t in tasks_today:
                priority_color = {"high": "#dc2626", "medium": "#d97706", "low": "#16a34a"}.get(t.get("priority", "medium"), "#6b7280")
                time_str = f" &middot; {t['due_time'][:5]}" if t.get("due_time") else ""
                tasks_rows += (
                    f'<tr><td style="padding:10px 0;border-bottom:1px solid #f1f5f9">'
                    f'<span style="font-weight:600;color:#1e293b">{t["title"]}</span>'
                    f'<span style="color:#94a3b8;font-size:13px">{time_str}</span></td>'
                    f'<td style="padding:10px 0;border-bottom:1px solid #f1f5f9;text-align:right">'
                    f'<span style="font-size:12px;font-weight:700;color:{priority_color};text-transform:uppercase">'
                    f'{t.get("priority","")}</span></td></tr>'
                )

            # Build reminders rows
            reminders_rows = ""
            for r in reminders_today:
                time_str = f" &middot; {r['due_time'][:5]}" if r.get("due_time") else ""
                note_str = f'<br><span style="color:#94a3b8;font-size:12px">{r["note"]}</span>' if r.get("note") else ""
                reminders_rows += (
                    f'<tr><td colspan="2" style="padding:10px 0;border-bottom:1px solid #f1f5f9">'
                    f'<span style="font-weight:600;color:#1e293b">&#9200; {r["title"]}</span>'
                    f'<span style="color:#94a3b8;font-size:13px">{time_str}</span>{note_str}</td></tr>'
                )

            tasks_section = ""
            if tasks_today:
                tasks_section = (
                    f'<h2 style="font-size:14px;font-weight:700;color:#475569;text-transform:uppercase;'
                    f'letter-spacing:0.05em;margin:0 0 8px">Tasks ({len(tasks_today)})</h2>'
                    f'<table style="width:100%;border-collapse:collapse;margin-bottom:24px">{tasks_rows}</table>'
                )

            reminders_section = ""
            if reminders_today:
                reminders_section = (
                    f'<h2 style="font-size:14px;font-weight:700;color:#475569;text-transform:uppercase;'
                    f'letter-spacing:0.05em;margin:0 0 8px">Reminders ({len(reminders_today)})</h2>'
                    f'<table style="width:100%;border-collapse:collapse;margin-bottom:24px">{reminders_rows}</table>'
                )

            first_name = u["name"].split()[0]
            subject = f"Your schedule for today \u2014 {total} item{'s' if total != 1 else ''}"

            html = (
                f'<div style="font-family:Inter,Arial,sans-serif;max-width:560px;margin:0 auto;padding:32px 24px">'
                f'<div style="background:linear-gradient(135deg,#1e3a5f,#2563eb);border-radius:12px;padding:24px;margin-bottom:28px">'
                f'<h1 style="color:#ffffff;margin:0;font-size:22px">Good morning, {first_name} &#128075;</h1>'
                f'<p style="color:#bfdbfe;margin:8px 0 0;font-size:14px">Here\'s your schedule for today &mdash; '
                f'<strong style="color:#ffffff">{total} item{"s" if total != 1 else ""}</strong> on your plate.</p>'
                f'</div>'
                f'{tasks_section}{reminders_section}'
                f'<p style="font-size:12px;color:#94a3b8;text-align:center;margin-top:32px;'
                f'border-top:1px solid #f1f5f9;padding-top:16px">Nexus CRM &nbsp;&middot;&nbsp; Automated daily digest</p>'
                f'</div>'
            )

            result = await send_email(u["email"], subject, html)
            logger.info(f"[digest] Sent to {u['email']}: {result.get('status')}")

        except Exception as e:
            logger.error(f"[digest] Failed for {u['email']}: {e}")

    # ── Weekly CEO summary — every Saturday only ─────────────
    if datetime.now(timezone.utc).weekday() == 5:  # 5 = Saturday
        try:
            admin_users = [u for u in users_list if u.get("role") == "admin"]
            for admin in admin_users:
                week_ago  = (datetime.now(timezone.utc).date() - timedelta(days=7)).isoformat()
                new_leads = await run(lambda: sb("leads").select("id", count="exact").gte("created_at", week_ago + "T00:00:00Z").execute())
                closed    = await run(lambda: sb("leads").select("id,deal_value").in_("status", ["closed","completed"]).gte("updated_at", week_ago + "T00:00:00Z").execute())
                subs      = await run(lambda: sb("candidate_submissions").select("id", count="exact").gte("created_at", week_ago + "T00:00:00Z").execute())
                active    = await run(lambda: sb("leads").select("deal_value").not_.in_("status", ["closed","completed","rejected","lost"]).execute())
                pipeline_v = sum(float(l.get("deal_value") or 0) for l in (active.data or []))
                closed_v   = sum(float(l.get("deal_value") or 0) for l in (closed.data or []))

                html_ceo = (
                    f'<div style="font-family:Inter,sans-serif;padding:24px;max-width:560px;background:#f8fafc">'
                    f'<h2 style="color:#131b2e;margin-bottom:4px">📊 Weekly Business Summary</h2>'
                    f'<p style="color:#737686;font-size:13px;margin-bottom:20px">Week ending {today}</p>'
                    f'<table style="width:100%;border-collapse:collapse">'
                    f'<tr style="background:#fff"><td style="padding:12px;border:1px solid #e2e8f0;font-weight:600">🆕 New Leads</td><td style="padding:12px;border:1px solid #e2e8f0;font-weight:700">{new_leads.count or 0}</td></tr>'
                    f'<tr style="background:#f8fafc"><td style="padding:12px;border:1px solid #e2e8f0;font-weight:600">💰 Active Pipeline</td><td style="padding:12px;border:1px solid #e2e8f0;font-weight:700">₹{pipeline_v:,.0f}</td></tr>'
                    f'<tr style="background:#fff"><td style="padding:12px;border:1px solid #e2e8f0;font-weight:600">✅ Deals Closed</td><td style="padding:12px;border:1px solid #e2e8f0;font-weight:700">{len(closed.data or [])} deals · ₹{closed_v:,.0f}</td></tr>'
                    f'<tr style="background:#f8fafc"><td style="padding:12px;border:1px solid #e2e8f0;font-weight:600">👤 Candidates Submitted</td><td style="padding:12px;border:1px solid #e2e8f0;font-weight:700">{subs.count or 0}</td></tr>'
                    f'</table>'
                    f'<p style="color:#737686;font-size:12px;margin-top:24px">Nexus CRM · Auto-generated every Saturday 8 AM</p>'
                    f'</div>'
                )
                await send_email(admin["email"], f"📊 Weekly Summary — {today}", html_ceo)
                logger.info(f"[digest] CEO weekly summary sent to {admin['email']}")
        except Exception as ce:
            logger.error(f"[digest] CEO weekly summary failed: {ce}")


# ============================================================
# LEAD ENRICHMENT — Apify LinkedIn scraper
# ============================================================
class EnrichRequest(BaseModel):
    linkedin_urls: List[str]

@api_router.post("/enrich/start")
async def enrich_start(body: EnrichRequest, request: Request):
    """Start an Apify enrichment run. Returns run_id to poll."""
    user = await get_current_user(request)
    _require_module(user, "sales")

    if not APIFY_API_KEY:
        raise HTTPException(503, "Apify API key not configured on server. Contact admin.")
    if not body.linkedin_urls:
        raise HTTPException(400, "No LinkedIn URLs provided.")

    urls = [u.strip().rstrip("/") for u in body.linkedin_urls if u.strip()]
    if not urls:
        raise HTTPException(400, "All LinkedIn URLs were empty.")

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.post(
            f"https://api.apify.com/v2/acts/{APIFY_ACTOR_ID}/runs?token={APIFY_API_KEY}",
            json={"profileUrls": urls},
            headers={"Content-Type": "application/json"},
        )
    if resp.status_code not in (200, 201):
        detail = resp.json().get("error", {}).get("message", resp.text[:200])
        raise HTTPException(502, f"Apify error: {detail}")

    run_id = resp.json().get("data", {}).get("id")
    if not run_id:
        raise HTTPException(502, "Apify did not return a run ID.")

    return {"run_id": run_id, "total": len(urls)}


@api_router.get("/enrich/status/{run_id}")
async def enrich_status(run_id: str, request: Request):
    """Poll Apify run status. Returns results when SUCCEEDED."""
    user = await get_current_user(request)
    _require_module(user, "sales")

    if not APIFY_API_KEY:
        raise HTTPException(503, "Apify API key not configured.")

    async with httpx.AsyncClient(timeout=20) as client:
        status_resp = await client.get(
            f"https://api.apify.com/v2/actor-runs/{run_id}?token={APIFY_API_KEY}"
        )

    if status_resp.status_code != 200:
        raise HTTPException(502, "Could not check Apify run status.")

    data   = status_resp.json().get("data", {})
    status = data.get("status", "UNKNOWN")
    stats  = data.get("stats", {})

    if status == "SUCCEEDED":
        async with httpx.AsyncClient(timeout=30) as client:
            items_resp = await client.get(
                f"https://api.apify.com/v2/actor-runs/{run_id}/dataset/items"
                f"?token={APIFY_API_KEY}&format=json"
            )
        items = items_resp.json() if items_resp.status_code == 200 else []

        # Normalise results: linkedin_url -> { email, phone }
        results = {}
        for item in (items if isinstance(items, list) else []):
            url = (item.get("profileUrl") or item.get("linkedinUrl") or item.get("url") or "")
            url = url.strip().rstrip("/").lower()
            if url:
                results[url] = {
                    "email": item.get("email") or (item.get("emails") or [None])[0],
                    "phone": item.get("phone") or (item.get("phoneNumbers") or [None])[0]
                             or (item.get("phones") or [None])[0],
                }
        return {"status": status, "results": results}

    return {
        "status": status,
        "processed": stats.get("itemsFinished", 0),
    }



# ============================================================
# CANDIDATE SUBMISSIONS (Lead ↔ Candidate link)
# ============================================================
@api_router.post("/submissions")
async def create_submission(body: SubmissionCreate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    doc = {
        "lead_id":      body.lead_id,
        "candidate_id": body.candidate_id,
        "status":       body.status,
        "notes":        body.notes,
        "created_by":   user["id"],
    }
    res = await run(lambda: sb("candidate_submissions").upsert(doc, on_conflict="lead_id,candidate_id").execute())
    await _log_activity(lead_id=body.lead_id, user=user, atype="note",
                        desc=f"Candidate submission created/updated")
    row = res.data[0] if res.data else {}
    asyncio.create_task(_audit("create", user=user, entity_type="submission",
                                entity_id=row.get("id"),
                                entity_name=f"Submission — lead:{body.lead_id}",
                                new_value={"lead_id": body.lead_id, "candidate_id": body.candidate_id, "status": body.status},
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return row

@api_router.get("/submissions")
async def get_submissions(request: Request, lead_id: Optional[str] = None, candidate_id: Optional[str] = None):
    user = await get_current_user(request)
    _require_module(user, "sales")
    q = sb("candidate_submissions").select(
        "*, candidate:candidate_id(id,full_name,candidate_role,status,email), lead:lead_id(id,full_name,company,status)"
    ).order("created_at", desc=True)
    if lead_id:      q = q.eq("lead_id", lead_id)
    if candidate_id: q = q.eq("candidate_id", candidate_id)
    res = await run(lambda: q.execute())
    return res.data or []

@api_router.put("/submissions/{submission_id}")
async def update_submission(submission_id: str, body: SubmissionUpdate, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    patch = {k: v for k, v in body.model_dump().items() if v is not None}
    await run(lambda: sb("candidate_submissions").update(patch).eq("id", submission_id).execute())
    asyncio.create_task(_audit("update", user=user, entity_type="submission",
                                entity_id=submission_id, new_value=patch,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Updated"}

@api_router.delete("/submissions/{submission_id}")
async def delete_submission(submission_id: str, request: Request):
    user = await get_current_user(request)
    _require_module(user, "sales")
    await run(lambda: sb("candidate_submissions").delete().eq("id", submission_id).execute())
    asyncio.create_task(_audit("delete", user=user, entity_type="submission",
                                entity_id=submission_id,
                                ip=_get_ip(request), ua=request.headers.get("user-agent","")))
    return {"message": "Deleted"}


# ============================================================
# TIMESHEET YEARLY SUMMARY  (CEO/Viewer chart data)
# ============================================================
@api_router.get("/timesheet/yearly-summary")
async def yearly_timesheet_summary(request: Request, year: Optional[int] = None):
    """
    Returns approved hours aggregated by employee + month for a given year.
    Used for the CEO monthly bar chart. One row per (user, month) — very light.
    """
    user = await get_current_user(request)
    if user.get("role") not in ("admin", "viewer"):
        raise HTTPException(403, "Access restricted to admin/viewer roles.")
    if not year:
        year = datetime.now(timezone.utc).year
    res = await run(lambda: supabase.rpc(
        "get_yearly_timesheet_summary", {"p_year": year}
    ).execute())
    return {"data": res.data or []}


# ============================================================
# TUTORIALS  (first-time onboarding — stored per user in Supabase)
# ============================================================
@api_router.get("/tutorials")
async def get_tutorials(request: Request):
    """Return list of tutorial page keys the current user has already seen."""
    user = await get_current_user(request)
    res = await run(lambda: sb("user_tutorials")
        .select("page")
        .eq("user_id", user["id"])
        .execute())
    return {"pages": [r["page"] for r in (res.data or [])]}


@api_router.post("/tutorials/{page}")
async def mark_tutorial_seen(page: str, request: Request):
    """Mark a tutorial page as seen for the current user (upsert — safe to call multiple times)."""
    user = await get_current_user(request)
    allowed_pages = {"timesheet", "leads", "candidates", "approvals"}
    if page not in allowed_pages:
        raise HTTPException(400, f"Unknown tutorial page '{page}'")
    await run(lambda: sb("user_tutorials")
        .upsert({"user_id": user["id"], "page": page}, on_conflict="user_id,page")
        .execute())
    return {"ok": True}


# ============================================================
# CEO DASHBOARD
# ============================================================
@api_router.get("/dashboard/ceo")
async def ceo_dashboard(request: Request):
    user = await get_current_user(request)
    if user.get("role") not in ("admin", "viewer"):
        raise HTTPException(403, "CEO dashboard is restricted to admin and CEO accounts")

    today     = datetime.now(timezone.utc).date().isoformat()
    week_ago  = (datetime.now(timezone.utc).date() - timedelta(days=7)).isoformat()
    month_ago = (datetime.now(timezone.utc).date() - timedelta(days=30)).isoformat()
    stale_cutoff = (datetime.now(timezone.utc).date() - timedelta(days=5)).isoformat()

    # Pipeline value (non-closed leads) — fire independent queries in parallel
    closed_statuses = ["closed", "completed", "rejected", "lost"]
    active_leads_res, total_cands, submissions_res, acts_res, recent_audit = await asyncio.gather(
        run(lambda: sb("leads").select("id,full_name,company,status,deal_value,next_follow_up,created_at").execute()),
        run(lambda: sb("candidates").select("id", count="exact").execute()),
        run(lambda: sb("candidate_submissions").select("id,status,created_at").gte("created_at", month_ago + "T00:00:00Z").execute()),
        run(lambda: sb("activities").select("lead_id,created_at").order("created_at", desc=True).execute()),
        run(lambda: sb("audit_logs").select("*").order("created_at", desc=True).limit(20).execute()),
    )
    all_leads = active_leads_res.data or []
    pipeline_leads = [l for l in all_leads if l.get("status") not in closed_statuses]
    closed_leads   = [l for l in all_leads if l.get("status") in ["closed", "completed"]]
    pipeline_value = sum(float(l.get("deal_value") or 0) for l in pipeline_leads)
    closed_value   = sum(float(l.get("deal_value") or 0) for l in closed_leads)

    # Leads by stage
    stage_counts = {}
    for l in all_leads:
        s = l.get("status", "new")
        stage_counts[s] = stage_counts.get(s, 0) + 1

    submissions = submissions_res.data or []
    acts = acts_res.data or []
    last_act = {}
    for a in acts:
        lid = a["lead_id"]
        if lid and lid not in last_act:
            last_act[lid] = a["created_at"][:10]  # just date

    stale_leads = []
    for l in pipeline_leads:
        lid = l["id"]
        last = last_act.get(lid)
        if not last or last < stale_cutoff:
            days_stale = (datetime.now(timezone.utc).date() - (
                datetime.fromisoformat(last).date() if last else
                datetime.fromisoformat(l["created_at"][:10]).date()
            )).days
            stale_leads.append({**l, "days_stale": days_stale, "last_activity": last})

    stale_leads.sort(key=lambda x: x["days_stale"], reverse=True)

    # This week's activity
    leads_this_week = [l for l in all_leads if l.get("created_at","")[:10] >= week_ago]
    subs_this_week  = [s for s in submissions if s.get("created_at","")[:10] >= week_ago]

    return {
        "pipeline_value":    pipeline_value,
        "closed_value":      closed_value,
        "total_leads":       len(all_leads),
        "pipeline_leads":    len(pipeline_leads),
        "closed_leads":      len(closed_leads),
        "stage_counts":      stage_counts,
        "total_candidates":  total_cands.count or 0,
        "submissions_month": len(submissions),
        "submissions_week":  len(subs_this_week),
        "stale_leads":       stale_leads[:10],
        "stale_count":       len(stale_leads),
        "leads_this_week":   len(leads_this_week),
        "recent_audit":      recent_audit.data or [],
    }


# ============================================================
# AUDIT LOG — read-only, admin only
# ============================================================
@api_router.get("/audit-logs")
async def get_audit_logs(
    request:     Request,
    action:      Optional[str] = None,
    entity_type: Optional[str] = None,
    user_id:     Optional[str] = None,
    user_name:   Optional[str] = None,   # search by user name
    date_from:   Optional[str] = None,   # YYYY-MM-DD
    date_to:     Optional[str] = None,   # YYYY-MM-DD
    entity_name: Optional[str] = None,   # search by record name
    limit:       int = 100,
    skip:        int = 0,
):
    caller = await get_current_user(request)
    if caller.get("role") not in ("admin", "viewer"):
        raise HTTPException(403, "Audit log is restricted to admin and CEO accounts")

    q = sb("audit_logs").select("*", count="exact").order("created_at", desc=True).range(skip, skip + limit - 1)
    if action:      q = q.eq("action", action)
    if entity_type: q = q.eq("entity_type", entity_type)
    if user_id:     q = q.eq("user_id", user_id)
    if user_name:   q = q.ilike("user_name", f"%{user_name}%")
    if entity_name: q = q.ilike("entity_name", f"%{entity_name}%")
    if date_from:   q = q.gte("created_at", f"{date_from}T00:00:00Z")
    if date_to:     q = q.lte("created_at", f"{date_to}T23:59:59Z")

    res = await run(lambda: q.execute())
    return {"logs": res.data or [], "total": res.count or 0}


# ============================================================
# AUDIT LOG — frontend event endpoint
# ============================================================
class FrontendAuditRequest(BaseModel):
    action:      str
    entity_type: Optional[str] = None
    entity_name: Optional[str] = None
    new_value:   Optional[dict] = None

@api_router.post("/audit/log")
async def log_frontend_audit(body: FrontendAuditRequest, request: Request):
    """Accept audit events triggered client-side (e.g. data exports)."""
    user = await get_current_user(request)
    await _audit(
        action=body.action, user=user,
        entity_type=body.entity_type, entity_name=body.entity_name,
        new_value=body.new_value,
        ip=_get_ip(request), ua=request.headers.get("user-agent", ""),
    )
    return {"ok": True}


# ============================================================
# HEALTH CHECK
# ============================================================
@api_router.api_route("/health", methods=["GET", "HEAD"])
async def health():
    return {"status": "ok", "service": "Nexus CRM + ATS", "version": "2.0.0"}


# ============================================================
# STARTUP — seed admin user
# ============================================================

# ============================================================
# TIMESHEET — Models + Endpoints
# ============================================================

class TimesheetEntryUpsert(BaseModel):
    entry_date: str          # ISO date string YYYY-MM-DD
    hours:      float = 0.0
    comments:   Optional[str] = None

class TimesheetCreate(BaseModel):
    week_start: str          # ISO date string YYYY-MM-DD (must be a Friday)
    entries:    list[TimesheetEntryUpsert] = []

class TimesheetReview(BaseModel):
    action: str              # "approve" | "reject"
    note:   Optional[str] = None


def _to_decimal_hours(value) -> Decimal:
    try:
        return Decimal(str(value or 0))
    except (InvalidOperation, TypeError, ValueError):
        raise HTTPException(400, "Invalid hours value.")


def _normalize_timesheet_hours(entry: dict) -> Decimal:
    """
    Accept both legacy decimal-hours payloads and the new hour/minute payload.
    Minutes are converted into decimal hours for storage in the existing numeric column.
    """
    if not isinstance(entry, dict):
        raise HTTPException(400, "Invalid timesheet entry.")

    if "minutes" in entry or "hours_value" in entry:
        raw_hours = entry.get("hours_value", 0)
        raw_minutes = entry.get("minutes", 0)
        try:
            hours_value = int(raw_hours or 0)
            minutes_value = int(raw_minutes or 0)
        except (TypeError, ValueError):
            raise HTTPException(400, "Hours and minutes must be whole numbers.")

        if hours_value < 0:
            raise HTTPException(400, "Hours cannot be negative.")
        if minutes_value < 0 or minutes_value > 59:
            raise HTTPException(400, "Minutes must be between 0 and 59.")

        total = Decimal(hours_value) + (Decimal(minutes_value) / Decimal(60))
    else:
        total = _to_decimal_hours(entry.get("hours", 0))

    if total < 0:
        raise HTTPException(400, "Hours cannot be negative.")

    return total.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def _week_friday(d: datetime) -> str:
    """Return ISO string of the Friday of the week containing d."""
    # Friday is weekday 4 (Mon=0..Sun=6). Roll back to most recent Friday.
    return (d.date() - timedelta(days=(d.weekday() - 4) % 7)).isoformat()


@api_router.get("/timesheets/me")
async def get_my_timesheets(request: Request, week_start: Optional[str] = None):
    """Get current user's timesheets. Optionally filter by week_start."""
    user = await get_current_user(request)
    q = sb("timesheets").select("*").eq("user_id", user["id"]).order("week_start", desc=True)
    if week_start:
        q = q.eq("week_start", week_start)
    res = await run(lambda: q.execute())
    timesheets = res.data or []
    # Attach entries for each timesheet
    for ts in timesheets:
        ent = await run(lambda tid=ts["id"]: sb("timesheet_entries").select("*").eq("timesheet_id", tid).order("entry_date").execute())
        ts["entries"] = ent.data or []
    return {"timesheets": timesheets, "total": len(timesheets)}


@api_router.get("/timesheets/me/current")
async def get_current_week_timesheet(request: Request):
    """Get (or create) the timesheet for the current week."""
    user = await get_current_user(request)
    week_start = _week_friday(datetime.now(timezone.utc))
    res = await run(lambda: sb("timesheets").select("*").eq("user_id", user["id"]).eq("week_start", week_start).execute())
    if res.data:
        ts = res.data[0]
    else:
        ins = await run(lambda: sb("timesheets").insert({"user_id": user["id"], "week_start": week_start, "status": "draft"}).execute())
        ts = ins.data[0]
    ent = await run(lambda tid=ts["id"]: sb("timesheet_entries").select("*").eq("timesheet_id", tid).order("entry_date").execute())
    ts["entries"] = ent.data or []
    return ts


@api_router.get("/timesheets/me/week")
async def get_week_timesheet(week_start: str, request: Request):
    """Get (or create) the timesheet for a specific week."""
    user = await get_current_user(request)
    res = await run(lambda: sb("timesheets").select("*").eq("user_id", user["id"]).eq("week_start", week_start).execute())
    if res.data:
        ts = res.data[0]
    else:
        ins = await run(lambda: sb("timesheets").insert({"user_id": user["id"], "week_start": week_start, "status": "draft"}).execute())
        ts = ins.data[0]
    ent = await run(lambda tid=ts["id"]: sb("timesheet_entries").select("*").eq("timesheet_id", tid).order("entry_date").execute())
    ts["entries"] = ent.data or []
    return ts


@api_router.put("/timesheets/{timesheet_id}/entries")
async def upsert_timesheet_entries(timesheet_id: str, body: dict, request: Request):
    """Save/update daily entries for a timesheet (draft only)."""
    user = await get_current_user(request)
    ts_res = await run(lambda: sb("timesheets").select("*").eq("id", timesheet_id).eq("user_id", user["id"]).execute())
    if not ts_res.data:
        raise HTTPException(404, "Timesheet not found")
    ts = ts_res.data[0]
    if ts["status"] == "submitted" or ts["status"] == "approved":
        raise HTTPException(400, f"Cannot edit a {ts['status']} timesheet")

    entries = body.get("entries", [])
    total_hours = Decimal("0.00")
    for e in entries:
        if not e.get("entry_date"):
            raise HTTPException(400, "Each timesheet entry must include entry_date.")
        hours_decimal = _normalize_timesheet_hours(e)
        total_hours += hours_decimal
        await run(lambda ed=e["entry_date"], eh=float(hours_decimal), ec=e.get("comments"): sb("timesheet_entries").upsert({
            "timesheet_id": timesheet_id,
            "entry_date":   ed,
            "hours":        eh,
            "comments":     ec,
        }, on_conflict="timesheet_id,entry_date").execute())

    await run(lambda: sb("timesheets").update({"total_hours": float(total_hours), "updated_at": datetime.now(timezone.utc).isoformat()}).eq("id", timesheet_id).execute())
    await _audit("timesheet_saved", entity_name=f"Week of {ts['week_start']}", user=user)
    return {"success": True, "total_hours": float(total_hours)}


@api_router.post("/timesheets/{timesheet_id}/submit")
async def submit_timesheet(timesheet_id: str, request: Request):
    """Submit a timesheet — sends email to CEO."""
    user = await get_current_user(request)
    ts_res = await run(lambda: sb("timesheets").select("*").eq("id", timesheet_id).eq("user_id", user["id"]).execute())
    if not ts_res.data:
        raise HTTPException(404, "Timesheet not found")
    ts = ts_res.data[0]
    if ts["status"] not in ("draft",):
        raise HTTPException(400, f"Timesheet is already {ts['status']}")

    now_iso = datetime.now(timezone.utc).isoformat()
    await run(lambda: sb("timesheets").update({"status": "submitted", "submitted_at": now_iso, "updated_at": now_iso}).eq("id", timesheet_id).execute())

    # Get entries for email
    ent_res = await run(lambda: sb("timesheet_entries").select("*").eq("timesheet_id", timesheet_id).order("entry_date").execute())
    entries = ent_res.data or []
    rows_html = ""
    for e in entries:
        h = float(e.get("hours") or 0)
        if h > 0 or e.get("comments"):
            day_name = datetime.fromisoformat(e["entry_date"]).strftime("%A, %b %d")
            rows_html += f"<tr><td style='padding:6px 12px;border-bottom:1px solid #f1f5f9'>{day_name}</td><td style='padding:6px 12px;border-bottom:1px solid #f1f5f9;font-weight:600'>{h:.1f}h</td><td style='padding:6px 12px;border-bottom:1px solid #f1f5f9;color:#64748b'>{e.get('comments','')}</td></tr>"

    total_h = float(ts.get("total_hours") or 0)
    html_ceo = f"""
<div style="font-family:Inter,sans-serif;max-width:600px;margin:0 auto">
  <div style="background:linear-gradient(135deg,#004ac6,#0066ff);padding:32px;border-radius:12px 12px 0 0;color:#fff">
    <h1 style="margin:0;font-size:1.5rem">📋 Timesheet Submitted</h1>
    <p style="margin:8px 0 0;opacity:0.85">{user['name']} submitted their timesheet for review</p>
  </div>
  <div style="background:#fff;padding:24px;border:1px solid #e2e8f0;border-top:none">
    <p><strong>Employee:</strong> {user['name']} ({user['email']})</p>
    <p><strong>Week of:</strong> {ts['week_start']}</p>
    <p><strong>Total Hours:</strong> {total_h:.1f}h</p>
    <table style="width:100%;border-collapse:collapse;margin-top:16px">
      <thead><tr style="background:#f8fafc">
        <th style="padding:8px 12px;text-align:left;font-size:0.75rem;text-transform:uppercase;letter-spacing:.05em">Day</th>
        <th style="padding:8px 12px;text-align:left;font-size:0.75rem;text-transform:uppercase;letter-spacing:.05em">Hours</th>
        <th style="padding:8px 12px;text-align:left;font-size:0.75rem;text-transform:uppercase;letter-spacing:.05em">Notes</th>
      </tr></thead>
      <tbody>{rows_html}</tbody>
    </table>
    <div style="margin-top:24px;padding:16px;background:#f0f4ff;border-radius:8px">
      <p style="margin:0;font-size:0.875rem">Please log into <strong>Nexus CRM</strong> to approve or reject this timesheet.</p>
    </div>
  </div>
</div>"""

    # Email the CEO (viewer role)
    ceos = await run(lambda: sb("users").select("email,name").eq("role", "viewer").execute())
    for ceo in (ceos.data or []):
        await send_email(ceo["email"], f"📋 Timesheet Submitted — {user['name']} (Week of {ts['week_start']})", html_ceo)

    await _audit("timesheet_submitted", entity_name=f"Week of {ts['week_start']}", user=user)
    return {"success": True, "status": "submitted"}


@api_router.post("/timesheets/{timesheet_id}/review")
async def review_timesheet(timesheet_id: str, body: TimesheetReview, request: Request):
    """CEO approves or rejects a timesheet."""
    reviewer = await get_current_user(request)
    if reviewer.get("role") != "viewer":
        raise HTTPException(403, "Only the CEO can review timesheets")

    action = body.action.lower()
    if action not in ("approve", "reject"):
        raise HTTPException(400, "Action must be 'approve' or 'reject'")

    ts_res = await run(lambda: sb("timesheets").select("*").eq("id", timesheet_id).execute())
    if not ts_res.data:
        raise HTTPException(404, "Timesheet not found")
    ts = ts_res.data[0]
    if ts["status"] != "submitted":
        raise HTTPException(400, f"Can only review submitted timesheets (current: {ts['status']})")

    new_status = "approved" if action == "approve" else "rejected"
    now_iso = datetime.now(timezone.utc).isoformat()
    await run(lambda: sb("timesheets").update({
        "status": new_status, "note": body.note,
        "reviewed_at": now_iso, "reviewed_by": reviewer["id"], "updated_at": now_iso
    }).eq("id", timesheet_id).execute())

    # Email the worker
    worker_res = await run(lambda: sb("users").select("email,name").eq("id", ts["user_id"]).execute())
    worker = (worker_res.data or [{}])[0]
    icon = "✅" if new_status == "approved" else "❌"
    color = "#16a34a" if new_status == "approved" else "#dc2626"
    note_html = f"<p style='background:#f8fafc;padding:12px;border-radius:8px;border-left:3px solid {color}'><strong>Note from CEO:</strong> {body.note}</p>" if body.note else ""
    html_worker = f"""
<div style="font-family:Inter,sans-serif;max-width:600px;margin:0 auto">
  <div style="background:{color};padding:32px;border-radius:12px 12px 0 0;color:#fff">
    <h1 style="margin:0;font-size:1.5rem">{icon} Timesheet {new_status.capitalize()}</h1>
    <p style="margin:8px 0 0;opacity:0.85">Your timesheet for week of {ts['week_start']} has been {new_status}</p>
  </div>
  <div style="background:#fff;padding:24px;border:1px solid #e2e8f0;border-top:none">
    <p>Hi {worker.get('name','there')},</p>
    <p>Your timesheet for the week of <strong>{ts['week_start']}</strong> ({float(ts.get('total_hours',0)):.1f} hours) has been <strong>{new_status}</strong> by {reviewer['name']}.</p>
    {note_html}
  </div>
</div>"""
    if worker.get("email"):
        await send_email(worker["email"], f"{icon} Your timesheet has been {new_status} — Week of {ts['week_start']}", html_worker)

    await _audit(f"timesheet_{new_status}", entity_name=f"Week of {ts['week_start']}", user=reviewer)
    return {"success": True, "status": new_status}


@api_router.get("/timesheets/all")
async def get_all_timesheets(
    request: Request,
    status: Optional[str] = None,
    user_id: Optional[str] = None,
    week_start: Optional[str] = None,
    limit: int = 50,
):
    """CEO view — all timesheets across all users."""
    reviewer = await get_current_user(request)
    if reviewer.get("role") not in ("viewer",):
        raise HTTPException(403, "Only the CEO can view all timesheets")

    q = sb("timesheets").select("*,users!timesheets_user_id_fkey(id,name,email,role)").order("week_start", desc=True).limit(limit)
    if status:   q = q.eq("status", status)
    if user_id:  q = q.eq("user_id", user_id)
    if week_start: q = q.eq("week_start", week_start)
    res = await run(lambda: q.execute())
    timesheets = res.data or []
    # Attach entries
    for ts in timesheets:
        ent = await run(lambda tid=ts["id"]: sb("timesheet_entries").select("*").eq("timesheet_id", tid).order("entry_date").execute())
        ts["entries"] = ent.data or []
    return {"timesheets": timesheets, "total": len(timesheets)}


@api_router.get("/timesheets/{timesheet_id}")
async def get_timesheet_detail(timesheet_id: str, request: Request):
    """Get a single timesheet with entries (CEO or owner)."""
    user = await get_current_user(request)
    q = sb("timesheets").select("*,users!timesheets_user_id_fkey(id,name,email)").eq("id", timesheet_id)
    res = await run(lambda: q.execute())
    if not res.data:
        raise HTTPException(404, "Not found")
    ts = res.data[0]
    # Only owner or admin/viewer can view
    if ts["user_id"] != user["id"] and user.get("role") not in ("admin", "viewer"):
        raise HTTPException(403, "Access denied")
    ent = await run(lambda: sb("timesheet_entries").select("*").eq("timesheet_id", timesheet_id).order("entry_date").execute())
    ts["entries"] = ent.data or []
    return ts


# ============================================================
# AUDIT LOG CLEANUP — keeps last 180 days, runs daily at 3 AM
# ============================================================
async def cleanup_audit_logs():
    """Delete audit_logs older than 180 days to keep the DB lean."""
    cutoff = (datetime.now(timezone.utc) - timedelta(days=180)).isoformat()
    try:
        res = await run(lambda: sb("audit_logs").delete().lt("created_at", cutoff).execute())
        deleted = len(res.data) if res.data else 0
        logger.info(f"[audit-cleanup] Deleted {deleted} log entries older than 180 days (cutoff: {cutoff[:10]})")
    except Exception as e:
        logger.error(f"[audit-cleanup] Failed: {e}")

async def send_timesheet_reminder():
    """Every Friday 5PM - remind all users to submit their timesheet."""
    today = datetime.now(timezone.utc)
    # Remind for the week that ended Thursday (week starts on Friday).
    current_week_start = _week_friday(today)
    prev_week_start = (datetime.fromisoformat(current_week_start).date() - timedelta(days=7)).isoformat()
    logger.info(f"[timesheet-reminder] Sending Friday reminders for week {prev_week_start}")
    try:
        users_res = await run(lambda: sb("users").select("id,email,name,role").execute())
        users_list = users_res.data or []
    except Exception as e:
        logger.error(f"[timesheet-reminder] Failed to fetch users: {e}")
        return

    for u in users_list:
        try:
            # Check if they've already submitted this week
            ts_res = await run(lambda uid=u["id"]: sb("timesheets").select("status").eq("user_id", uid).eq("week_start", prev_week_start).execute())
            existing = ts_res.data[0] if ts_res.data else None
            if existing and existing.get("status") in ("submitted", "approved"):
                continue  # Already submitted - no need to remind

            html = f'''
<div style="font-family:Inter,sans-serif;max-width:600px;margin:0 auto">
  <div style="background:linear-gradient(135deg,#f59e0b,#f97316);padding:32px;border-radius:12px 12px 0 0;color:#fff">
    <h1 style="margin:0;font-size:1.5rem">Timesheet Reminder</h1>
    <p style="margin:8px 0 0;opacity:0.85">Please submit your timesheet before end of day today</p>
  </div>
  <div style="background:#fff;padding:24px;border:1px solid #e2e8f0;border-top:none">
    <p>Hi {u['name']},</p>
    <p>This is a friendly reminder to submit your timesheet for the week of <strong>{prev_week_start}</strong>.</p>
    <p>Log your hours worked each day, add your activity notes, and click <strong>Submit for Approval</strong>.</p>
    <div style="margin-top:24px;padding:16px;background:#fffbeb;border-radius:8px;border-left:4px solid #f59e0b">
      <p style="margin:0;font-size:0.875rem">Please submit by <strong>end of day Friday</strong> so your manager can review over the weekend.</p>
    </div>
    <p style="margin-top:16px;font-size:0.8rem;color:#94a3b8">Nexus CRM - Automated weekly reminder</p>
  </div>
</div>'''
            await send_email(u["email"], f"Please submit your timesheet - Week of {prev_week_start}", html)
            logger.info(f"[timesheet-reminder] Sent reminder to {u['email']}")
        except Exception as e:
            logger.error(f"[timesheet-reminder] Failed for {u['email']}: {e}")



# ── Constants ────────────────────────────────────────────────
EXPENSE_CATEGORIES = [
    "payroll", "subscriptions", "infrastructure", "travel",
    "accommodation", "meals", "office_supplies", "marketing", "other",
]
EXPENSE_CURRENCIES = ["EUR", "INR"]
EXPENSE_RECEIPT_MIME_TYPES = {
    "application/pdf": "pdf",
    "image/jpeg":      "jpg",
    "image/jpg":       "jpg",
    "image/png":       "png",
    "image/webp":      "webp",
}
EXPENSE_MAX_BYTES = 15 * 1024 * 1024  # 15 MB
 
# Website application resume folder (created by Ravi)
PUBLIC_RESUME_FOLDER_ID = os.environ.get(
    "PUBLIC_RESUME_FOLDER_ID", "1II9tu-fCUqAs63vWzoiamS_2YnXdE0g2"
)
 
# Simple in-memory FX rate cache (1 hour TTL)
_fx_cache: dict = {}
 
 
# ── Pydantic models ──────────────────────────────────────────
 
class ExpenseUpdate(BaseModel):
    title:                Optional[str]   = None
    vendor:               Optional[str]   = None
    amount:               Optional[float] = None
    currency:             Optional[str]   = None
    category:             Optional[str]   = None
    expense_date:         Optional[str]   = None
    description:          Optional[str]   = None
    worker_name:          Optional[str]   = None
    skill:                Optional[str]   = None
    hours_worked:         Optional[float] = None
    rate_per_hour:        Optional[float] = None
    total_amount_in_words: Optional[str]  = None
    invoice_period_start:  Optional[str]  = None
    invoice_period_end:    Optional[str]  = None
 
 
# ── FX helper (Frankfurter API) ──────────────────────────────
 
async def _get_fx_rate(base: str = "EUR", target: str = "INR", date: str = None) -> tuple:
    """
    Fetch exchange rate from api.frankfurter.dev.

    If date is given (YYYY-MM-DD), returns the historical rate for that exact date.
    If date is None, returns today's latest rate.

    Cache strategy:
    - Historical dates: cached permanently (rate never changes for a past date)
    - Latest rate: 1-hour TTL cache

    Returns (rate: float, date: str).
    Falls back to ~91.5 if API is unreachable.
    """
    # Normalise: use "latest" key for today, exact date for historical
    cache_key = f"{base}_{target}_{date or 'latest'}"
    now = datetime.now(timezone.utc).timestamp()
    cached = _fx_cache.get(cache_key)

    # Historical rates never change — cache permanently
    # Latest rate expires after 1 hour
    if cached:
        is_historical = bool(date)
        if is_historical or (now - cached["ts"] < 3600):
            return cached["rate"], cached["date"]

    url_date = date if date else "latest"
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            r = await client.get(
                f"https://api.frankfurter.dev/v1/{url_date}",
                params={"base": base, "symbols": target},
            )
            r.raise_for_status()
            data = r.json()
        rate = float(data["rates"][target])
        actual_date = data.get("date", date or "")
        _fx_cache[cache_key] = {"rate": rate, "date": actual_date, "ts": now}
        logger.info(f"[FX] {base}→{target} = {rate} (requested={url_date}, actual={actual_date})")
        return rate, actual_date
    except Exception as exc:
        logger.warning(f"[FX] Frankfurter API failed for {url_date}: {exc}. Using cache/fallback.")
        if cached:
            return cached["rate"], cached.get("date", "")
        fallback = 91.5 if base == "EUR" else (1.0 / 91.5)
        return fallback, "fallback"
 
 
# ── Phone validator ──────────────────────────────────────────
 
def _validate_phone(phone: str) -> str:
    """
    Validates international phone number.
    Must start with + and contain 8–15 digits total.
    Strips spaces, dashes, parentheses before validating.
    """
    cleaned = _re.sub(r"[\s\-\(\)\.]", "", phone.strip())
    if not cleaned.startswith("+"):
        raise ValueError(
            "Phone number must start with country code (e.g. +91-9876543210 or +353851234567). "
            "The + prefix is required."
        )
    digit_count = len(_re.sub(r"\D", "", cleaned))
    if digit_count < 8:
        raise ValueError(
            f"Phone number too short ({digit_count} digits). "
            "Please include your full number with country code."
        )
    if digit_count > 15:
        raise ValueError(
            f"Phone number too long ({digit_count} digits). "
            "Maximum 15 digits including country code."
        )
    return cleaned
 
 
# ============================================================
# PUBLIC JOB API  (no authentication required — secured via
#                  per-job apply_key validated server-side)
#
# Security model:
#   • apply_key is a random URL-safe string stored in the jobs table.
#   • It is generated server-side at job creation (secrets.token_urlsafe).
#   • Even if a candidate tampers the URL, a wrong key returns 404 —
#     no CRM data is exposed and no credentials touch the browser.
#   • job_title is NEVER taken from user input; it is always fetched
#     from the database using the validated key.
#   • Rate-limited: 5 submissions per minute per IP (slowapi).
# ============================================================
 
 
@api_router.get("/public/jobs")
async def public_list_jobs(
    search:          Optional[str] = None,
    department:      Optional[str] = None,
    employment_type: Optional[str] = None,
):
    """
    PUBLIC — No authentication required.
    Returns all active job listings for display on jmdatatalent.com/jobs.
 
    Query params (all optional):
      search          — Full-text search in title, department, description
      department      — Filter by department name (exact, case-insensitive)
      employment_type — Filter by type: Full-time | Part-time | Contract | Internship
    """
    result = await run(
        lambda: sb("jobs")
        .select("id,title,department,location,employment_type,description,"
                "requirements,salary_range,skills,is_urgent,created_at,updated_at")
        .eq("is_active", True)
        .order("is_urgent", desc=True)
        .order("created_at", desc=True)
        .execute()
    )
    jobs = result.data or []
 
    if search:
        s = search.lower()
        jobs = [
            j for j in jobs
            if s in (j.get("title") or "").lower()
            or s in (j.get("department") or "").lower()
            or s in (j.get("description") or "").lower()
            or s in " ".join(j.get("skills") or []).lower()
        ]
    if department:
        jobs = [j for j in jobs if (j.get("department") or "").lower() == department.lower()]
    if employment_type:
        jobs = [
            j for j in jobs
            if (j.get("employment_type") or "").lower() == employment_type.lower()
        ]
 
    return {
        "success":    True,
        "count":      len(jobs),
        "jobs":       jobs,
        "generated":  datetime.now(timezone.utc).isoformat(),
    }
 
 
@api_router.get("/public/jobs/{job_id}")
async def public_get_job(job_id: str):
    """
    PUBLIC — No authentication required.
    Returns full details of a single active job.
    Returns 404 if the job does not exist or is not active.
    """
    job = await safe_single(
        lambda: sb("jobs")
        .select("id,title,department,location,employment_type,description,"
                "requirements,salary_range,skills,is_urgent,created_at,updated_at")
        .eq("id", job_id)
        .eq("is_active", True)
        .single()
        .execute()
    )
    if not job:
        raise HTTPException(
            404,
            detail={
                "error":   "JOB_NOT_FOUND",
                "message": "This job posting does not exist or is no longer active.",
                "job_id":  job_id,
            },
        )
    return {"success": True, "job": job}


@api_router.get("/public/jobs/by-key/{apply_key}")
async def public_get_job_by_key(apply_key: str):
    """
    PUBLIC — No authentication required.
    Validates the apply_key and returns safe job details for the application form.
    Returns 404 if the key is wrong, the job is closed, or the job doesn't exist.

    Only returns data the candidate is allowed to see (title, dept, location, type).
    Never exposes internal IDs, credentials, or CRM data.
    """
    apply_key = apply_key.strip()
    if not apply_key or len(apply_key) > 64:
        raise HTTPException(404, detail={"error": "JOB_NOT_FOUND", "message": "Job not found."})

    job = await safe_single(
        lambda: sb("jobs")
        .select("id,title,department,location,employment_type,is_urgent")
        .eq("apply_key", apply_key)
        .eq("is_active", True)
        .single()
        .execute()
    )
    if not job:
        raise HTTPException(
            404,
            detail={
                "error":   "JOB_NOT_FOUND",
                "message": "This job is no longer available or the link is invalid.",
            },
        )
    # Return only what the form needs — internal job UUID is intentionally excluded
    return {
        "success":         True,
        "title":           job["title"],
        "department":      job.get("department") or "",
        "location":        job.get("location") or "",
        "employment_type": job.get("employment_type") or "",
        "is_urgent":       job.get("is_urgent") or False,
    }


@api_router.post("/public/apply")
@limiter.limit("5/minute")          # 5 submissions per IP per minute — prevents spam bots
async def public_apply_job(
    request:          Request,
    # Required fields
    first_name:       str            = Form(..., description="Applicant's first name"),
    last_name:        str            = Form(..., description="Applicant's last name"),
    email:            str            = Form(..., description="Applicant's email address"),
    phone:            str            = Form(..., description="Phone with country code, e.g. +91-9876543210"),
    apply_key:        str            = Form(..., description="Unique job key from the URL (?key=…)"),
    resume:           UploadFile     = File(...,  description="Resume file — PDF, DOC, or DOCX, max 10 MB"),
    # Optional fields
    current_company:  Optional[str]  = Form(None, description="Current employer"),
    candidate_role:   Optional[str]  = Form(None, description="Current job title / role"),
    experience_years: Optional[str]  = Form(None, description="Total years of experience (numeric)"),
    linkedin_url:     Optional[str]  = Form(None, description="LinkedIn profile URL"),
    portfolio_url:    Optional[str]  = Form(None, description="Portfolio, GitHub, or personal site URL"),
):
    """
    PUBLIC — No authentication header required.

    Security: job_id and job_title are NEVER taken from user input.
    The apply_key is validated against the database; the real job data
    is fetched server-side. A tampered key returns 404 — no CRM data exposed.

    Returns 409 if the same email has already applied for the same job.
    Returns 429 if the IP exceeds 5 submissions per minute.
    """

    # ── 0. Compute full name ─────────────────────────────────────
    first_name = first_name.strip()
    last_name  = last_name.strip()
    if not first_name:
        raise HTTPException(422, detail={"error": "MISSING_FIRST_NAME", "message": "First name is required.", "field": "first_name"})
    if not last_name:
        raise HTTPException(422, detail={"error": "MISSING_LAST_NAME", "message": "Last name is required.", "field": "last_name"})
    full_name = f"{first_name} {last_name}"

    # ── 1. Validate apply_key → fetch real job data server-side ──
    #   This is the core security step. job_title is NEVER sourced from user input.
    apply_key = apply_key.strip()
    if not apply_key or len(apply_key) > 64:
        raise HTTPException(404, detail={"error": "JOB_NOT_FOUND", "message": "Invalid application link."})

    job_row = await safe_single(
        lambda: sb("jobs")
        .select("id,title,department")
        .eq("apply_key", apply_key)
        .eq("is_active", True)
        .single()
        .execute()
    )
    if not job_row:
        raise HTTPException(
            404,
            detail={
                "error":   "JOB_NOT_FOUND",
                "message": "This job is no longer available. The position may have been filled.",
            },
        )
    # Real values from the DB — not from the user
    real_job_id    = job_row["id"]
    real_job_title = job_row["title"]

    # ── 2. Validate required text fields ─────────────────────────
    # (No job_id / job_title validation needed — they come from DB)

    # ── 3. Validate email ─────────────────────────────────────────
    email = email.strip().lower()
    if not _re.match(r"^[^\s@]+@[^\s@]+\.[^\s@]+$", email):
        raise HTTPException(
            422,
            detail={
                "error":   "INVALID_EMAIL",
                "message": "Please provide a valid email address (e.g. name@example.com).",
                "field":   "email",
            },
        )

    # ── 4. Validate phone ─────────────────────────────────────────
    try:
        phone_clean = _validate_phone(phone)
    except ValueError as exc:
        raise HTTPException(
            422,
            detail={
                "error":   "INVALID_PHONE",
                "message": str(exc),
                "field":   "phone",
                "example": "+91-9876543210 or +353851234567",
            },
        )

    # ── 5. Validate resume file ───────────────────────────────────
    if resume.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            422,
            detail={
                "error":    "INVALID_FILE_TYPE",
                "message":  f"Resume must be PDF, DOC, or DOCX. Received: {resume.content_type}",
                "field":    "resume",
                "accepted": ["application/pdf",
                             "application/msword",
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
            },
        )
    resume_bytes = await resume.read()
    if len(resume_bytes) > MAX_FILE_BYTES:
        raise HTTPException(
            422,
            detail={
                "error":   "FILE_TOO_LARGE",
                "message": f"Resume must be under 10 MB. Received: {round(len(resume_bytes)/1024/1024, 1)} MB.",
                "field":   "resume",
                "max_mb":  10,
            },
        )

    # ── 6. Duplicate check — same email + real job UUID ──────────
    dup_check = await run(
        lambda: sb("candidates")
        .select("id")
        .eq("email", email)
        .eq("job_id", real_job_id)
        .eq("source", "website")
        .execute()
    )
    if dup_check.data:
        raise HTTPException(
            409,
            detail={
                "error":   "ALREADY_APPLIED",
                "message": "An application with this email already exists for this position. "
                           "Our team will review it and be in touch.",
                "email":   email,
            },
        )

    # ── 7. Upload resume to Google Drive ─────────────────────────
    ext       = ALLOWED_MIME_TYPES[resume.content_type]
    safe_name = _re.sub(r"[^\w\s-]", "", full_name.strip()).replace(" ", "_")[:30]
    short_id  = str(uuid.uuid4()).replace("-", "")[:8]
    filename  = f"WebApp_{safe_name}_{short_id}.{ext}"

    resume_url = None
    if upload_resume is not None:
        try:
            drive_result = await run(
                lambda: upload_resume(
                    resume_bytes,
                    filename,
                    resume.content_type,
                    folder_id=PUBLIC_RESUME_FOLDER_ID,
                )
            )
            resume_url = drive_result["preview_url"]
        except Exception as exc:
            logger.error(f"[public-apply] Drive upload failed: {exc}")
    else:
        logger.warning("[public-apply] Google Drive not configured — resume not uploaded")

    # ── 8. Parse optional numeric fields ─────────────────────────
    exp_years = None
    if experience_years:
        try:
            exp_years = int(float(experience_years.strip()))
            if exp_years < 0 or exp_years > 60:
                exp_years = None
        except (ValueError, AttributeError):
            pass

    # ── 9. Create candidate record ────────────────────────────────
    candidate_payload = {
        "full_name":        full_name,
        "email":            email,
        "phone":            phone_clean,
        "current_company":  (current_company  or "").strip() or None,
        "candidate_role":   (candidate_role   or "").strip() or None,
        "experience_years": exp_years,
        "linkedin_url":     (linkedin_url     or "").strip() or None,
        "portfolio_url":    (portfolio_url    or "").strip() or None,
        "source":           "website",
        "status":           "sourced",
        "job_id":           real_job_id,       # from DB — not from user
        "job_title":        real_job_title,    # from DB — not from user
        "notes":            f"Applied via website for: {real_job_title}",
        "resume_url":       resume_url,
    }

    result = await run(lambda: sb("candidates").insert(candidate_payload).execute())
    if not result.data:
        raise HTTPException(
            500,
            detail={
                "error":   "SERVER_ERROR",
                "message": "Failed to record your application. Please try again or contact us directly.",
            },
        )

    new_candidate = result.data[0]

    # ── 10. Audit log ─────────────────────────────────────────────
    asyncio.create_task(_audit(
        action="create",
        user={"id": None, "email": email, "name": full_name},
        entity_type="candidate",
        entity_id=new_candidate["id"],
        entity_name=f"{full_name} → {real_job_title} (website)",
    ))

    logger.info(
        f"[public-apply] New application: {full_name} <{email}> "
        f"→ {real_job_title} (job_id={real_job_id}) candidate_id={new_candidate['id']}"
    )

    return {
        "success":        True,
        "application_id": new_candidate["id"],
        "message":        "Your application has been submitted successfully. "
                          "Our team will review it and be in touch.",
    }


# ============================================================
# EXPENSE TRACKER
# ============================================================
 
def _require_expense_access(user: dict) -> None:
    """Admin and viewer (CEO) roles only."""
    if user.get("role") not in ("admin", "viewer"):
        raise HTTPException(
            403,
            detail={
                "error":   "ACCESS_DENIED",
                "message": "Expense tracker is restricted to admin and CEO roles.",
            },
        )
 
 
@api_router.get("/expenses/summary")
async def get_expense_summary(
    request: Request,
    year: Optional[int] = None,
):
    """
    Dashboard summary — returns monthly data + category breakdown in BOTH currencies.
    The frontend uses this for the chart + KPI cards + currency toggle (no extra call needed).
    """
    user = await get_current_user(request)
    _require_expense_access(user)
 
    target_year = year or datetime.now(timezone.utc).year
 
    # Fetch today's live rate for the dashboard pill (display reference only)
    try:
        live_eur_to_inr, live_fx_date = await _get_fx_rate("EUR", "INR")
    except Exception:
        live_eur_to_inr, live_fx_date = 91.5, "fallback"

    # Fetch all expenses for target year (+ previous year for YoY comparison)
    raw = await run(
        lambda: sb("expenses")
        .select("id,title,vendor,amount,currency,category,expense_date,"
                "worker_name,skill,source,created_at")
        .gte("expense_date", f"{target_year - 1}-01-01")
        .lte("expense_date", f"{target_year}-12-31")
        .execute()
    )
    rows = raw.data or []

    # Pre-fetch historical FX rate for every unique expense date that has a cross-currency row.
    # Each expense is converted at the rate that was live on its own transaction date — not today's rate.
    # Historical rates are cached permanently (they never change), so this is fast on repeat calls.
    unique_dates = {str(r["expense_date"]) for r in rows if r.get("currency") in ("EUR", "INR")}

    import asyncio as _asyncio
    async def _prefetch_rate(d: str):
        try:
            rate, _ = await _get_fx_rate("EUR", "INR", date=d)
            return d, rate
        except Exception:
            return d, live_eur_to_inr  # fall back to live rate if historical unavailable

    date_rate_pairs = await _asyncio.gather(*[_prefetch_rate(d) for d in unique_dates])
    date_eur_to_inr: dict = dict(date_rate_pairs)  # {"2025-03-15": 89.4, "2025-04-01": 91.2, ...}

    def _rates_for(expense_date: str) -> tuple:
        e2i = date_eur_to_inr.get(str(expense_date), live_eur_to_inr)
        i2e = 1.0 / e2i if e2i else 1.0 / 91.5
        return e2i, i2e

    def to_eur(amount: float, ccy: str, expense_date: str) -> float:
        if ccy == "EUR":
            return float(amount)
        _, i2e = _rates_for(expense_date)
        return round(float(amount) * i2e, 2)

    def to_inr(amount: float, ccy: str, expense_date: str) -> float:
        if ccy == "INR":
            return float(amount)
        e2i, _ = _rates_for(expense_date)
        return round(float(amount) * e2i, 2)
 
    now_utc       = datetime.now(timezone.utc)
    cur_month_pfx = f"{target_year}-{now_utc.month:02d}"
    this_yr_rows  = [r for r in rows if str(r["expense_date"]).startswith(str(target_year))]
    this_mo_rows  = [r for r in this_yr_rows if str(r["expense_date"]).startswith(cur_month_pfx)]
    prev_yr_rows  = [r for r in rows if str(r["expense_date"]).startswith(str(target_year - 1))]
 
    # KPI values
    this_mo_eur   = round(sum(to_eur(r["amount"], r["currency"], r["expense_date"]) for r in this_mo_rows), 2)
    this_mo_inr   = round(sum(to_inr(r["amount"], r["currency"], r["expense_date"]) for r in this_mo_rows), 2)
    ytd_eur       = round(sum(to_eur(r["amount"], r["currency"], r["expense_date"]) for r in this_yr_rows), 2)
    ytd_inr       = round(sum(to_inr(r["amount"], r["currency"], r["expense_date"]) for r in this_yr_rows), 2)
    prev_ytd_eur  = round(sum(to_eur(r["amount"], r["currency"], r["expense_date"]) for r in prev_yr_rows), 2)
 
    subs_mo       = [r for r in this_mo_rows if r["category"] == "subscriptions"]
    subs_mo_eur   = round(sum(to_eur(r["amount"], r["currency"], r["expense_date"]) for r in subs_mo), 2)
    subs_mo_inr   = round(sum(to_inr(r["amount"], r["currency"], r["expense_date"]) for r in subs_mo), 2)
 
    # Monthly breakdown for bar chart
    months_map: dict = {}
    for r in this_yr_rows:
        m = str(r["expense_date"])[:7]  # YYYY-MM
        if m not in months_map:
            months_map[m] = {cat: {"eur": 0.0, "inr": 0.0} for cat in EXPENSE_CATEGORIES}
            months_map[m]["_total"] = {"eur": 0.0, "inr": 0.0, "count": 0}
        cat = r["category"]
        if cat in months_map[m]:
            months_map[m][cat]["eur"] += to_eur(r["amount"], r["currency"], r["expense_date"])
            months_map[m][cat]["inr"] += to_inr(r["amount"], r["currency"], r["expense_date"])
        months_map[m]["_total"]["eur"]   += to_eur(r["amount"], r["currency"], r["expense_date"])
        months_map[m]["_total"]["inr"]   += to_inr(r["amount"], r["currency"], r["expense_date"])
        months_map[m]["_total"]["count"] += 1
 
    # Build all 12 months (fill zeros for missing months)
    monthly_list = []
    for mo in range(1, 13):
        key  = f"{target_year}-{mo:02d}"
        data = months_map.get(key, {})
        entry = {"month": key}
        for cat in EXPENSE_CATEGORIES:
            catdata = data.get(cat, {"eur": 0.0, "inr": 0.0})
            entry[f"{cat}_eur"] = round(catdata["eur"], 2)
            entry[f"{cat}_inr"] = round(catdata["inr"], 2)
        total = data.get("_total", {"eur": 0.0, "inr": 0.0, "count": 0})
        entry["total_eur"] = round(total["eur"], 2)
        entry["total_inr"] = round(total["inr"], 2)
        entry["count"]     = total.get("count", 0)
        monthly_list.append(entry)
 
    # Category breakdown for donut
    cat_map: dict = {}
    for r in this_yr_rows:
        c = r["category"]
        if c not in cat_map:
            cat_map[c] = {"category": c, "eur": 0.0, "inr": 0.0, "count": 0}
        cat_map[c]["eur"]   += to_eur(r["amount"], r["currency"], r["expense_date"])
        cat_map[c]["inr"]   += to_inr(r["amount"], r["currency"], r["expense_date"])
        cat_map[c]["count"] += 1
    for v in cat_map.values():
        v["eur"] = round(v["eur"], 2)
        v["inr"] = round(v["inr"], 2)
    cat_list = sorted(cat_map.values(), key=lambda x: x["eur"], reverse=True)
 
    # Top vendor breakdown
    vendor_map: dict = {}
    for r in this_yr_rows:
        v = (r.get("vendor") or r.get("title") or "Unknown").strip()
        if v not in vendor_map:
            vendor_map[v] = {"vendor": v, "eur": 0.0, "inr": 0.0, "count": 0}
        vendor_map[v]["eur"]   += to_eur(r["amount"], r["currency"], r["expense_date"])
        vendor_map[v]["inr"]   += to_inr(r["amount"], r["currency"], r["expense_date"])
        vendor_map[v]["count"] += 1
    top_vendors = sorted(vendor_map.values(), key=lambda x: x["eur"], reverse=True)[:10]
    for v in top_vendors:
        v["eur"] = round(v["eur"], 2)
        v["inr"] = round(v["inr"], 2)
 
    return {
        "year":              target_year,
        "fx": {
            "eur_to_inr":    round(live_eur_to_inr, 4),
            "inr_to_eur":    round(1.0 / live_eur_to_inr, 6) if live_eur_to_inr else 0,
            "date":          live_fx_date,
            "note":          "Live rate shown for reference. All amounts converted at historical rate for each transaction date.",
        },
        "this_month": {
            "eur":   this_mo_eur,
            "inr":   this_mo_inr,
            "count": len(this_mo_rows),
        },
        "ytd": {
            "eur":       ytd_eur,
            "inr":       ytd_inr,
            "count":     len(this_yr_rows),
            "prev_year": prev_ytd_eur,
            "yoy_pct":   round(((ytd_eur - prev_ytd_eur) / prev_ytd_eur * 100) if prev_ytd_eur else 0, 1),
        },
        "subscriptions_this_month": {
            "eur":   subs_mo_eur,
            "inr":   subs_mo_inr,
            "count": len(subs_mo),
        },
        "monthly_data":  monthly_list,
        "by_category":   cat_list,
        "top_vendors":   top_vendors,
    }
 
 
@api_router.get("/expenses")
async def list_expenses(
    request:  Request,
    month:    Optional[str] = None,   # YYYY-MM
    year:     Optional[int] = None,
    category: Optional[str] = None,
    currency: Optional[str] = None,
    search:   Optional[str] = None,
    limit:    int = 200,
    offset:   int = 0,
):
    """
    List expenses with optional filters.
    Returns expenses in reverse chronological order.
    Only admin and viewer roles can access.
    """
    user = await get_current_user(request)
    _require_expense_access(user)
 
    q = (
        sb("expenses")
        .select("*")
        .order("expense_date", desc=True)
        .order("created_at", desc=True)
    )
 
    if month:
        # YYYY-MM → filter to that calendar month
        try:
            yr, mo = month.split("-")
            import calendar
            last_day = calendar.monthrange(int(yr), int(mo))[1]
            q = q.gte("expense_date", f"{month}-01").lte("expense_date", f"{month}-{last_day:02d}")
        except Exception:
            raise HTTPException(400, "month must be in YYYY-MM format (e.g. 2025-03)")
    elif year:
        q = q.gte("expense_date", f"{year}-01-01").lte("expense_date", f"{year}-12-31")
 
    if category and category in EXPENSE_CATEGORIES:
        q = q.eq("category", category)
    if currency and currency.upper() in EXPENSE_CURRENCIES:
        q = q.eq("currency", currency.upper())
 
    result = await run(lambda: q.limit(limit).offset(offset).execute())
    rows = result.data or []
 
    if search:
        s = search.lower()
        rows = [
            r for r in rows
            if s in (r.get("title")       or "").lower()
            or s in (r.get("vendor")      or "").lower()
            or s in (r.get("description") or "").lower()
            or s in (r.get("worker_name") or "").lower()
            or s in (r.get("skill")       or "").lower()
        ]
 
    return {"success": True, "count": len(rows), "expenses": rows}
 
 
@api_router.post("/expenses")
async def create_expense(
    request:               Request,
    # Core required fields (multipart/form-data)
    title:                 str           = Form(...),
    amount:                float         = Form(...),
    currency:              str           = Form("INR"),
    category:              str           = Form(...),
    expense_date:          str           = Form(...),
    # Optional text fields
    vendor:                Optional[str] = Form(None),
    description:           Optional[str] = Form(None),
    # Payroll-specific (only if category == payroll)
    worker_name:           Optional[str]   = Form(None),
    skill:                 Optional[str]   = Form(None),
    hours_worked:          Optional[float] = Form(None),
    rate_per_hour:         Optional[float] = Form(None),
    total_amount_in_words: Optional[str]   = Form(None),
    invoice_period_start:  Optional[str]   = Form(None),
    invoice_period_end:    Optional[str]   = Form(None),
    # Optional receipt file
    receipt: Optional[UploadFile] = File(None),
):
    """
    Create a new expense record with optional receipt attachment.
    Accepts multipart/form-data.
    Receipt: PDF, JPG, PNG, or WEBP — max 15 MB.
    Only admin and viewer (CEO) can create expenses.
    """
    user = await get_current_user(request)
    _require_expense_access(user)
 
    # ── Validate ──────────────────────────────────────────────
    currency = currency.upper()
    if currency not in EXPENSE_CURRENCIES:
        raise HTTPException(422, f"Currency must be EUR or INR. Got: '{currency}'")
    if category not in EXPENSE_CATEGORIES:
        raise HTTPException(
            422,
            f"Invalid category '{category}'. "
            f"Must be one of: {', '.join(EXPENSE_CATEGORIES)}",
        )
    if amount <= 0:
        raise HTTPException(422, "Amount must be greater than 0.")
    try:
        datetime.strptime(expense_date, "%Y-%m-%d")
    except ValueError:
        raise HTTPException(422, "expense_date must be YYYY-MM-DD (e.g. 2025-03-15)")
 
    # ── Handle receipt upload ─────────────────────────────────
    receipt_url   = None
    drive_file_id = None
 
    if receipt and receipt.filename and receipt.filename.strip():
        if receipt.content_type not in EXPENSE_RECEIPT_MIME_TYPES:
            raise HTTPException(
                422,
                f"Receipt must be PDF, JPG, PNG, or WEBP. "
                f"Received: {receipt.content_type}",
            )
        receipt_bytes = await receipt.read()
        if len(receipt_bytes) > EXPENSE_MAX_BYTES:
            raise HTTPException(
                422,
                f"Receipt too large ({round(len(receipt_bytes)/1024/1024, 1)} MB). "
                "Maximum allowed is 15 MB.",
            )
        if upload_resume is not None:
            try:
                ext        = EXPENSE_RECEIPT_MIME_TYPES[receipt.content_type]
                safe_title = _re.sub(r"[^\w\s-]", "", title)[:20].strip().replace(" ", "_")
                fname      = f"Receipt_{safe_title}_{str(uuid.uuid4())[:8]}.{ext}"
                folder     = os.environ.get("GOOGLE_DRIVE_EXPENSE_FOLDER_ID", "1pkprkI4q1PKOeczdv6ST1iljVj28Evuh").strip() or None
                drive_res  = await run(
                    lambda: upload_resume(receipt_bytes, fname, receipt.content_type, folder_id=folder)
                )
                receipt_url   = drive_res["preview_url"]
                drive_file_id = drive_res["file_id"]
            except RuntimeError as exc:
                raise HTTPException(503, f"Receipt upload failed: {exc}")
            except Exception as exc:
                logger.exception("[expenses] Unexpected receipt upload failure")
                raise HTTPException(500, "Unexpected error uploading receipt. Please try again.")
        else:
            logger.warning("[expenses] Google Drive not configured — receipt not uploaded")
 
    # ── Insert record ─────────────────────────────────────────
    record = {
        "title":                 title.strip(),
        "vendor":                (vendor or "").strip() or None,
        "amount":                float(amount),
        "currency":              currency,
        "category":              category,
        "expense_date":          expense_date,
        "description":           (description or "").strip() or None,
        "worker_name":           (worker_name or "").strip() or None,
        "skill":                 (skill or "").strip() or None,
        "hours_worked":          hours_worked,
        "rate_per_hour":         rate_per_hour,
        "total_amount_in_words": (total_amount_in_words or "").strip() or None,
        "invoice_period_start":  (invoice_period_start or "").strip() or None,
        "invoice_period_end":    (invoice_period_end or "").strip() or None,
        "receipt_url":           receipt_url,
        "drive_file_id":         drive_file_id,
        "source":                "manual",
        "uploaded_by":           user["id"],
    }
 
    result = await run(lambda: sb("expenses").insert(record).execute())
    if not result.data:
        raise HTTPException(500, "Failed to save expense record. Please try again.")
 
    saved = result.data[0]
 
    asyncio.create_task(_audit(
        action="create",
        user=user,
        entity_type="expense",
        entity_id=saved["id"],
        entity_name=f"{title} | {category} | {currency} {amount:.2f}",
        new_value={"category": category, "amount": amount, "currency": currency},
    ))
 
    return {"success": True, "expense": saved}
 
 
@api_router.put("/expenses/{expense_id}")
async def update_expense(
    expense_id: str,
    body:       ExpenseUpdate,
    request:    Request,
):
    """
    Update an expense record (metadata only — receipt cannot be changed here).
    Use DELETE + POST to replace an expense with a different receipt.
    Only admin and viewer (CEO) can update.
    """
    user = await get_current_user(request)
    _require_expense_access(user)
 
    existing = await safe_single(
        lambda: sb("expenses")
        .select("id,title,currency,category,amount")
        .eq("id", expense_id)
        .single()
        .execute()
    )
    if not existing:
        raise HTTPException(404, f"Expense '{expense_id}' not found.")
 
    patch = {k: v for k, v in body.model_dump(exclude_none=True).items() if v is not None}
    if not patch:
        raise HTTPException(400, "No fields provided to update.")
 
    # Validate mutated fields
    if "currency" in patch:
        patch["currency"] = patch["currency"].upper()
        if patch["currency"] not in EXPENSE_CURRENCIES:
            raise HTTPException(422, "Currency must be EUR or INR.")
    if "category" in patch and patch["category"] not in EXPENSE_CATEGORIES:
        raise HTTPException(422, f"Invalid category. Choose from: {', '.join(EXPENSE_CATEGORIES)}")
    if "amount" in patch and patch["amount"] <= 0:
        raise HTTPException(422, "Amount must be greater than 0.")
    if "expense_date" in patch:
        try:
            datetime.strptime(patch["expense_date"], "%Y-%m-%d")
        except ValueError:
            raise HTTPException(422, "expense_date must be YYYY-MM-DD")
 
    result = await run(lambda: sb("expenses").update(patch).eq("id", expense_id).execute())
 
    asyncio.create_task(_audit(
        action="update",
        user=user,
        entity_type="expense",
        entity_id=expense_id,
        entity_name=existing.get("title", expense_id),
        old_value={k: existing.get(k) for k in patch if k in existing},
        new_value=patch,
    ))
 
    return {"success": True, "expense": result.data[0] if result.data else {}}
 
 
@api_router.delete("/expenses/{expense_id}")
async def delete_expense(expense_id: str, request: Request):
    """
    Delete an expense record and its receipt from Google Drive (if any).
    Admin only.
    """
    user = await get_current_user(request)
    if user.get("role") not in ("admin", "viewer"):
        raise HTTPException(403, "Only admin and CEO can delete expense records.")
 
    existing = await safe_single(
        lambda: sb("expenses")
        .select("id,title,receipt_url,drive_file_id")
        .eq("id", expense_id)
        .single()
        .execute()
    )
    if not existing:
        raise HTTPException(404, f"Expense '{expense_id}' not found.")
 
    # Remove receipt from Drive
    receipt_url = existing.get("receipt_url") or ""
    if receipt_url and "drive.google.com" in receipt_url and delete_resume is not None:
        try:
            await run(lambda: delete_resume(receipt_url))
        except Exception as exc:
            logger.warning(f"[expenses] Drive delete failed for {expense_id}: {exc}")
 
    await run(lambda: sb("expenses").delete().eq("id", expense_id).execute())
 
    asyncio.create_task(_audit(
        action="delete",
        user=user,
        entity_type="expense",
        entity_id=expense_id,
        entity_name=existing.get("title", expense_id),
    ))
 
    return {"success": True, "deleted": expense_id, "message": "Expense deleted successfully."}

@app.on_event("startup")
async def startup():
    # Start daily digest scheduler
    digest_time = os.environ.get("DIGEST_TIME", "08:00")
    hour, minute = digest_time.split(":")
    scheduler.add_job(send_daily_digest, CronTrigger(hour=int(hour), minute=int(minute)), id="daily_digest", replace_existing=True)
    scheduler.add_job(cleanup_audit_logs, CronTrigger(hour=3, minute=0), id="audit_cleanup", replace_existing=True)
    scheduler.add_job(send_timesheet_reminder, CronTrigger(day_of_week="fri", hour=17, minute=0), id="timesheet_reminder", replace_existing=True)
    scheduler.start()
    logger.info("[scheduler] Audit log cleanup scheduled at 03:00 daily (keeps 180 days)")
    logger.info(f"[scheduler] Daily digest scheduled at {digest_time}")
    logger.info("[scheduler] Timesheet reminder scheduled at Friday 17:00")

    admin_email    = os.environ.get("ADMIN_EMAIL", "admin@nexuscrm.com")
    admin_password = os.environ.get("ADMIN_PASSWORD", "Admin123!")
    admin_name     = os.environ.get("ADMIN_NAME", "Admin")

    existing = await run(lambda: sb("users").select("id").eq("email", admin_email).execute())
    if not existing.data:
        await run(lambda: sb("users").insert({
            "email":         admin_email,
            "password_hash": hash_password(admin_password),
            "name":          admin_name,
            "role":          "admin",
        }).execute())
        logger.info(f"Admin user created: {admin_email}")
    else:
        # Ensure existing admin always has the admin role (fixes missing role on older installs)
        await run(lambda: sb("users").update({"role": "admin"}).eq("email", admin_email).execute())
        logger.info(f"Admin role confirmed for: {admin_email}")


# ============================================================
# WIRE UP
# ============================================================
@app.api_route("/health", methods=["GET", "HEAD"])
async def root_health():
    return {"status": "ok", "service": "Nexus CRM + ATS", "version": "2.0.0"}

app.include_router(api_router)
app.add_middleware(GZipMiddleware, minimum_size=512)  # compress JSON > 512 bytes
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    # allow_origins=os.environ.get("CORS_ORIGINS", "http://localhost:3000").split(","),
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)
